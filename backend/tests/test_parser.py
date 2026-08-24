"""docx_parser 固定模板解析单元测试。"""
import base64
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches

from app.services.docx_parser import (
    build_import_template,
    is_report_docx,
    parse_any_docx,
    parse_docx,
    parse_report_filename,
)

# 1x1 红色 PNG
_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
    "AAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def _make_docx(path: Path, with_image: bool = True) -> None:
    doc = Document()
    doc.add_heading("渗透测试结果", level=1)
    doc.add_paragraph("以下为本次测试发现的漏洞。")

    rows = [
        ("漏洞名称", "登录接口存在SQL注入"),
        ("漏洞等级", "高危"),
        ("漏洞类型", "SQL注入"),
        ("影响URL", "https://example.com/api/login"),
        ("漏洞描述", "登录接口 username 参数未过滤。"),
        ("复现步骤", "输入 ' or 1=1-- 后返回全部数据。"),
        ("修复建议", "使用参数化查询。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    if with_image:
        cell = table.rows[5].cells[1]  # 复现步骤单元格追加截图
        run = cell.add_paragraph().add_run()
        run.add_picture(BytesIO(_PNG), width=Inches(0.2))

    # 干扰表格：无「漏洞名称」行，应被跳过
    other = doc.add_table(rows=1, cols=2)
    other.rows[0].cells[0].text = "备注"
    other.rows[0].cells[1].text = "非漏洞表格"

    doc.save(str(path))


def test_parse_docx_basic(tmp_path: Path):
    docx_file = tmp_path / "sample.docx"
    _make_docx(docx_file)
    image_dir = tmp_path / "images"

    records = parse_docx(str(docx_file), str(image_dir), "/storage/uploads/test")

    assert len(records) == 1
    rec = records[0]
    assert rec["title"] == "登录接口存在SQL注入"
    assert rec["level"] == 20  # 高危
    assert rec["affected_url"] == "https://example.com/api/login"
    assert "username 参数未过滤" in rec["description_html"]
    assert "参数化查询" in rec["solution_html"]
    assert rec["errors"] == []

    # 图片被提取落盘并替换为 <img> 引用
    assert '<img src="/storage/uploads/test/' in rec["reproduce_html"]
    saved = list(image_dir.glob("*"))
    assert len(saved) == 1
    assert saved[0].read_bytes() == _PNG


def test_parse_docx_unknown_level(tmp_path: Path):
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "漏洞名称"
    table.rows[0].cells[1].text = "弱口令"
    table.rows[1].cells[0].text = "漏洞等级"
    table.rows[1].cells[1].text = "未知等级"
    docx_file = tmp_path / "bad_level.docx"
    doc.save(str(docx_file))

    records = parse_docx(str(docx_file), str(tmp_path / "img"), "/x")
    assert len(records) == 1
    assert records[0]["level"] == 30  # 回落中危
    assert any("无法识别漏洞等级" in e for e in records[0]["errors"])


def test_parse_docx_no_table(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("这是一篇没有表格的文档。")
    docx_file = tmp_path / "empty.docx"
    doc.save(str(docx_file))

    assert parse_docx(str(docx_file), str(tmp_path / "img"), "/x") == []


def test_import_template_roundtrip(tmp_path: Path):
    """官方模板应能被解析器识别。"""
    docx_file = tmp_path / "template.docx"
    build_import_template().save(str(docx_file))

    records = parse_docx(str(docx_file), str(tmp_path / "img"), "/x")
    assert len(records) == 1
    assert records[0]["level"] == 20  # 模板示例为高危
    assert records[0]["errors"] == []


# ---------- 平台报告格式解析 ----------

_SAMPLE_REPORT = Path(__file__).resolve().parents[1] / "storage" / "uploads" / "imports" / "8776123568ad49f885c239c454641a24.docx"


def test_parse_report_filename():
    info = parse_report_filename("20260729综合办公系统渗透测试复测报告.docx")
    assert info == {
        "report_date": "2026-07-29", "system_name": "综合办公系统",
        "is_retest": True, "retest_round_seq": 1,
    }

    info = parse_report_filename("门户系统渗透测试报告.docx")
    assert info["system_name"] == "门户系统"
    assert info["is_retest"] is False
    assert info["report_date"] == ""
    assert info["retest_round_seq"] == 0

    # 同日重复复测带 -N 后缀：-1 表示第二轮复测（round_seq=2）
    info = parse_report_filename("20251011中移系统集成有限公司综合办公系统渗透测试复测报告-1.docx")
    assert info == {
        "report_date": "2025-10-11", "system_name": "综合办公系统",
        "is_retest": True, "retest_round_seq": 2,
    }


def test_normalize_vul_title():
    """标题归一化：剥掉（部分未修复）等修复状态后缀，供跨报告去重合并；fixed 判定严谨。"""
    from app.services.docx_parser import _normalize_vul_title

    # 初测：无后缀，未修复
    assert _normalize_vul_title("越权-劳动合同变更审批") == ("越权-劳动合同变更审批", False)
    # 已修复
    assert _normalize_vul_title("越权-劳动合同变更审批（已修复）") == ("越权-劳动合同变更审批", True)
    # 未修复
    assert _normalize_vul_title("越权-劳动合同变更审批（未修复）") == ("越权-劳动合同变更审批", False)
    # 部分未修复 → 剥后缀且不算已修复
    assert _normalize_vul_title("越权-劳动合同变更审批（部分未修复）") == ("越权-劳动合同变更审批", False)
    # 部分已修复 → 剥后缀且不算已修复（存在未闭环部分）
    assert _normalize_vul_title("越权-劳动合同变更审批（部分已修复）") == ("越权-劳动合同变更审批", False)
    # 基本已修复 → 剥后缀且不算已修复
    assert _normalize_vul_title("越权-劳动合同变更审批（基本已修复）") == ("越权-劳动合同变更审批", False)
    # 半角括号
    assert _normalize_vul_title("越权-劳动合同变更审批(已修复)") == ("越权-劳动合同变更审批", True)
    # 标题含括号但非修复状态 → 不误删
    assert _normalize_vul_title("后台接口存在SQL注入（GET参数）") == ("后台接口存在SQL注入（GET参数）", False)


def test_is_report_docx_on_template(tmp_path: Path):
    """固定模板不应被误判为报告格式。"""
    docx_file = tmp_path / "template.docx"
    build_import_template().save(str(docx_file))
    assert is_report_docx(Document(str(docx_file))) is False


def test_parse_schedule_table():
    """「时间与人员」表：提取测试周期与参测人员；普通文档不误解析。"""
    from app.services.docx_parser import _parse_schedule_table

    doc = Document()
    t = doc.add_table(rows=7, cols=4)
    t.rows[0].cells[0].text = "测试工作时间段"
    t.rows[1].cells[0].text, t.rows[1].cells[1].text = "起始时间", "2026-06-30"
    t.rows[1].cells[2].text, t.rows[1].cells[3].text = "结束时间", "2026-07-01"
    for i, h in enumerate(("参测人员", "所属部门", "人员角色", "人员分工")):
        t.rows[3].cells[i].text = h
    for i, name in enumerate(("邢博宇", "许宁安", "薛田泽")):
        t.rows[4 + i].cells[0].text = name
    assert _parse_schedule_table(doc) == {
        "test_start": "2026-06-30",
        "test_end": "2026-07-01",
        "testers": ["邢博宇", "许宁安", "薛田泽"],
    }

    # 无「时间与人员」表的普通文档返回空值
    plain = Document()
    plain.add_paragraph("普通文档")
    assert _parse_schedule_table(plain) == {"test_start": "", "test_end": "", "testers": []}


def test_parse_target_table():
    """「测试目标」表：解析系统名 / 被测URL / 被测IP / 被测测试账号。"""
    from app.services.docx_parser import _parse_target_table

    doc = Document()
    t = doc.add_table(rows=5, cols=2)
    rows = (
        ("业务系统名称", "综合办公系统"),
        ("被测系统URL", "https://oa.example.com"),
        ("被测系统域名", "oa.example.com"),
        ("被测系统IP", "10.0.0.1"),
        ("被测测试账号", "admin/Admin@123"),
    )
    for i, (label, value) in enumerate(rows):
        t.rows[i].cells[0].text, t.rows[i].cells[1].text = label, value
    assert _parse_target_table(doc) == {
        "system_name": "综合办公系统",
        "target_url": "https://oa.example.com",
        "target_ip": "10.0.0.1",
        "test_account": "admin/Admin@123",
    }


def test_parse_report_docx_sample(tmp_path: Path):
    """样例复测报告：meta 与漏洞记录解析。"""
    if not _SAMPLE_REPORT.exists():
        import pytest
        pytest.skip("样例报告文件不存在")

    assert is_report_docx(Document(str(_SAMPLE_REPORT))) is True

    kind, meta, records = parse_any_docx(
        str(_SAMPLE_REPORT), str(tmp_path / "img"), "/x",
        "20260729综合办公系统渗透测试复测报告.docx",
    )
    assert kind == "report"
    assert meta["system_name"] == "综合办公系统"
    assert meta["report_date"] == "2026-07-29"
    assert meta["is_retest"] is True
    assert meta["target_ip"] == "10.55.133.37"

    assert len(records) == 1
    rec = records[0]
    assert rec["title"].startswith("平行越权")
    assert "已修复" not in rec["title"]  # 标题尾缀被剔除
    assert rec["level"] == 20  # 汇总表：高危
    assert rec["fixed"] is True
    assert rec["affected_url"].startswith("http://10.55.133.37")
    assert rec["description_html"]
    assert rec["retest_html"]  # 「20260729漏洞复测：」段落
    assert rec["errors"] == []
