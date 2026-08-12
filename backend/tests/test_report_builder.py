"""report_builder 与报告 HTML 助手的纯函数单元测试（不依赖数据库 / 队列）。

覆盖计划测试项：
- 汇总统计仅计入未修复漏洞、已修复状态单元格绿色
- 详情「测试状态：」按漏洞最新 is_retest 重写
- settings 中 w:updateFields 位于 w:hdrShapeDefaults 之前
- 封面第二行为 project_name（系统名称）
- _affected_urls_html / _vuln_section_html 多 URL 渲染
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph

from PIL import Image

from app.api.v1.reports import _affected_urls_html, _vuln_section_html
from app.constants import VulStatus
from app.core.config import settings
from app.core.timeutil import now as tznow
from app.services.report_builder import (
    _localize_images,
    FIXED_STATUS_COLOR,
    build_report_docx,
)


def _build(tmp_path, meta=None, vulns=None, sections=None):
    meta = {"title": "报告标题", "project_name": "统一门户系统", **(meta or {})}
    out = str(tmp_path / "out.docx")
    build_report_docx(meta, vulns or [], sections or [], out)
    return Document(out)


def _code_block_paras(doc: Document):
    """带 pBdr 边框的代码块段落。"""
    return [
        p for p in doc.paragraphs
        if (pPr := p._p.find(qn("w:pPr"))) is not None
        and pPr.find(qn("w:pBdr")) is not None
    ]


def _shape_in_table(shape) -> bool:
    node = shape._inline
    while node is not None:
        if node.tag == qn("w:tbl"):
            return True
        node = node.getparent()
    return False


def _cell_run_colors(cell):
    return [run.font.color.rgb for p in cell.paragraphs for run in p.runs]


def test_affected_urls_html_multi_value():
    # 多 URL 以换行分隔存储，逐条转义后以 <br/> 连接
    html = _affected_urls_html("https://a.com/x\nhttps://b.com/y?a=1&b=2\n")
    assert html == "https://a.com/x<br/>https://b.com/y?a=1&amp;b=2"
    # 空值回退占位符
    assert _affected_urls_html("") == "-"
    assert _affected_urls_html(None) == "-"


def test_vuln_section_html_multi_url():
    class _V:
        is_retest = False
        level = 20
        affected_url = "https://a.com/1\nhttps://a.com/2"
        description_html = ""
        reproduce_html = ""
        solution_html = ""
        retest_html = ""

    html = _vuln_section_html(_V())
    assert "漏洞链接：</strong>https://a.com/1<br/>https://a.com/2" in html


def test_summary_counts_exclude_fixed_and_green_cell(tmp_path):
    vulns = [
        {"id": 1, "title": "高危未修复", "level": 20, "vul_type": 10,
         "status": VulStatus.UNFIXED, "is_retest": False},
        {"id": 2, "title": "高危已修复", "level": 20, "vul_type": 10,
         "status": VulStatus.FIXED, "is_retest": False},
    ]
    doc = _build(tmp_path, vulns=vulns)

    # 统计段：高危漏洞仅计未修复 = 1 个
    para_texts = [p.text.strip() for p in doc.paragraphs]
    assert any(t.startswith("高危漏洞1个") for t in para_texts), para_texts

    # 汇总表列出全部漏洞（表头 + 2 行）
    table = doc.tables[6]
    assert len(table.rows) == 3
    # 已修复行状态列为绿色，未修复行状态列非绿色
    fixed_row = next(r for r in table.rows[1:] if r.cells[2].text.strip() == "高危已修复")
    unfixed_row = next(r for r in table.rows[1:] if r.cells[2].text.strip() == "高危未修复")
    assert FIXED_STATUS_COLOR in _cell_run_colors(fixed_row.cells[3])
    assert FIXED_STATUS_COLOR not in _cell_run_colors(unfixed_row.cells[3])


def test_detail_test_state_rewritten_by_is_retest(tmp_path):
    # 章节快照写「初测」，但漏洞已复测，导出应改写为「复测」
    vulns = [{"id": 1, "title": "复测漏洞", "level": 20, "vul_type": 10,
              "status": VulStatus.FIXING, "is_retest": True}]
    sections = [{"title": "复测漏洞", "vul_id": 1,
                 "content_html": "<p><strong>测试状态：</strong>初测</p>"}]
    doc = _build(tmp_path, vulns=vulns, sections=sections)
    body = [p.text for p in doc.paragraphs]
    assert any("测试状态：复测" in t for t in body), body
    # 修复中 + 复测 = 复测未通过（标题后缀）
    h3 = [p.text for p in doc.paragraphs if p.style.name == "Heading 3"]
    assert any(t.endswith("（复测未通过）") for t in h3), h3


def test_cover_second_line_is_project_name(tmp_path):
    doc = _build(tmp_path, meta={"project_name": "网上营业厅"})
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "网上营业厅" in texts


def test_update_fields_precedes_compat_anchor(tmp_path):
    doc = _build(tmp_path)
    settings_el = doc.settings.element
    children = list(settings_el)
    update = settings_el.find(qn("w:updateFields"))
    assert update is not None
    # updateFields 必须排在 hdrShapeDefaults / compat / rsids 等锚点之前
    anchor_tags = {qn(t) for t in (
        "w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr",
        "w:compat", "w:rsids", "m:mathPr", "w:themeFontLang",
    )}
    anchors = [i for i, c in enumerate(children) if c.tag in anchor_tags]
    if anchors:
        assert children.index(update) < min(anchors)


def test_code_block_width_14cm_and_left_aligned(tmp_path):
    # 代码块显示宽度统一 14cm（左右缩进均分、块体居中），内容保持左对齐
    sections = [{
        "title": "代码块章节",
        "vul_id": None,
        "content_html": "<pre><code>print('hello')\nfor i in range(3):\n    print(i)</code></pre>",
    }]
    doc = _build(tmp_path, sections=sections)
    paras = _code_block_paras(doc)
    assert paras, "导出文档中未找到代码块段落"
    section = doc.sections[-1]
    content_w = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    target = int(Cm(14))
    for para in paras:
        left = int(para.paragraph_format.left_indent)
        right = int(para.paragraph_format.right_indent)
        # 左右缩进对称（水平居中）
        assert abs(left - right) <= 635, f"左右缩进不对称 {left}/{right} EMU"
        width = content_w - left - right
        # docx 缩进按 twip（1twip=635EMU）存储，允许 2 twip 舍入误差
        assert abs(width - target) <= 1270, f"代码块宽度 {width} EMU，目标 {target} EMU"
        # 块体居中靠对称缩进实现，内容不居中（保持左对齐）
        assert para.alignment != WD_ALIGN_PARAGRAPH.CENTER


def test_code_block_line_spacing_kept(tmp_path):
    # 章节行距统一 1.5 时跳过代码块，保持紧凑 1.0 行距
    sections = [{
        "title": "代码",
        "vul_id": None,
        "content_html": "<pre>line1\nline2</pre>",
    }]
    doc = _build(tmp_path, sections=sections)
    paras = _code_block_paras(doc)
    assert paras
    for para in paras:
        assert para.paragraph_format.line_spacing == 1.0


def test_body_image_centered_and_width_14cm(tmp_path):
    # 正文图片统一 14cm 宽且所在段落水平居中（跳过封面表格内图片）
    img = tmp_path / "shot.png"
    Image.new("RGB", (800, 600), "white").save(img)
    sections = [{
        "title": "图片章节",
        "vul_id": None,
        "content_html": f'<p><img src="{img}"></p>',
    }]
    doc = _build(tmp_path, sections=sections)
    body_shapes = [s for s in doc.inline_shapes if not _shape_in_table(s)]
    assert body_shapes, "未找到正文图片"
    for shape in body_shapes:
        node = shape._inline
        while node is not None and node.tag != qn("w:p"):
            node = node.getparent()
        assert Paragraph(node, doc).alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert abs(int(shape.width) - int(Cm(14))) <= 1, f"图片宽度 {shape.width} EMU，目标 {int(Cm(14))} EMU"


# ---------- 复测封面标题 ----------
def test_cover_retest_title(tmp_path):
    # 复测报告：首页第三行标题自动变更为「渗透测试复测报告」
    doc = _build(tmp_path, meta={"project_name": "统一门户系统", "is_retest": True})
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "渗透测试复测报告" in texts
    assert "渗透测试报告" not in texts


def test_cover_retest_title_with_customer(tmp_path):
    # 设置了客户时不得提前中断封面遍历：复测标题与系统名均须替换（回归：曾因 break 跳过后续占位）
    doc = _build(tmp_path, meta={
        "project_name": "统一门户系统", "customer": "测试客户", "is_retest": True,
    })
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "测试客户" in texts
    assert "统一门户系统" in texts
    assert "渗透测试复测报告" in texts
    assert "标准名称（邮件、台账）-系统名称（网页）" not in texts


def test_cover_normal_title_unchanged(tmp_path):
    # 初测报告：封面标题保持「渗透测试报告」
    doc = _build(tmp_path, meta={"project_name": "统一门户系统"})
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "渗透测试报告" in texts
    assert "渗透测试复测报告" not in texts


# ---------- 版本变更记录：时间=报告自身日期、修改人=发起导出账号 ----------
def test_version_records_use_export_time_and_generator(tmp_path):
    meta = {
        "project_name": "统一门户系统",
        "author": "报告作者甲",
        "generator": "导出账号乙",
        "report_records": [
            {"is_retest": False, "creator_name": "报告作者甲", "date": ""},
            {"is_retest": True, "creator_name": "轮次创建人", "date": ""},
        ],
    }
    doc = _build(tmp_path, meta=meta)
    table = doc.tables[1]
    rows = table.rows[2:4]  # 表头 r0/r1，数据行 V1.0 + V2.0
    dates = [r.cells[0].text.strip() for r in rows]
    modifiers = [r.cells[4].text.strip() for r in rows]
    today = tznow().strftime("%Y-%m-%d")
    assert all(d == today for d in dates), dates
    assert all(m == "导出账号乙" for m in modifiers), modifiers


def test_version_records_each_version_own_export_date(tmp_path):
    # 各版记录时间取各自对应报告自身的日期，不得被最后一次导出时间覆盖
    meta = {
        "project_name": "统一门户系统",
        "author": "报告作者甲",
        "generator": "导出账号乙",
        "report_records": [
            {"is_retest": False, "creator_name": "报告作者甲", "date": "2026-08-01"},
            {"is_retest": True, "creator_name": "轮次创建人", "date": "2026-08-07"},
        ],
    }
    doc = _build(tmp_path, meta=meta)
    table = doc.tables[1]
    rows = table.rows[2:4]  # V1.0 + V2.0
    dates = [r.cells[0].text.strip() for r in rows]
    versions = [r.cells[1].text.strip() for r in rows]
    assert dates == ["2026-08-01", "2026-08-07"], dates
    assert versions == ["V1.0", "V2.0"], versions


def test_version_records_fallback_when_dates_empty(tmp_path):
    # 报告 date 缺失（空串）：回退当前导出时间（today）
    meta = {
        "project_name": "统一门户系统",
        "author": "报告作者甲",
        "report_records": [
            {"is_retest": False, "creator_name": "报告作者甲", "date": "2026-08-01"},
            {"is_retest": True, "creator_name": "轮次创建人", "date": ""},
        ],
    }
    doc = _build(tmp_path, meta=meta)
    table = doc.tables[1]
    rows = table.rows[2:4]
    dates = [r.cells[0].text.strip() for r in rows]
    today = tznow().strftime("%Y-%m-%d")
    assert dates[0] == "2026-08-01"
    assert dates[1] == today, dates


def test_version_records_count_matches_retest_reports(tmp_path):
    # 回归：版本号严格按复测报告数量输出，与复测轮次/幽灵数据无关。
    # 仅 2 份复测报告（同创建人）→ V1.0 / V2.0 / V2.1，不得虚增到 V2.5
    meta = {
        "project_name": "统一门户系统",
        "report_records": [
            {"is_retest": False, "creator_name": "甲", "date": "2026-07-21"},
            {"is_retest": True, "creator_name": "乙", "date": "2026-07-28"},
            {"is_retest": True, "creator_name": "乙", "date": "2026-08-04"},
        ],
    }
    doc = _build(tmp_path, meta=meta)
    table = doc.tables[1]
    versions = [r.cells[1].text.strip() for r in table.rows[2:5]]
    notes = [r.cells[2].text.strip() for r in table.rows[2:5]]
    assert versions == ["V1.0", "V2.0", "V2.1"], versions
    assert notes == ["初测创建", "复测创建", "复测更新"], notes


def test_version_records_creator_change_bumps_major(tmp_path):
    # 复测报告创建人与上一份不同 → 主版本 +1（V3.0）
    meta = {
        "project_name": "统一门户系统",
        "report_records": [
            {"is_retest": False, "creator_name": "甲", "date": ""},
            {"is_retest": True, "creator_name": "乙", "date": ""},
            {"is_retest": True, "creator_name": "丙", "date": ""},
        ],
    }
    doc = _build(tmp_path, meta=meta)
    table = doc.tables[1]
    versions = [r.cells[1].text.strip() for r in table.rows[2:5]]
    assert versions == ["V1.0", "V2.0", "V3.0"], versions


# ---------- 测试时间与人员：参测人员首行取发起导出账号，其余取报告作者 ----------
def test_schedule_table_first_row_generator(tmp_path):
    # 发起人不在作者中：首行取 generator，其余按作者顺序
    doc = _build(tmp_path, meta={"project_name": "X", "author": "张三、李四", "generator": "王博宇"})
    table = doc.tables[5]
    names = [r.cells[0].text.strip() for r in table.rows[4:] if r.cells[0].text.strip()]
    assert names == ["王博宇", "张三", "李四"]


def test_schedule_table_generator_in_author_no_dup(tmp_path):
    # 发起人与作者重复：不重复，发起人仍为首行
    doc = _build(tmp_path, meta={"project_name": "X", "author": "王博宇、李四", "generator": "王博宇"})
    table = doc.tables[5]
    names = [r.cells[0].text.strip() for r in table.rows[4:] if r.cells[0].text.strip()]
    assert names == ["王博宇", "李四"]


def test_schedule_table_expand_rows_beyond_3(tmp_path):
    # 发起人 + 3 作者 = 4 人，模板仅 3 行，应自动克隆新增 1 行
    doc = _build(tmp_path, meta={"project_name": "X", "author": "张三、李四、王五", "generator": "赵六"})
    table = doc.tables[5]
    names = [r.cells[0].text.strip() for r in table.rows[4:] if r.cells[0].text.strip()]
    assert names == ["赵六", "张三", "李四", "王五"]
    # 发起人缺失时首行取报告作者
    doc2 = _build(tmp_path, meta={"project_name": "X", "author": "张三、李四"})
    table2 = doc2.tables[5]
    names2 = [r.cells[0].text.strip() for r in table2.rows[4:] if r.cells[0].text.strip()]
    assert names2 == ["张三", "李四"]


# ---------- 无漏洞报告：风险问题汇总以 \ 填充 ----------
def test_no_vuln_summary_backslash_fill(tmp_path):
    doc = _build(tmp_path, meta={"project_name": "X"})
    table = doc.tables[6]
    row = table.rows[1]
    assert [c.text for c in row.cells] == ["\\", "\\", "\\", "\\"]


# ---------- 图片导出：缺失图片不产生 <image:> 链接占位 ----------
def test_missing_image_removed_not_placeholder(tmp_path):
    sections = [{
        "title": "缺失图片章节",
        "vul_id": None,
        "content_html": '<p>正文内容<img src="/storage/uploads/images/not_exist_xyz.png"></p>',
    }]
    doc = _build(tmp_path, sections=sections)
    body = [p.text for p in doc.paragraphs]
    assert not any(t.strip().startswith("<image:") for t in body), body
    assert any("正文内容" in t for t in body), body


# ---------- 图片导出回归：/storage/ 真实图片必须本地化并嵌入 ----------
def test_storage_image_localized_and_embedded(tmp_path, monkeypatch):
    """回归：v1.1.0 _localize_images 取错正则分组（group(1) 是引号）且 repl 返回 src="..."
    导致 /storage/ 图片全部丢失。真实存在的 storage 图片必须本地化并内嵌到 docx。"""
    monkeypatch.setattr(settings, "STORAGE_DIR", str(tmp_path))
    img_dir = tmp_path / "uploads" / "images"
    img_dir.mkdir(parents=True)
    Image.new("RGB", (800, 600), "white").save(img_dir / "ab12cd34.png")
    sections = [{
        "title": "图片章节",
        "vul_id": None,
        "content_html": '<p>正文<img src="/storage/uploads/images/ab12cd34.png"></p>',
    }]
    doc = _build(tmp_path, sections=sections)
    body_shapes = [s for s in doc.inline_shapes if not _shape_in_table(s)]
    assert body_shapes, "未找到正文图片：/storage/ 图片未被本地化嵌入 docx"


def test_localize_images_keeps_img_tag_and_matches_real_file(tmp_path, monkeypatch):
    """回归：本地化后必须保留完整 <img> 标签（此前 repl 返回 src="..." 丢 <img 与 >），
    且 src 指向 storage 下真实存在的文件。"""
    monkeypatch.setattr(settings, "STORAGE_DIR", str(tmp_path))
    img_dir = tmp_path / "uploads" / "images"
    img_dir.mkdir(parents=True)
    Image.new("RGB", (10, 10), "white").save(img_dir / "a1b2c3.png")
    out = _localize_images('<img src="/storage/uploads/images/a1b2c3.png">')
    assert out.startswith("<img "), out
    assert out.endswith(">"), out
    local_src = out[out.find('src="') + 5:out.rfind('"')]
    assert Path(local_src).is_file(), out
