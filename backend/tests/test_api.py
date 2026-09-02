"""API 集成测试：认证 → 业务 CRUD → 状态机 → Word 导入 → 报告与导出全链路。"""
import asyncio
import json
import re
from io import BytesIO

import pytest
from httpx import AsyncClient

pytestmark = pytest.mark.asyncio(loop_scope="session")


async def test_health(client: AsyncClient):
    resp = await client.get("/api/health")
    assert resp.status_code == 200


async def test_login_and_me(client: AsyncClient, auth: dict):
    resp = await client.get("/api/v1/auth/me", headers=auth)
    assert resp.status_code == 200
    body = resp.json()
    assert body["username"] == "admin"

    bad = await client.post(
        "/api/v1/auth/login", data={"username": "admin", "password": "wrong"}
    )
    assert bad.status_code == 401


async def test_meta(client: AsyncClient, auth: dict):
    resp = await client.get("/api/v1/meta", headers=auth)
    assert resp.status_code == 200
    meta = resp.json()
    assert "vul_level" in meta and "vul_status" in meta
    # 字典单源：色值命名空间、导入/导出状态与 nonpen 命名空间随 /meta 下发
    assert meta["colors"]["vul_level"]["10"] == "#DC2626"
    assert meta["import_batch_status"]["parsed"] == "待确认"
    assert meta["export_job_status"]["done"] == "已完成"
    assert meta["report_status"]["draft"] == "草稿"
    # 名称字典对应的色值必须同步下发：前端 applyDictMeta 无条件注入，
    # 任一 key 缺失会令导出/导入状态标签渲染崩溃（报告区域整体消失）
    assert meta["colors"]["import_batch_status"]["parsed"] == "#0284C7"
    assert meta["colors"]["import_record_status"]["confirmed"] == "#059669"
    assert meta["colors"]["export_job_status"]["done"] == "#059669"
    nonpen_items = {item["key"] for item in meta["nonpen"]["items"]}
    assert nonpen_items == {"baseline", "host", "web"}
    assert meta["nonpen"]["actions"]["not_started"] == ["start", "ignore"]
    assert meta["nonpen"]["action_names"]["start"] == "开始初测"


async def test_asset_and_vuln_lifecycle(client: AsyncClient, auth: dict):
    # 建资产（系统级：负责人 / URL 标签 / 端口服务对 / 中间件与数据库多条目）
    resp = await client.post(
        "/api/v1/assets", headers=auth,
        json={
            "name": "测试商城",
            "sub_system": "订单中心",
            "department": "电商事业部",
            "public_urls": [{"url": "https://shop.example.com", "tag": 10}],
            "internal_urls": ["http://10.0.0.8:8080"],
            "port_services": [
                {"port": "80", "service": "Web服务"},
                {"port": "443", "service": "HTTPS"},
            ],
            "middlewares": [{"name": "Nginx", "version": "1.24"}],
            "databases": [{"name": "MySQL", "version": "8.0"}],
            "owners": [{"name": "张三", "phone": "13800000000", "email": "zhangsan@example.com"}],
        },
    )
    assert resp.status_code == 200, resp.text
    asset = resp.json()
    asset_id = asset["id"]
    assert asset["owners"][0]["name"] == "张三"
    assert asset["public_urls"][0]["tag"] == 10
    assert asset["port_services"][1] == {"port": "443", "service": "HTTPS"}
    assert asset["middlewares"][0]["version"] == "1.24"
    assert asset["databases"][0] == {"name": "MySQL", "version": "8.0"}

    # 系统命名必填
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": ""})
    assert resp.status_code == 422

    # 建漏洞（未修复 10），多对多关联资产
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={
            "title": "订单接口越权", "level": 20, "vul_type": 30,
            "affected_url": "https://shop.example.com/api/order",
            "description_html": "<p>横向越权读取他人订单</p>",
            "asset_ids": [asset_id],
        },
    )
    assert resp.status_code == 200, resp.text
    vul = resp.json()
    assert vul["status"] == 10
    assert vul["asset_ids"] == [asset_id]
    assert vul["assets"][0]["name"] == "测试商城"
    vul_id = vul["id"]

    # 按资产筛选漏洞
    resp = await client.get("/api/v1/vulns", headers=auth, params={"asset_id": asset_id})
    assert vul_id in [v["id"] for v in resp.json()["items"]]

    # 已关联漏洞的资产不能删除
    resp = await client.delete(f"/api/v1/assets/{asset_id}", headers=auth)
    assert resp.status_code == 400

    # 未修复可流转到：已忽略/暂不处理/修复中/复测中（复测中为修复状态冗余兜底路径）
    resp = await client.get(f"/api/v1/vulns/{vul_id}/transitions", headers=auth)
    assert {t["status"] for t in resp.json()} == {20, 35, 50, 55}

    # 10 -> 50 进入修复中，打 notice_time
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/transition", headers=auth,
        json={"status": 50, "comment": "已出报告"},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["status"] == 50
    assert resp.json()["notice_time"] is not None

    # 非法流转 50 -> 60 应被拒绝（必须经过复测）
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/transition", headers=auth,
        json={"status": 60},
    )
    assert resp.status_code == 400

    # 50 -> 55 复测中 -> 60 已修复（终态，须携带复测内容）
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": 55},
    )
    assert resp.status_code == 200
    assert resp.json()["is_retest"] is True
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/transition", headers=auth,
        json={"status": 60, "comment": "复测通过", "retest_html": "<p>复测通过</p>"},
    )
    assert resp.status_code == 200
    assert resp.json()["fix_time"] is not None
    # 已修复为可回退状态：可重新打开为未修复(10)或直接重新复测(55)
    resp = await client.get(f"/api/v1/vulns/{vul_id}/transitions", headers=auth)
    assert {t["status"] for t in resp.json()} == {10, 55}

    # 日志包含创建与状态流转
    resp = await client.get(f"/api/v1/vulns/{vul_id}/logs", headers=auth)
    actions = [log["action"] for log in resp.json()]
    assert "创建漏洞" in actions
    assert len(actions) >= 2


async def test_vuln_batch_create(client: AsyncClient, auth: dict):
    resp = await client.post(
        "/api/v1/assets", headers=auth, json={"name": "批量目标系统"},
    )
    asset_id = resp.json()["id"]

    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={
            "asset_ids": [asset_id],
            "vulns": [
                {"title": "批量漏洞A", "level": 20, "vul_type": 10},
                {"title": "批量漏洞B", "level": 30, "vul_type": 15},
                {"title": "批量漏洞C", "level": 40, "vul_type": 55},
            ],
        },
    )
    assert resp.status_code == 200, resp.text
    vulns = resp.json()
    assert len(vulns) == 3
    assert all(v["asset_ids"] == [asset_id] for v in vulns)

    # 无效资产ID被拒绝
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [999999], "vulns": [{"title": "无效资产"}]},
    )
    assert resp.status_code == 400


async def test_asset_excel_import_export(client: AsyncClient, auth: dict):
    from openpyxl import Workbook, load_workbook

    # 构造导入文件：1 行合法 + 1 行缺系统命名
    wb = Workbook()
    ws = wb.active
    # 表头与模板一致（assets.py EXCEL_HEADERS，含「系统类型」列）
    ws.append(["系统命名*", "子系统名称", "部门", "系统类型", "公网URL", "内网URL",
               "开放端口与服务", "中间件", "数据库", "系统负责人", "状态", "备注"])
    ws.append(["Excel导入系统", "支付子系统", "金融部", "自有系统",
               "https://pay.example.com|互联网;https://oa.example.com|办公网",
               "http://192.168.1.10", "443:支付服务;8443:管理后台", "Tomcat/9.0", "Oracle/19c",
               "李四/13900000000/lisi@example.com;王五//wangwu@example.com",
               "线上", "备注内容"])
    ws.append(["", "无名子系统", "", "", "", "", "", "", "", "", "", ""])
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    resp = await client.post(
        "/api/v1/assets/import", headers=auth,
        files={"file": ("assets.xlsx", buf,
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert resp.status_code == 200, resp.text
    result = resp.json()
    assert result["total"] == 2
    assert result["success"] == 1
    assert result["failed"] == 1
    assert "系统命名为必填项" in result["errors"][0]

    # 非 xlsx 拒绝
    resp = await client.post(
        "/api/v1/assets/import", headers=auth,
        files={"file": ("bad.txt", b"hello", "text/plain")},
    )
    assert resp.status_code == 400

    # 导入结果校验（多值字段解析）
    resp = await client.get("/api/v1/assets", headers=auth, params={"search": "Excel导入系统"})
    items = resp.json()["items"]
    assert len(items) == 1
    imported = items[0]
    assert imported["public_urls"] == [
        {"url": "https://pay.example.com", "tag": 10},
        {"url": "https://oa.example.com", "tag": 20},
    ]
    assert imported["port_services"] == [
        {"port": "443", "service": "支付服务"},
        {"port": "8443", "service": "管理后台"},
    ]
    assert imported["middlewares"] == [{"name": "Tomcat", "version": "9.0"}]
    assert imported["databases"] == [{"name": "Oracle", "version": "19c"}]
    assert len(imported["owners"]) == 2
    assert imported["owners"][1] == {"name": "王五", "phone": "", "email": "wangwu@example.com"}

    # 导出并回读
    resp = await client.get("/api/v1/assets/export", headers=auth, params={"search": "Excel导入系统"})
    assert resp.status_code == 200
    wb = load_workbook(BytesIO(resp.content))
    rows = list(wb.active.iter_rows(values_only=True))
    assert len(rows) == 2  # 表头 + 1 行数据
    assert rows[1][0] == "Excel导入系统"
    assert rows[1][3] == "自有系统"  # 系统类型列
    assert "互联网" in rows[1][4]  # 公网URL列
    assert rows[1][6] == "443:支付服务;8443:管理后台"
    assert rows[1][7] == "Tomcat/9.0"

    # 模板下载
    resp = await client.get("/api/v1/assets/import/template", headers=auth)
    assert resp.status_code == 200
    assert resp.content[:2] == b"PK"


async def _wait_batch(client: AsyncClient, auth: dict, batch_id: int) -> dict:
    for _ in range(50):
        resp = await client.get(f"/api/v1/imports/{batch_id}", headers=auth)
        assert resp.status_code == 200
        detail = resp.json()
        if detail["batch"]["status"] in ("parsed", "failed", "confirmed"):
            return detail
        await asyncio.sleep(0.2)
    raise AssertionError("导入批次解析超时")


async def test_word_import_flow(client: AsyncClient, auth: dict):
    # 下载官方模板作为上传样例
    resp = await client.get("/api/v1/imports/template", headers=auth)
    assert resp.status_code == 200
    docx_bytes = resp.content

    # 非 docx 拒绝
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": ("bad.txt", b"hello", "text/plain")},
    )
    assert resp.status_code == 400

    # 上传并等待后台解析（DISABLE_QUEUE 下进程内执行）
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": ("样例报告.docx", BytesIO(docx_bytes),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    batch_id = resp.json()["id"]

    detail = await _wait_batch(client, auth, batch_id)
    assert detail["batch"]["status"] == "parsed", detail
    records = detail["records"]
    assert len(records) == 1
    rec = records[0]
    assert rec["status"] == "parsed"
    assert rec["level"] == 20

    # 预览修正标题后确认入库
    resp = await client.put(
        f"/api/v1/imports/records/{rec['id']}", headers=auth,
        json={"title": "后台登录接口存在SQL注入(已确认)"},
    )
    assert resp.status_code == 200

    resp = await client.post(
        f"/api/v1/imports/{batch_id}/confirm", headers=auth,
        json={"record_ids": [rec["id"]]},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["created"] == 1

    # 入库漏洞 source=0（来源未选择；Word导入不再单列来源）
    resp = await client.get(
        "/api/v1/vulns", headers=auth, params={"search": "已确认"}
    )
    items = resp.json()["items"]
    assert len(items) == 1
    assert items[0]["source"] == 0


def _build_report_docx(system_name: str, target_url: str, target_ip: str, sections: list[tuple[str, str, str, bool]],
                       status_texts: list[str] | None = None,
                       testers: list[str] | None = None, test_start: str = "", test_end: str = "",
                       test_account: str = ""):
    """构造平台报告格式 docx：测试目标表 + 时间与人员表（可选）+ 风险问题汇总表 + 风险问题详情章节。

    sections: [(标题, 等级文本, 类型文本, 是否已修复)]
    status_texts: 可选，逐条覆盖汇总表「修复状态」文本（如「部分未修复」），默认已修复/未修复。
    testers/test_start/test_end/test_account: 可选，构造「时间与人员」表与测试账号，
        用于验证参测人员/测试周期/测试账号解析回填。
    """
    from docx import Document

    doc = Document()
    # 测试目标表
    t = doc.add_table(rows=5, cols=2)
    t.rows[0].cells[0].text, t.rows[0].cells[1].text = "业务系统名称", system_name
    t.rows[1].cells[0].text, t.rows[1].cells[1].text = "被测系统URL", target_url
    t.rows[2].cells[0].text, t.rows[2].cells[1].text = "被测系统域名", ""
    t.rows[3].cells[0].text, t.rows[3].cells[1].text = "被测系统IP", target_ip
    t.rows[4].cells[0].text, t.rows[4].cells[1].text = "被测测试账号", test_account
    # 时间与人员表（结构与解析器 _parse_schedule_table 匹配：表头在第4行，参测人员从第5行起）
    if testers:
        sche = doc.add_table(rows=4 + len(testers), cols=4)
        sche.rows[0].cells[0].text = "测试工作时间段"
        sche.rows[1].cells[0].text, sche.rows[1].cells[1].text = "起始时间", test_start
        sche.rows[1].cells[2].text, sche.rows[1].cells[3].text = "结束时间", test_end
        for i, h in enumerate(("参测人员", "所属部门", "人员角色", "人员分工")):
            sche.rows[3].cells[i].text = h
        for i, name in enumerate(testers):
            sche.rows[4 + i].cells[0].text = name
    # 风险问题汇总表
    doc.add_heading("风险问题汇总", level=1)
    s = doc.add_table(rows=1 + len(sections), cols=4)
    for i, h in enumerate(("问题等级", "风险类型", "风险问题", "修复状态")):
        s.rows[0].cells[i].text = h
    for ri, (title, level_text, type_text, fixed) in enumerate(sections, start=1):
        s.rows[ri].cells[0].text = level_text
        s.rows[ri].cells[1].text = type_text
        s.rows[ri].cells[2].text = title
        if status_texts is not None and ri - 1 < len(status_texts):
            s.rows[ri].cells[3].text = status_texts[ri - 1]
        else:
            s.rows[ri].cells[3].text = "已修复" if fixed else "未修复"
    # 风险问题详情
    doc.add_heading("风险问题详情", level=1)
    for title, _lvl, _typ, fixed in sections:
        suffix = "（已修复）" if fixed else "（未修复）"
        doc.add_heading(f"{title}{suffix}", level=3)
        doc.add_paragraph("漏洞链接")
        doc.add_paragraph(target_url)
        doc.add_paragraph("漏洞描述")
        doc.add_paragraph("此处为漏洞描述内容。")
        doc.add_paragraph("漏洞证明")
        doc.add_paragraph("此处为漏洞证明内容。")
        doc.add_paragraph("修复建议")
        doc.add_paragraph("此处为修复建议内容。")
        doc.add_paragraph("20260728漏洞复测")
        doc.add_paragraph("复测详情：已按建议整改。")
    return doc


async def _import_report(client: AsyncClient, auth: dict, filename: str, doc):
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": (filename, buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    batch_id = resp.json()["id"]
    detail = await _wait_batch(client, auth, batch_id)
    assert detail["batch"]["status"] == "parsed", detail
    assert detail["batch"]["doc_kind"] == "report", detail
    rec_ids = [r["id"] for r in detail["records"]]
    resp = await client.post(
        f"/api/v1/imports/{batch_id}/confirm", headers=auth,
        json={"record_ids": rec_ids},
    )
    assert resp.status_code == 200, resp.text
    return detail["records"], resp.json()


async def _upload_report_batch(client: AsyncClient, auth: dict, system_name: str, filename: str) -> int:
    """上传报告格式 docx 并等待解析完成（不确认入库），返回 batch_id。"""
    doc = _build_report_docx(
        system_name, f"http://10.9.9.9/{system_name}", "10.9.9.9",
        sections=[(f"{system_name}漏洞A", "高危", "逻辑漏洞", False)],
    )
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": (filename, buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    batch_id = resp.json()["id"]
    detail = await _wait_batch(client, auth, batch_id)
    assert detail["batch"]["status"] == "parsed", detail
    assert detail["batch"]["doc_kind"] == "report", detail
    return batch_id


async def _find_vuln(client: AsyncClient, auth: dict, keyword: str) -> dict:
    resp = await client.get("/api/v1/vulns", headers=auth, params={"search": keyword})
    items = resp.json()["items"]
    assert items, f"未找到漏洞：{keyword}"
    return items[0]


async def test_report_import_partial_fixed_flow(client: AsyncClient, auth: dict):
    """复测报告含未修复漏洞：计划应为复测中(50)、未修复漏洞置「复测未修复」(50+is_retest)、自动建资产与报告。"""
    system_name = "综合办公系统ZZ"
    target_url = "http://10.9.9.9/officezz"
    doc = _build_report_docx(
        system_name, target_url, "10.9.9.9",
        sections=[
            ("平行越权访问漏洞ZZ", "高危", "逻辑漏洞", True),
            ("敏感信息泄露漏洞ZZ", "中危", "信息泄露", False),
        ],
    )
    _records, result = await _import_report(
        client, auth, "20260728综合办公系统ZZ渗透测试复测报告.docx", doc,
    )
    assert result["created"] == 2

    # 计划：复测中(50)，不是复测完成(60)
    resp = await client.get("/api/v1/testing-plans", headers=auth, params={"search": system_name})
    plans = resp.json()["items"]
    assert len(plans) == 1, plans
    assert plans[0]["status"] == 50

    # 已修复漏洞 → 60；未修复漏洞 → 50 且 is_retest=True（展示层为「复测未修复/复测未通过」）
    fixed_vuln = await _find_vuln(client, auth, "平行越权访问漏洞ZZ")
    assert fixed_vuln["status"] == 60
    unfixed_vuln = await _find_vuln(client, auth, "敏感信息泄露漏洞ZZ")
    assert unfixed_vuln["status"] == 50
    assert unfixed_vuln["is_retest"] is True
    assert any(a["name"] == system_name for a in unfixed_vuln["assets"])

    # 自动创建资产：系统名匹配，被测 URL 入内网地址
    resp = await client.get("/api/v1/assets", headers=auth, params={"search": system_name})
    assets = [a for a in resp.json()["items"] if a["name"] == system_name]
    assert assets, "未自动创建资产"
    assert target_url in (assets[0]["internal_urls"] or [])

    # 自动创建报告：显示在报告中心，草稿态，章节数与漏洞数一致
    resp = await client.get("/api/v1/reports", headers=auth, params={"search": system_name})
    reports = [r for r in resp.json()["items"] if r["project_name"] == system_name]
    assert reports, "未自动创建报告"
    report = reports[0]
    assert report["status"] == "draft"
    assert report["testing_plan_id"] == plans[0]["id"]
    detail = await client.get(f"/api/v1/reports/{report['id']}", headers=auth)
    assert len(detail.json()["sections"]) == 2


async def test_report_import_all_fixed_flow(client: AsyncClient, auth: dict):
    """复测报告全部修复：计划复测完成(60)、报告保持草稿(draft，定稿由导出 Word 驱动)。"""
    system_name = "门户系统ZZ"
    doc = _build_report_docx(
        system_name, "http://10.8.8.8/portalzz", "10.8.8.8",
        sections=[("命令执行漏洞ZZ", "高危", "命令执行漏洞", True)],
    )
    _records, result = await _import_report(
        client, auth, "20260728门户系统ZZ渗透测试复测报告.docx", doc,
    )
    assert result["created"] == 1

    resp = await client.get("/api/v1/testing-plans", headers=auth, params={"search": system_name})
    plans = resp.json()["items"]
    assert len(plans) == 1 and plans[0]["status"] == 60

    resp = await client.get("/api/v1/reports", headers=auth, params={"search": system_name})
    reports = [r for r in resp.json()["items"] if r["project_name"] == system_name]
    assert reports and reports[0]["status"] == "draft"


async def test_report_import_full_rounds_flow(client: AsyncClient, auth: dict):
    """初测 + 两轮复测报告依次导入：漏洞归并为同一条、状态按 未修复→复测未修复→已修复 流转，
    复测轮次=2、计划复测完成(60)、三份报告均自动创建。"""
    system_name = "综合办公系统RR"
    target_url = "http://10.7.7.7/oarr"
    vuln_title = "越权-劳动合同变更审批RR"

    # 1) 初测报告：发现漏洞，未修复
    _records, result = await _import_report(
        client, auth, "20250917中移系统集成有限公司综合办公系统RR渗透测试报告.docx",
        _build_report_docx(system_name, target_url, "10.7.7.7",
                           sections=[(vuln_title, "高危", "逻辑漏洞", False)]),
    )
    assert result["created"] == 1
    plan = (await client.get("/api/v1/testing-plans", headers=auth,
                             params={"search": system_name})).json()["items"][0]
    assert plan["status"] == 30  # 初测完成，等待复测
    vuln = await _find_vuln(client, auth, vuln_title)
    assert vuln["status"] == 10 and vuln["is_retest"] is False  # 初测发现漏洞，未修复

    # 2) 第一轮复测（无后缀）：部分未修复 → 漏洞置「复测未修复」(50+is_retest)
    #    标题带「（部分未修复）」后缀 + 汇总表状态「部分未修复」，验证归一化后与初测漏洞去重合并
    _records, result = await _import_report(
        client, auth, "20251011中移系统集成有限公司综合办公系统RR渗透测试复测报告.docx",
        _build_report_docx(system_name, target_url, "10.7.7.7",
                           sections=[(vuln_title + "（部分未修复）", "高危", "逻辑漏洞", False)],
                           status_texts=["部分未修复"]),
    )
    assert result["created"] == 1
    plan = (await client.get("/api/v1/testing-plans", headers=auth,
                             params={"search": system_name})).json()["items"][0]
    assert plan["status"] == 50  # 复测中
    vuln = await _find_vuln(client, auth, vuln_title)
    assert vuln["status"] == 50 and vuln["is_retest"] is True

    # 3) 第二轮复测（-1 后缀）：已修复 → 漏洞闭环，计划复测完成
    _records, result = await _import_report(
        client, auth, "20251011中移系统集成有限公司综合办公系统RR渗透测试复测报告-1.docx",
        _build_report_docx(system_name, target_url, "10.7.7.7",
                           sections=[(vuln_title, "高危", "逻辑漏洞", True)]),
    )
    assert result["created"] == 1

    # 漏洞始终只有一条，贯穿三轮
    resp = await client.get("/api/v1/vulns", headers=auth, params={"search": vuln_title})
    assert len(resp.json()["items"]) == 1
    vuln = resp.json()["items"][0]
    assert vuln["status"] == 60 and vuln["is_retest"] is True

    # 计划：复测完成(60)，两轮复测轮次记录
    plan = (await client.get("/api/v1/testing-plans", headers=auth,
                             params={"search": system_name})).json()["items"][0]
    assert plan["status"] == 60
    assert plan["retest_round_count"] == 2

    # 三份报告均自动创建，且都关联到同一工单
    resp = await client.get("/api/v1/reports", headers=auth, params={"search": system_name})
    reports = [r for r in resp.json()["items"] if r["project_name"] == system_name]
    assert len(reports) == 3, reports
    assert {r["testing_plan_id"] for r in reports} == {plan["id"]}


async def _wait_job(client: AsyncClient, auth: dict, report_id: int, job_id: int) -> dict:
    for _ in range(50):
        resp = await client.get(f"/api/v1/reports/{report_id}/exports", headers=auth)
        job = next(j for j in resp.json() if j["id"] == job_id)
        if job["status"] in ("done", "failed"):
            return job
        await asyncio.sleep(0.2)
    raise AssertionError("导出任务超时")


async def test_report_edit_and_export(client: AsyncClient, auth: dict):
    # 取已有漏洞生成报告
    resp = await client.get("/api/v1/vulns", headers=auth)
    vul_ids = [v["id"] for v in resp.json()["items"]][:2]
    assert vul_ids

    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "季度渗透测试报告", "vul_ids": vul_ids},
    )
    assert resp.status_code == 200, resp.text
    report = resp.json()
    assert len(report["sections"]) == len(vul_ids)
    report_id = report["id"]

    # 错误 revision 保存应 409（乐观锁）
    save_body = {
        "title": "季度渗透测试报告 V2",
        "project_name": "商城安全测试",
        "revision": report["revision"] + 99,
        "sections": report["sections"],
    }
    resp = await client.put(f"/api/v1/reports/{report_id}", headers=auth, json=save_body)
    assert resp.status_code == 409

    # 正确 revision 保存成功：revision +1，导出 version 不受编辑影响
    save_body["revision"] = report["revision"]
    save_body["target_ip"] = "10.0.0.8"
    resp = await client.put(f"/api/v1/reports/{report_id}", headers=auth, json=save_body)
    assert resp.status_code == 200, resp.text
    saved = resp.json()
    assert saved["revision"] == report["revision"] + 1
    assert saved["version"] == report["version"]  # 保存不改导出版本号
    assert saved["title"] == "季度渗透测试报告 V2"
    assert saved["target_ip"] == "10.0.0.8"

    # 导出 docx（pdf 依赖 Gotenberg，容器环境验证）
    resp = await client.post(
        f"/api/v1/reports/{report_id}/export", headers=auth, json={"fmt": "docx"}
    )
    assert resp.status_code == 200, resp.text
    job_id = resp.json()["id"]

    job = await _wait_job(client, auth, report_id, job_id)
    assert job["status"] == "done", job

    resp = await client.get(f"/api/v1/reports/exports/{job_id}/download", headers=auth)
    assert resp.status_code == 200
    assert resp.content[:2] == b"PK"  # docx 是 zip 容器

    # 导出成功后导出版本 +1；测试周期在字段为空时自动预填（开始=最早提交日期，结束=当天）
    from app.core.timeutil import now as _tnow

    after = (await client.get(f"/api/v1/reports/{report_id}", headers=auth)).json()
    assert after["version"] == saved["version"] + 1
    assert after["test_start"]
    assert after["test_end"] == _tnow().date().isoformat()

    # 产物基于渗透测试报告模板：验证封面系统名称、测试目标 IP、汇总表与详情标题
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(resp.content))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    assert "商城安全测试" in texts  # 封面第二行=系统名称(project_name)，报告标题作为文件名
    assert any(t == "风险问题详情" for t in texts)
    # 测试目标表：被测系统 IP
    target_tbl = doc.tables[4]
    assert target_tbl.rows[3].cells[1].text.strip() == "10.0.0.8"
    # 汇总表：表头 + 每个关联漏洞一行，且样例行已被替换
    summary_tbl = doc.tables[6]
    assert len(summary_tbl.rows) == 1 + len(vul_ids)
    assert "平行越权" not in summary_tbl.rows[1].cells[2].text
    # 详情段：每个章节一个 Heading 3，标题含修复状态后缀
    h3 = [p.text for p in doc.paragraphs if p.style.name == "Heading 3"]
    assert len(h3) == len(vul_ids)
    assert all(t.endswith("）") for t in h3)


async def test_report_from_vulns_infer_plan_id(client: AsyncClient, auth: dict):
    """from-vulns 未显式指定计划时：漏洞归属唯一计划则自动回写，多计划则不回写。"""
    # 计划一 + 归属该计划的两个漏洞
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "回写系统A", "test_type": "渗透测试"},
    )
    plan_a = resp.json()["id"]
    # 严格认领：需先认领计划才能批量建计划漏洞
    await client.post(f"/api/v1/testing-plans/{plan_a}/claim", headers=auth)
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "回写资产A"})
    asset_a = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [asset_a], "vulns": [
            {"title": "回写漏洞A1", "level": 20, "testing_plan_id": plan_a},
            {"title": "回写漏洞A2", "level": 30, "testing_plan_id": plan_a},
        ]},
    )
    a1, a2 = [v["id"] for v in resp.json()]

    # 未传 testing_plan_id：唯一归属计划应被自动回写
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "自动回写报告", "vul_ids": [a1, a2]},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["testing_plan_id"] == plan_a

    # 计划二 + 归属该计划的漏洞
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "回写系统B", "test_type": "渗透测试"},
    )
    plan_b = resp.json()["id"]
    await client.post(f"/api/v1/testing-plans/{plan_b}/claim", headers=auth)
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "回写资产B"})
    asset_b = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [asset_b], "vulns": [
            {"title": "回写漏洞B1", "level": 20, "testing_plan_id": plan_b},
        ]},
    )
    b1 = resp.json()[0]["id"]

    # 跨两个计划的漏洞：归属不唯一，不回写
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "多计划报告", "vul_ids": [a1, b1]},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["testing_plan_id"] is None


async def test_report_export_target_urls_from_plan(client: AsyncClient, auth: dict):
    """工单「被测系统URL」为报告测试目标表优先数据源：资产未录URL时导出仍能带出URL/域名。"""
    from io import BytesIO

    from docx import Document

    # 工单维护被测系统URL；关联资产未录URL
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "目标系统", "test_type": "渗透测试",
              "target_urls": ["https://target.example.com/app", "http://10.20.1.10:8080"]},
    )
    plan_id = resp.json()["id"]
    await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "无URL资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [asset_id], "vulns": [{"title": "目标漏洞", "level": 30}]},
    )
    vul_id = resp.json()[0]["id"]

    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "目标URL报告", "vul_ids": [vul_id], "testing_plan_id": plan_id},
    )
    report_id = resp.json()["id"]

    resp = await client.post(
        f"/api/v1/reports/{report_id}/export", headers=auth, json={"fmt": "docx"}
    )
    job_id = resp.json()["id"]
    job = await _wait_job(client, auth, report_id, job_id)
    assert job["status"] == "done", job

    resp = await client.get(f"/api/v1/reports/exports/{job_id}/download", headers=auth)
    doc = Document(BytesIO(resp.content))
    target_tbl = doc.tables[4]
    # URL格取工单 target_urls，域名格由URL推导（纯 IP hostname 不计入域名）
    assert target_tbl.rows[1].cells[1].text.strip() == "https://target.example.com/app\nhttp://10.20.1.10:8080"
    assert target_tbl.rows[2].cells[1].text.strip() == "target.example.com"


async def test_report_similarity_check(client: AsyncClient, auth: dict):
    """相似性检查：基础信息（标题+归属计划+漏洞集合）与所选漏洞最后编辑时间完全一致才判相似。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "相似系统", "test_type": "渗透测试"},
    )
    plan_id = resp.json()["id"]
    # 录入漏洞需先认领计划（严格认领校验，管理员未认领也不放行）
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    assert resp.status_code == 200, resp.text
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "相似资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [asset_id], "vulns": [
            {"title": "相似漏洞A", "level": 20, "testing_plan_id": plan_id},
            {"title": "相似漏洞B", "level": 30, "testing_plan_id": plan_id},
        ]},
    )
    vul_a, vul_b = [v["id"] for v in resp.json()]

    # 生成报告后，同配置再次检查 → 高度相似
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "相似性检查报告", "vul_ids": [vul_a, vul_b], "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    report_id = resp.json()["id"]

    async def check(**overrides):
        payload = {
            "title": "相似性检查报告",
            "vul_ids": [vul_a, vul_b],
            "testing_plan_id": plan_id,
            **overrides,
        }
        return (await client.post("/api/v1/reports/similarity-check", headers=auth, json=payload)).json()

    data = await check()
    assert data["similar"] is True
    assert [r["id"] for r in data["matched_reports"]] == [report_id]

    # 漏洞顺序无关（集合比较）
    assert (await check(vul_ids=[vul_b, vul_a]))["similar"] is True

    # 未显式传计划时，由漏洞归属唯一计划推导 → 同样命中
    payload = {"title": "相似性检查报告", "vul_ids": [vul_a, vul_b]}
    resp = await client.post("/api/v1/reports/similarity-check", headers=auth, json=payload)
    assert resp.json()["similar"] is True

    # 所选漏洞最后编辑时间变化（编辑漏洞）→ 不再相似
    resp = await client.put(f"/api/v1/vulns/{vul_a}", headers=auth, json={"title": "相似漏洞A-已修订"})
    assert resp.status_code == 200, resp.text
    assert (await check())["similar"] is False

    # 标题不同 / 漏洞集合不同 / 空集合 → 不相似
    assert (await check(title="完全不同标题"))["similar"] is False
    assert (await check(vul_ids=[vul_a]))["similar"] is False
    assert (await check(vul_ids=[]))["similar"] is False


async def test_report_similarity_check_legacy_no_snapshot(client: AsyncClient, auth: dict):
    """存量报告（无 vul_edit_snapshot）再次生成时：漏洞集合一致则回填快照并判相似。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "存量系统", "test_type": "渗透测试"},
    )
    plan_id = resp.json()["id"]
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    assert resp.status_code == 200, resp.text
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "存量资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [asset_id], "vulns": [
            {"title": "存量漏洞A", "level": 20, "testing_plan_id": plan_id},
            {"title": "存量漏洞B", "level": 30, "testing_plan_id": plan_id},
        ]},
    )
    vul_a, vul_b = [v["id"] for v in resp.json()]
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "存量相似报告", "vul_ids": [vul_a, vul_b], "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    report_id = resp.json()["id"]

    # 模拟存量数据：清空快照
    from sqlalchemy import select as sa_select

    from app.db import async_session_maker
    from app.models import Report

    async with async_session_maker() as session:
        report = (
            await session.execute(sa_select(Report).where(Report.id == report_id))
        ).scalar_one()
        report.vul_edit_snapshot = None
        await session.commit()

    async def check(**overrides):
        payload = {
            "title": "存量相似报告",
            "vul_ids": [vul_a, vul_b],
            "testing_plan_id": plan_id,
            **overrides,
        }
        resp = await client.post("/api/v1/reports/similarity-check", headers=auth, json=payload)
        return resp.json()

    # 无快照但基础信息与漏洞集合一致 → 判相似，且回填快照
    data = await check()
    assert data["similar"] is True
    assert [r["id"] for r in data["matched_reports"]] == [report_id]
    async with async_session_maker() as session:
        report = (
            await session.execute(sa_select(Report).where(Report.id == report_id))
        ).scalar_one()
        assert report.vul_edit_snapshot is not None

    # 回填后再次检查仍相似
    assert (await check())["similar"] is True

    # 漏洞集合不一致时不误匹配
    assert (await check(vul_ids=[vul_a]))["similar"] is False

    # 漏洞被编辑（快照已回填旧值）→ 不再相似
    resp = await client.put(f"/api/v1/vulns/{vul_a}", headers=auth, json={"title": "存量漏洞A-已修订"})
    assert resp.status_code == 200, resp.text
    assert (await check())["similar"] is False


async def test_report_export_duplicate_check(client: AsyncClient, auth: dict):
    """导出前重复判断：内容指纹与上次成功导出一致则 duplicate，编辑报告后不再重复。"""
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "重复导出资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [asset_id], "vulns": [
            {"title": "重复导出漏洞A", "level": 20},
            {"title": "重复导出漏洞B", "level": 30},
        ]},
    )
    vul_ids = [v["id"] for v in resp.json()]
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "重复导出检查报告", "vul_ids": vul_ids},
    )
    assert resp.status_code == 200, resp.text
    report_id = resp.json()["id"]

    async def check_export(fmt: str = "docx"):
        resp = await client.post(
            f"/api/v1/reports/{report_id}/export-check", headers=auth, json={"fmt": fmt},
        )
        assert resp.status_code == 200, resp.text
        return resp.json()

    # 尚无成功导出 → 不重复
    assert (await check_export())["duplicate"] is False

    # 导出 docx 并等待完成
    resp = await client.post(
        f"/api/v1/reports/{report_id}/export", headers=auth, json={"fmt": "docx"},
    )
    assert resp.status_code == 200, resp.text
    job = await _wait_job(client, auth, report_id, resp.json()["id"])
    assert job["status"] == "done", job

    # 内容未变 → 重复（附完整提示信息）；其他格式无历史 → 不重复
    data = await check_export()
    assert data["duplicate"] is True
    assert data["fmt"] == "docx"
    assert data["report_title"] == "重复导出检查报告"
    assert data["last_status"] == "done"
    assert data["last_time"] is not None
    assert data["last_version"] == 2  # 初始 v1 + 导出成功后版本 +1
    assert data["last_file_name"].endswith(".docx")
    assert (await check_export("pdf"))["duplicate"] is False

    # 编辑报告（revision/update_time 变化）→ 不再重复
    detail = (await client.get(f"/api/v1/reports/{report_id}", headers=auth)).json()
    save_body = {
        "title": "重复导出检查报告",
        "revision": detail["revision"],
        "sections": detail["sections"],
    }
    resp = await client.put(f"/api/v1/reports/{report_id}", headers=auth, json=save_body)
    assert resp.status_code == 200, resp.text
    assert (await check_export())["duplicate"] is False


async def test_report_vuln_state_automation(client: AsyncClient, auth: dict):
    """报告联动状态机：生成报告→修复中，发起复测→复测中，全部已修复/已忽略→报告已完成。"""
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "联动测试系统"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={
            "asset_ids": [asset_id],
            "vulns": [{"title": "联动漏洞A", "level": 20}, {"title": "联动漏洞B", "level": 30}],
        },
    )
    vul_a, vul_b = [v["id"] for v in resp.json()]

    # 生成报告后关联漏洞自动进入修复中
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "联动测试报告", "vul_ids": [vul_a, vul_b]},
    )
    assert resp.status_code == 200, resp.text
    report_id = resp.json()["id"]
    for vid in (vul_a, vul_b):
        vul = (await client.get(f"/api/v1/vulns/{vid}", headers=auth)).json()
        assert vul["status"] == 50
        assert vul["notice_time"] is not None

    # 点击复测：关联漏洞自动进入复测中
    resp = await client.post(f"/api/v1/reports/{report_id}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    for vid in (vul_a, vul_b):
        vul = (await client.get(f"/api/v1/vulns/{vid}", headers=auth)).json()
        assert vul["status"] == 55

    # 复测面板数据源：报告关联漏洞状态列表
    resp = await client.get(f"/api/v1/reports/{report_id}/vuln-states", headers=auth)
    assert resp.status_code == 200, resp.text
    states = resp.json()
    assert {s["vul_id"] for s in states} == {vul_a, vul_b}
    assert all(s["status"] == 55 for s in states)

    # 复测中可选：已修复/复测未通过(回修复中)/已忽略/暂不处理
    resp = await client.get(f"/api/v1/vulns/{vul_a}/transitions", headers=auth)
    assert {t["status"] for t in resp.json()} == {20, 35, 50, 60}

    # A 已修复（携带复测详情）：报告尚未全部处理完，不应自动完成
    resp = await client.post(
        f"/api/v1/vulns/{vul_a}/transition", headers=auth,
        json={
            "status": 60,
            "retest_html": "<p>复测通过，漏洞已修复</p>",
            "retest_json": {"type": "doc", "content": []},
        },
    )
    assert resp.status_code == 200
    assert resp.json()["retest_html"] == "<p>复测通过，漏洞已修复</p>"
    vul = (await client.get(f"/api/v1/vulns/{vul_a}", headers=auth)).json()
    assert vul["retest_html"] == "<p>复测通过，漏洞已修复</p>"
    assert vul["retest_json"] == {"type": "doc", "content": []}
    report = (await client.get(f"/api/v1/reports/{report_id}", headers=auth)).json()
    assert report["status"] == "draft"

    # B 已忽略：全部为已修复/已忽略，计划复测完成，报告保持草稿（需求6：报告状态由导出定稿驱动）
    resp = await client.post(
        f"/api/v1/vulns/{vul_b}/transition", headers=auth, json={"status": 20},
    )
    assert resp.status_code == 200
    report = (await client.get(f"/api/v1/reports/{report_id}", headers=auth)).json()
    assert report["status"] == "draft"

    # 无关联漏洞的报告不能发起复测
    resp = await client.post("/api/v1/reports", headers=auth, json={"title": "空报告", "sections": []})
    empty_id = resp.json()["id"]
    resp = await client.post(f"/api/v1/reports/{empty_id}/retest", headers=auth)
    assert resp.status_code == 400


async def test_retest_keeps_fixed_vuln_status(client: AsyncClient, auth: dict):
    """发起复测仅对未修复漏洞生效：已修复(60)漏洞保持原状态，不重新进入复测中。"""
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "复测筛选测试系统"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={
            "asset_ids": [asset_id],
            "vulns": [{"title": "筛选漏洞A", "level": 20}, {"title": "筛选漏洞B", "level": 30}],
        },
    )
    vul_a, vul_b = [v["id"] for v in resp.json()]

    # 生成报告 → 修复中(50)；首次发起复测 → 复测中(55)
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "复测筛选报告", "vul_ids": [vul_a, vul_b]},
    )
    assert resp.status_code == 200, resp.text
    report_id = resp.json()["id"]
    resp = await client.post(f"/api/v1/reports/{report_id}/retest", headers=auth)
    assert resp.status_code == 200, resp.text

    # A 已修复（携带复测详情）
    resp = await client.post(
        f"/api/v1/vulns/{vul_a}/transition", headers=auth,
        json={"status": 60, "retest_html": "<p>复测通过</p>"},
    )
    assert resp.status_code == 200
    assert (await client.get(f"/api/v1/vulns/{vul_a}", headers=auth)).json()["status"] == 60

    # 再次发起复测：已修复 A 保持 60，未修复 B 保持 55，均不重复进入复测中
    resp = await client.post(f"/api/v1/reports/{report_id}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    assert (await client.get(f"/api/v1/vulns/{vul_a}", headers=auth)).json()["status"] == 60
    assert (await client.get(f"/api/v1/vulns/{vul_b}", headers=auth)).json()["status"] == 55

    # 全部漏洞已修复后再次发起复测：均不再进入复测中，保持已修复
    resp = await client.post(
        f"/api/v1/vulns/{vul_b}/transition", headers=auth,
        json={"status": 60, "retest_html": "<p>复测通过</p>"},
    )
    assert resp.status_code == 200
    resp = await client.post(f"/api/v1/reports/{report_id}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    assert (await client.get(f"/api/v1/vulns/{vul_a}", headers=auth)).json()["status"] == 60
    assert (await client.get(f"/api/v1/vulns/{vul_b}", headers=auth)).json()["status"] == 60


async def test_retest_blocked_without_update(client: AsyncClient, auth: dict):
    """复测报告生成防重：漏洞状态/内容未更新时再次发起复测被阻止（不新增报告、不增加轮次）；
    状态有更新后允许再次发起复测，同日标题重复自动追加 -1 后缀。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "复测防重系统", "test_type": "渗透测试"},
    )
    assert resp.status_code == 200, resp.text
    plan_id = resp.json()["id"]
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    assert resp.status_code == 200, resp.text
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "复测防重资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [asset_id], "vulns": [
            {"title": "防重漏洞A", "level": 20, "testing_plan_id": plan_id},
            {"title": "防重漏洞B", "level": 30, "testing_plan_id": plan_id},
        ]},
    )
    assert resp.status_code == 200, resp.text
    vul_a, vul_b = [v["id"] for v in resp.json()]

    # 生成初测报告 → 漏洞进入修复中(50)
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "防重系统渗透测试报告", "vul_ids": [vul_a, vul_b], "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    src_report = resp.json()["id"]

    # 第一次发起复测 → 生成复测报告 R1，漏洞进入复测中(55)
    resp = await client.post(f"/api/v1/reports/{src_report}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    r1_id = resp.json()["id"]
    assert (await client.get(f"/api/v1/vulns/{vul_a}", headers=auth)).json()["status"] == 55

    # 状态/内容未变化再次发起复测 → 阻止并提示，不新增报告、不增加轮次
    resp = await client.post(f"/api/v1/reports/{src_report}/retest", headers=auth)
    assert resp.status_code == 400, resp.text
    assert "复测结果未更新" in resp.json()["detail"]
    plan = (await client.get(f"/api/v1/testing-plans/{plan_id}", headers=auth)).json()
    assert len(plan["reports"]) == 2
    assert plan["retest_round_count"] == 1

    # 处理复测：A 已修复（携带复测内容）→ 状态有更新，允许再次发起复测
    resp = await client.post(
        f"/api/v1/vulns/{vul_a}/transition", headers=auth,
        json={"status": 60, "retest_html": "<p>复测通过</p>"},
    )
    assert resp.status_code == 200
    resp = await client.post(f"/api/v1/reports/{src_report}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    r2_id = resp.json()["id"]
    assert r2_id != r1_id
    # 同日标题重复 → 自动追加 -1 后缀
    assert resp.json()["title"].endswith("-1")

    # 最终：1 初测报告 + 2 复测报告；复测轮数 2（被阻止的一次未计入）
    plan = (await client.get(f"/api/v1/testing-plans/{plan_id}", headers=auth)).json()
    assert len(plan["reports"]) == 3
    assert plan["retest_round_count"] == 2


async def test_import_confirm_into_report(client: AsyncClient, auth: dict):
    """Word 导入确认时关联报告：自动追加章节、漏洞进入修复中。"""
    resp = await client.post(
        "/api/v1/reports", headers=auth, json={"title": "导入关联报告", "sections": []},
    )
    report = resp.json()
    report_id = report["id"]

    resp = await client.get("/api/v1/imports/template", headers=auth)
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": ("关联报告样例.docx", BytesIO(resp.content),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    batch_id = resp.json()["id"]
    detail = await _wait_batch(client, auth, batch_id)
    assert detail["batch"]["status"] == "parsed", detail
    rec = detail["records"][0]

    # 不存在的报告被拒绝
    resp = await client.post(
        f"/api/v1/imports/{batch_id}/confirm", headers=auth,
        json={"record_ids": [rec["id"]], "report_id": 999999},
    )
    assert resp.status_code == 400

    resp = await client.post(
        f"/api/v1/imports/{batch_id}/confirm", headers=auth,
        json={"record_ids": [rec["id"]], "report_id": report_id},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["created"] == 1

    # 报告新增关联章节且编辑锁 revision +1（导出版本号 version 不受编辑影响）
    saved = (await client.get(f"/api/v1/reports/{report_id}", headers=auth)).json()
    assert len(saved["sections"]) == 1
    assert saved["sections"][0]["vul_id"] is not None
    assert saved["revision"] == report["revision"] + 1

    # 入库漏洞自动进入修复中
    vul = (await client.get(f"/api/v1/vulns/{saved['sections'][0]['vul_id']}", headers=auth)).json()
    assert vul["status"] == 50


async def test_import_report_author_from_testers(client: AsyncClient, auth: dict):
    """显式关联渗透测试工单导入报告：自动创建的报告作者取自工单测试人员姓名。"""
    system_name = "作者同步系统TA"
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": system_name, "test_type": "渗透测试"},
    )
    assert resp.status_code == 200, resp.text
    plan_id = resp.json()["id"]
    # admin 认领工单 → testers=[管理员]
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    assert resp.status_code == 200, resp.text
    assert resp.json()["testers"][0]["realname"] == "管理员"

    doc = _build_report_docx(
        system_name, "http://10.6.6.6/authorTA", "10.6.6.6",
        sections=[("报告作者同步漏洞TA", "高危", "逻辑漏洞", False)],
    )
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": ("20260728作者同步系统TA渗透测试报告.docx", buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    batch_id = resp.json()["id"]
    detail = await _wait_batch(client, auth, batch_id)
    assert detail["batch"]["doc_kind"] == "report", detail
    rec_ids = [r["id"] for r in detail["records"]]

    resp = await client.post(
        f"/api/v1/imports/{batch_id}/confirm", headers=auth,
        json={"record_ids": rec_ids, "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["created"] >= 1

    # 自动创建的报告作者 = 工单测试人员姓名（「、」拼接）
    resp = await client.get("/api/v1/reports", headers=auth, params={"search": system_name})
    reports = [r for r in resp.json()["items"] if r["project_name"] == system_name]
    assert reports, "未自动创建报告"
    report = (await client.get(f"/api/v1/reports/{reports[0]['id']}", headers=auth)).json()
    assert report["author"] == "管理员"
    assert report["testing_plan_id"] == plan_id


async def test_import_report_fields_and_auto_export(client: AsyncClient, auth: dict):
    """导入报告含参测人员与测试周期：映射系统账号关联工单、回填报告字段
    （作者/测试周期/测试账号/实际人天）、被测URL更新资产、报告时间取标题日期、
    自动生成可下载的导出记录（时间=报告日期 14:00）。"""
    # 创建参测人员对应系统账号（对应生产环境 admin/xna/xtz 三人场景）
    for username, realname in (("xna", "许宁安"), ("xtz", "薛田泽")):
        resp = await client.post(
            "/api/v1/users", headers=auth,
            json={"username": username, "password": "Tester@123", "realname": realname,
                  "email": "", "phone": "", "is_active": True},
        )
        assert resp.status_code == 200, resp.text

    system_name = "字段回填系统FB"
    doc = _build_report_docx(
        system_name, "http://10.7.7.7/fieldfb", "10.7.7.7",
        sections=[("字段回填漏洞FB", "中危", "信息泄露", False)],
        testers=["管理员", "许宁安", "薛田泽"], test_start="2026-06-30", test_end="2026-07-01",
        test_account="admin/Admin@123",
    )
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": ("20260701字段回填系统FB渗透测试报告.docx", buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    batch_id = resp.json()["id"]
    detail = await _wait_batch(client, auth, batch_id)
    assert detail["batch"]["doc_kind"] == "report", detail
    rec_ids = [r["id"] for r in detail["records"]]

    resp = await client.post(
        f"/api/v1/imports/{batch_id}/confirm", headers=auth,
        json={"record_ids": rec_ids},
    )
    assert resp.status_code == 200, resp.text

    # 自动创建的工单已按姓名关联参测人员账号（admin=管理员、xna=许宁安、xtz=薛田泽）
    resp = await client.get("/api/v1/reports", headers=auth, params={"search": system_name})
    reports = [r for r in resp.json()["items"] if r["project_name"] == system_name]
    assert reports, "未自动创建报告"
    report = (await client.get(f"/api/v1/reports/{reports[0]['id']}", headers=auth)).json()
    plan = (await client.get(f"/api/v1/testing-plans/{report['testing_plan_id']}", headers=auth)).json()
    assert {u["realname"] for u in plan["testers"]} == {"管理员", "许宁安", "薛田泽"}
    assert {u["username"] for u in plan["testers"]} >= {"admin", "xna", "xtz"}

    # 报告字段回填：作者 / 测试周期 / 测试账号 / 实际人天 / 报告时间（取标题日期 2026-07-01 14:00）
    assert report["author"] == "管理员、许宁安、薛田泽"
    assert report["test_start"] == "2026-06-30"
    assert report["test_end"] == "2026-07-01"
    assert report["test_account"] == "admin/Admin@123"
    assert report["actual_mandays"] == 2  # 2026-06-30 ~ 2026-07-01
    assert report["create_time"].startswith("2026-07-01T14:00")
    # 自动导出成功：导出版本 +1；报告保持草稿（定稿仍由人工导出 Word 驱动）
    assert report["version"] == 2
    assert report["status"] == "draft"

    # 工单实际人天同步刷新（仅纳入初测报告）
    assert plan["actual_mandays"] == 2

    # 漏洞提交时间 = 报告时间（标题日期 14:00），保证按月统计口径一致
    resp = await client.get("/api/v1/vulns", headers=auth, params={"search": "字段回填漏洞FB"})
    vul = resp.json()["items"][0]
    assert vul["submit_time"].startswith("2026-07-01T14:00"), vul["submit_time"]

    # 被测系统 URL 自动更新到资产（internal_urls 去重）
    resp = await client.get("/api/v1/assets", headers=auth, params={"search": system_name})
    asset = [a for a in resp.json()["items"] if a["name"] == system_name][0]
    assert "http://10.7.7.7/fieldfb" in asset["internal_urls"]

    # 自动生成一条导出记录：时间=报告日期 14:00，含实际文件可下载
    resp = await client.get(f"/api/v1/reports/{report['id']}/exports", headers=auth)
    assert resp.status_code == 200, resp.text
    assert len(resp.json()) == 1, resp.text
    job = resp.json()[0]
    assert job["status"] == "done"
    assert job["create_time"].startswith("2026-07-01T14:00")
    assert job["finish_time"].startswith("2026-07-01T14:00")
    assert job["has_file"] is True
    dl = await client.get(f"/api/v1/reports/exports/{job['id']}/download", headers=auth)
    assert dl.status_code == 200, dl.text
    assert len(dl.content) > 0


async def test_import_report_public_urls_and_doc_time(client: AsyncClient, auth: dict):
    """多条被测系统 URL 分别录入资产（公网→public_urls、内网→internal_urls）；
    下载的 docx 封面与版本变更记录时间与导入报告时间一致。"""
    system_name = "多URL分类系统MU"
    doc = _build_report_docx(
        system_name, "https://www.a-mu.com\nhttps://www.b-mu.com\n10.30.30.30", "10.30.30.30",
        sections=[("多URL漏洞MU", "中危", "信息泄露", False)],
        testers=["管理员"], test_start="2026-07-10", test_end="2026-07-15",
    )
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": ("20260715多URL分类系统MU渗透测试报告.docx", buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    batch_id = resp.json()["id"]
    detail = await _wait_batch(client, auth, batch_id)
    rec_ids = [r["id"] for r in detail["records"]]
    resp = await client.post(
        f"/api/v1/imports/{batch_id}/confirm", headers=auth,
        json={"record_ids": rec_ids},
    )
    assert resp.status_code == 200, resp.text

    # 公网 URL 分别录入 public_urls（tag=10 互联网），内网 IP 录入 internal_urls
    resp = await client.get("/api/v1/assets", headers=auth, params={"search": system_name})
    asset = [a for a in resp.json()["items"] if a["name"] == system_name][0]
    pub = {u["url"]: u["tag"] for u in (asset["public_urls"] or [])}
    assert pub.get("https://www.a-mu.com") == 10
    assert pub.get("https://www.b-mu.com") == 10
    assert "10.30.30.30" in (asset["internal_urls"] or [])

    # 下载自动导出的 docx：封面日期与版本变更记录 V1.0 日期均为导入报告时间 2026-07-15
    resp = await client.get("/api/v1/reports", headers=auth, params={"search": system_name})
    report = [r for r in resp.json()["items"] if r["project_name"] == system_name][0]
    jobs = (await client.get(f"/api/v1/reports/{report['id']}/exports", headers=auth)).json()
    assert len(jobs) == 1, jobs
    dl = await client.get(f"/api/v1/reports/exports/{jobs[0]['id']}/download", headers=auth)
    assert dl.status_code == 200, dl.text

    from docx import Document

    docx = Document(BytesIO(dl.content))
    texts = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
    assert any("2026年07月15日" in t for t in texts), texts
    cells = [c.text.strip() for c in docx.tables[1].rows[2].cells]
    assert cells[0] == "2026-07-15", cells


async def test_special_modules_crud(client: AsyncClient, auth: dict):
    """三个专项模块：远程检测 / 测试计划 / 春耕行动 CRUD。"""
    # ---- 远程检测（2026-08-14 按通报口径重构：申诉报告改为附件上传） ----
    resp = await client.post(
        "/api/v1/remote-testings", headers=auth,
        json={"system_name": "门户系统", "notice_time": "2026-01", "department": "信息部",
              "notified_unit": "省公司", "is_external": False, "vuln_name": "SQL注入",
              "vuln_type": "SQL注入", "appeal_status": "", "appeal_method": "",
              "appeal_file_name": "", "appeal_file_path": "", "appeal_file_size": 0},
    )
    assert resp.status_code == 200, resp.text
    rt_id = resp.json()["id"]
    assert resp.json()["system_name"] == "门户系统"

    # 申诉报告附件上传（返回文件元信息供表单绑定）
    resp = await client.post(
        "/api/v1/remote-testings/upload-appeal", headers=auth,
        files={"file": ("appeal.pdf", b"%PDF-1.4 test", "application/pdf")},
    )
    assert resp.status_code == 200, resp.text
    up = resp.json()
    assert up["name"] == "appeal.pdf"
    assert up["path"].startswith("uploads/remote_appeal/")
    assert up["size"] == len(b"%PDF-1.4 test")

    resp = await client.put(
        f"/api/v1/remote-testings/{rt_id}", headers=auth,
        json={"system_name": "门户系统", "notice_time": "2026-02", "department": "信息部",
              "notified_unit": "省公司", "is_external": True, "vuln_name": "SQL注入",
              "vuln_type": "SQL注入", "appeal_status": "success", "appeal_method": "线下申诉",
              "appeal_file_name": up["name"], "appeal_file_path": up["path"],
              "appeal_file_size": up["size"]},
    )
    assert resp.status_code == 200
    assert resp.json()["appeal_status"] == "success"

    # 申诉报告附件下载
    resp = await client.get(f"/api/v1/remote-testings/{rt_id}/appeal", headers=auth)
    assert resp.status_code == 200
    assert resp.content == b"%PDF-1.4 test"

    resp = await client.get("/api/v1/remote-testings", headers=auth, params={"search": "门户"})
    assert rt_id in [r["id"] for r in resp.json()["items"]]

    # ---- 测试计划 ----
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "计划系统", "test_type": "渗透测试", "department": "研发部",
              "receive_time": "2026-01-01", "first_test_done_time": "2026-01-05",
              "status": 20, "stat_critical": 1, "stat_high": 2, "stat_medium": 3,
              "stat_low": 4, "target_urls": ["https://plan.example.com"],
              "detail": "测试人员：张三"},
    )
    assert resp.status_code == 200, resp.text
    plan = resp.json()
    assert plan["stat_high"] == 2
    assert plan["target_urls"] == ["https://plan.example.com"]

    # 编辑：被测系统URL支持增删（替换）
    resp = await client.put(
        f"/api/v1/testing-plans/{plan['id']}", headers=auth,
        json={"system_name": "计划系统", "test_type": "渗透测试", "department": "研发部",
              "receive_time": "2026-01-01", "status": 20,
              "target_urls": ["https://a.example.com", "http://10.20.1.10:8080"]},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["target_urls"] == ["https://a.example.com", "http://10.20.1.10:8080"]

    # status 筛选
    resp = await client.get("/api/v1/testing-plans", headers=auth, params={"status": 20})
    assert plan["id"] in [p["id"] for p in resp.json()["items"]]
    resp = await client.get("/api/v1/testing-plans", headers=auth, params={"status": 50})
    assert plan["id"] not in [p["id"] for p in resp.json()["items"]]

    # meta 提供六档状态字典
    meta = (await client.get("/api/v1/meta", headers=auth)).json()
    assert meta["testing_plan_status"]["10"] == "未测试"
    assert meta["testing_plan_status"]["60"] == "复测完成"

    # ---- 春耕行动 ----
    resp = await client.post(
        "/api/v1/vulns", headers=auth, json={"title": "春耕漏洞", "level": 20},
    )
    vul_id = resp.json()["id"]

    # 不存在的漏洞被拒绝
    resp = await client.post(
        "/api/v1/spring-actions", headers=auth,
        json={"report_no": "RPT-BAD", "vul_ids": [999999]},
    )
    assert resp.status_code == 400

    resp = await client.post(
        "/api/v1/spring-actions", headers=auth,
        json={"report_no": "RPT-2026-001", "system_name": "春耕系统",
              "asset_reason": "备案归属本单位", "appeal_success": True,
              "est_score_deduction": 4, "score_deduction": 2.5,
              "doc_no": "公文〔2026〕1号", "vul_ids": [vul_id]},
    )
    assert resp.status_code == 200, resp.text
    sa = resp.json()
    assert sa["vul_ids"] == [vul_id]
    assert sa["vuls"][0]["title"] == "春耕漏洞"
    assert sa["asset_reason"] == "备案归属本单位"
    assert sa["est_score_deduction"] == 4
    assert sa["score_deduction"] == 2.5

    # 更新：清空漏洞关联
    resp = await client.put(
        f"/api/v1/spring-actions/{sa['id']}", headers=auth,
        json={"report_no": "RPT-2026-001", "system_name": "春耕系统",
              "asset_reason": "", "appeal_success": False,
              "est_score_deduction": 0, "score_deduction": 0,
              "doc_no": "", "vul_ids": []},
    )
    assert resp.status_code == 200
    assert resp.json()["vuls"] == []

    # 列表返回漏洞摘要
    resp = await client.get("/api/v1/spring-actions", headers=auth, params={"search": "RPT-2026"})
    assert len(resp.json()["items"]) == 1

    # ---- 原始报告上传导入 ----
    from docx import Document as _Docx

    # 非 docx 被拒绝
    resp = await client.post(
        "/api/v1/spring-actions/upload-report", headers=auth,
        files={"file": ("原始报告.txt", b"not a docx", "text/plain")},
    )
    assert resp.status_code == 400

    # 无漏洞表的 docx 也可上传留档，解析返回空草稿
    buf = BytesIO()
    _Docx().save(buf)
    buf.seek(0)
    resp = await client.post(
        "/api/v1/spring-actions/upload-report", headers=auth,
        files={"file": ("原始报告.docx", buf,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert resp.status_code == 200, resp.text
    upload = resp.json()
    assert upload["name"] == "原始报告.docx"
    assert upload["path"].startswith("uploads/spring_report/")
    assert upload["size"] > 0
    assert upload["vuls"] == []

    # 保存：附件绑定 + 漏洞草稿随保存创建并关联（来源固定为春耕行动）
    resp = await client.post(
        "/api/v1/spring-actions", headers=auth,
        json={"report_no": "SA-2026-009", "system_name": "春耕附件系统",
              "report_file_name": upload["name"], "report_file_path": upload["path"],
              "report_file_size": upload["size"],
              "new_vuls": [{"title": "报告导入漏洞", "level": 20,
                            "description_html": "<p>报告描述内容</p>"}]},
    )
    assert resp.status_code == 200, resp.text
    sa2 = resp.json()
    assert sa2["report_file_name"] == "原始报告.docx"
    assert [v["title"] for v in sa2["vuls"]] == ["报告导入漏洞"]
    # 漏洞摘要带所在层（网络层级列聚合展示依赖该字段）
    assert sa2["vuls"][0]["layer"] == 10

    # 附件下载
    resp = await client.get(f"/api/v1/spring-actions/{sa2['id']}/report", headers=auth)
    assert resp.status_code == 200

    # 收尾清理导入的漏洞（source=20 会污染后续 dashboard 来源筛选计数）
    imported_id = sa2["vuls"][0]["id"]
    resp = await client.post("/api/v1/vulns/batch-delete", headers=auth, json={"ids": [imported_id]})
    assert resp.status_code == 200

    # ---- 删除 ----
    for path in (f"/api/v1/remote-testings/{rt_id}",
                 f"/api/v1/testing-plans/{plan['id']}",
                 f"/api/v1/spring-actions/{sa['id']}",
                 f"/api/v1/spring-actions/{sa2['id']}"):
        resp = await client.delete(path, headers=auth)
        assert resp.status_code == 200


async def _get_plan(client: AsyncClient, auth: dict, plan_id: int) -> dict:
    resp = await client.get("/api/v1/testing-plans", headers=auth, params={"size": 100})
    return next(p for p in resp.json()["items"] if p["id"] == plan_id)


async def _list_plan_names(client: AsyncClient, auth: dict, params: dict) -> set[str]:
    resp = await client.get("/api/v1/testing-plans", headers=auth, params={"size": 100, **params})
    assert resp.status_code == 200, resp.text
    return {p["system_name"] for p in resp.json()["items"]}


async def test_testing_plan_filters(client: AsyncClient, auth: dict):
    """聚合筛选：多字段规则 + and/or 连接 + not 取反 + 派生/关联字段 + 非法入参。

    系统名/部门均使用唯一前缀，且断言限定在专属部门范围内，避免共享会话数据库
    中其他测试残留计划干扰精确集合比较。
    """
    DEPT = "筛选专用部门"
    A = "筛选专用系统-渗透A"
    B = "筛选专用系统-审计B"
    C = "筛选专用系统-渗透C"
    bodies = [
        {"system_name": A, "test_type": "渗透测试", "department": DEPT,
         "receive_time": "2026-01-01", "status": 10, "est_mandays": 3},
        {"system_name": B, "test_type": "代码审计", "department": "筛选专用部门B",
         "receive_time": "2026-02-01", "status": 20, "est_mandays": 5},
        {"system_name": C, "test_type": "渗透测试", "department": DEPT,
         "receive_time": "2026-03-01", "status": 60, "est_mandays": 2},
    ]
    plans = []
    for body in bodies:
        resp = await client.post("/api/v1/testing-plans", headers=auth, json=body)
        assert resp.status_code == 200, resp.text
        plans.append(resp.json())
    try:
        def rule(field, op, value=None, *, not_=False, connector="and"):
            return {"field": field, "op": op, "value": value, "not": not_, "connector": connector}

        def q(*rules):
            return {"filters": json.dumps({"rules": list(rules)})}

        dept = rule("department", "eq", DEPT)

        # 单条件：部门等于
        assert await _list_plan_names(client, auth, q(dept)) == {A, C}
        # 文本包含
        assert await _list_plan_names(client, auth, q(rule("system_name", "contains", "筛选专用系统-渗透"))) == {A, C}
        # AND 组合：部门=专属部门 且 状态=10
        assert await _list_plan_names(client, auth, q(dept, rule("status", "eq", 10))) == {A}
        # NOT 取反：部门=专属部门 且 状态≠10
        assert await _list_plan_names(client, auth, q(dept, rule("status", "ne", 10))) == {C}
        # OR 组合：规则间按顺序左结合（AND 优先），用唯一系统名避免受其他测试数据影响
        assert await _list_plan_names(
            client, auth,
            q(rule("system_name", "eq", A), rule("system_name", "eq", C, connector="or")),
        ) == {A, C}
        # 数字比较与区间（限定专属部门）
        assert await _list_plan_names(client, auth, q(dept, rule("est_mandays", "gte", 3))) == {A}
        assert await _list_plan_names(client, auth, q(dept, rule("est_mandays", "between", [2, 3]))) == {A, C}
        # 日期字符串上界（严格小于），空值被排除
        assert await _list_plan_names(client, auth, q(dept, rule("receive_time", "lt", "2026-02-01"))) == {A}
        # 空值筛选
        assert await _list_plan_names(client, auth, q(dept, rule("retest_notice_time", "is_empty"))) == {A, C}

        # 派生字段：自动生成工单ID精确匹配
        target = next(p for p in plans if p["system_name"] == A)
        assert target["ticket_id"]
        assert await _list_plan_names(client, auth, q(rule("ticket_id", "eq", target["ticket_id"]))) == {A}

        # 关联字段：认领后按测试人员筛选（限定专属部门）
        await client.post(f"/api/v1/testing-plans/{target['id']}/claim", headers=auth)
        assert await _list_plan_names(client, auth, q(dept, rule("testers", "contains", "admin"))) == {A}
        assert await _list_plan_names(client, auth, q(dept, rule("testers", "contains", "admin", not_=True))) == {C}

        # 非法字段 / 不匹配操作符返回 400
        resp = await client.get("/api/v1/testing-plans", headers=auth,
                                params=q(rule("not_exist", "eq", 1)))
        assert resp.status_code == 400
        resp = await client.get("/api/v1/testing-plans", headers=auth,
                                params=q(rule("system_name", "gt", 1)))
        assert resp.status_code == 400

        # stats 端点应用同一套筛选条件
        resp = await client.get("/api/v1/testing-plans/stats", headers=auth,
                                params=q(dept))
        assert resp.status_code == 200
        assert resp.json()["total_plans"] == 2
    finally:
        for p in plans:
            await client.delete(f"/api/v1/testing-plans/{p['id']}", headers=auth)


async def test_testing_plan_workflow(client: AsyncClient, auth: dict):
    """测试计划工作台：认领/退出、录入漏洞统计重算、报告关联三方状态联动。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "工作台系统", "test_type": "渗透测试", "department": "研发部"},
    )
    assert resp.status_code == 200, resp.text
    plan_id = resp.json()["id"]
    assert resp.json()["status"] == 10

    # 认领：当前用户加入测试人员，未测试自动进入初测中
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    assert resp.status_code == 200, resp.text
    plan = resp.json()
    assert plan["status"] == 20
    assert "admin" in [u["username"] for u in plan["testers"]]

    # 认领幂等：重复调用不重复添加
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    assert len(resp.json()["testers"]) == 1

    # 带计划录入漏洞：统计按等级自动重算
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "工作台资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={
            "asset_ids": [asset_id],
            "vulns": [
                {"title": "计划漏洞A", "level": 20, "testing_plan_id": plan_id},
                {"title": "计划漏洞B", "level": 30, "testing_plan_id": plan_id},
            ],
        },
    )
    assert resp.status_code == 200, resp.text
    vul_a, vul_b = [v["id"] for v in resp.json()]
    plan = await _get_plan(client, auth, plan_id)
    assert (plan["stat_high"], plan["stat_medium"]) == (1, 1)
    assert {v["id"] for v in plan["vuls"]} == {vul_a, vul_b}

    # 不存在的计划被拒绝
    resp = await client.post(
        "/api/v1/vulns", headers=auth, json={"title": "坏计划", "testing_plan_id": 999999},
    )
    assert resp.status_code == 400

    # 生成报告并关联计划：报告带 testing_plan_id，计划进入等待复测
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "工作台报告", "vul_ids": [vul_a, vul_b], "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    report_id = resp.json()["id"]
    assert resp.json()["testing_plan_id"] == plan_id
    plan = await _get_plan(client, auth, plan_id)
    assert plan["status"] == 30
    assert plan["first_test_done_time"]

    # 发起复测：计划进入复测中，漏洞进入复测中
    resp = await client.post(f"/api/v1/reports/{report_id}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    plan = await _get_plan(client, auth, plan_id)
    assert plan["status"] == 50
    for vid in (vul_a, vul_b):
        vul = (await client.get(f"/api/v1/vulns/{vid}", headers=auth)).json()
        assert vul["status"] == 55

    # 部分处理完：报告与计划均未完成
    resp = await client.post(
        f"/api/v1/vulns/{vul_a}/transition", headers=auth,
        json={"status": 60, "retest_html": "<p>复测通过</p>"},
    )
    assert resp.status_code == 200
    plan = await _get_plan(client, auth, plan_id)
    assert plan["status"] == 50

    # 全部已修复/已忽略：报告保持草稿（需求6），计划复测完成并记完成时间
    resp = await client.post(
        f"/api/v1/vulns/{vul_b}/transition", headers=auth, json={"status": 20},
    )
    assert resp.status_code == 200
    report = (await client.get(f"/api/v1/reports/{report_id}", headers=auth)).json()
    assert report["status"] == "draft"
    plan = await _get_plan(client, auth, plan_id)
    assert plan["status"] == 60
    assert plan["retest_done_time"]

    # 非管理员权限边界：状态修改与录入漏洞仅限认领者
    resp = await client.post(
        "/api/v1/roles", headers=auth,
        json={"name": "计划测试员", "permissions": ["special:manage", "vuln:submit"], "remark": ""},
    )
    role_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/users", headers=auth,
        json={"username": "plan_tester", "password": "Tester@123", "realname": "计划测试员",
              "email": "", "phone": "", "is_active": True, "role_id": role_id},
    )
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        "/api/v1/auth/login", data={"username": "plan_tester", "password": "Tester@123"},
    )
    auth2 = {"Authorization": f"Bearer {resp.json()['access_token']}"}

    plan_body = {k: plan[k] for k in (
        "system_name", "test_type", "department", "receive_time", "first_test_done_time",
        "status", "retest_notice_time", "retest_done_time",
        "stat_critical", "stat_high", "stat_medium", "stat_low", "target_urls", "detail",
    )}

    # 未认领：修改状态 403，录入漏洞 403，但基本信息编辑放行
    resp = await client.put(
        f"/api/v1/testing-plans/{plan_id}", headers=auth2, json={**plan_body, "status": 40},
    )
    assert resp.status_code == 403
    resp = await client.post(
        "/api/v1/vulns", headers=auth2, json={"title": "越权录入", "testing_plan_id": plan_id},
    )
    assert resp.status_code == 403
    resp = await client.put(
        f"/api/v1/testing-plans/{plan_id}", headers=auth2,
        json={**plan_body, "department": "研发一部"},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["department"] == "研发一部"

    # 认领后可修改状态；退出认领后再改状态回到 403
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth2)
    assert len(resp.json()["testers"]) == 2
    # 当前计划已处于复测完成(60)，认领后改回复测中(50)为合法流转
    resp = await client.put(
        f"/api/v1/testing-plans/{plan_id}", headers=auth2, json={**plan_body, "status": 50},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["status"] == 50
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/quit", headers=auth2)
    assert "plan_tester" not in [u["username"] for u in resp.json()["testers"]]
    # 退出认领后修改状态（60 为合法目标但无认领权限）回到 403
    resp = await client.put(
        f"/api/v1/testing-plans/{plan_id}", headers=auth2, json={**plan_body, "status": 60},
    )
    assert resp.status_code == 403


async def test_testing_plan_detail(client: AsyncClient, auth: dict):
    """单条计划详情端点：返回关联字段；不存在 404；无 special:manage 权限 403。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "详情端点系统", "test_type": "渗透测试"},
    )
    plan_id = resp.json()["id"]
    resp = await client.get(f"/api/v1/testing-plans/{plan_id}", headers=auth)
    assert resp.status_code == 200, resp.text
    detail = resp.json()
    assert detail["id"] == plan_id
    for key in ("testers", "vuls", "reports", "retest_rounds", "retest_round_count"):
        assert key in detail

    resp = await client.get("/api/v1/testing-plans/999999", headers=auth)
    assert resp.status_code == 404

    # 无 special:manage 权限的用户被拒绝
    resp = await client.post(
        "/api/v1/roles", headers=auth,
        json={"name": "无专项权限", "permissions": ["vuln:submit"], "remark": ""},
    )
    role_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/users", headers=auth,
        json={"username": "no_special", "password": "Tester@123", "realname": "无专项",
              "email": "", "phone": "", "is_active": True, "role_id": role_id},
    )
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        "/api/v1/auth/login", data={"username": "no_special", "password": "Tester@123"},
    )
    auth2 = {"Authorization": f"Bearer {resp.json()['access_token']}"}
    resp = await client.get(f"/api/v1/testing-plans/{plan_id}", headers=auth2)
    assert resp.status_code == 403


async def test_dashboard(client: AsyncClient, auth: dict):
    resp = await client.get("/api/v1/dashboard/stats", headers=auth)
    assert resp.status_code == 200
    stats = resp.json()
    assert stats["total_vulns"] >= 2
    assert "by_level" in stats and "trend" in stats
    assert len(stats["trend"]) == 12


async def test_permission_denied(client: AsyncClient):
    resp = await client.get("/api/v1/vulns")
    assert resp.status_code == 401


async def test_role_permission_catalog(client: AsyncClient, auth: dict):
    """权限目录接口：按功能模块分组返回，key 与扁平 PERMISSIONS 一致。"""
    resp = await client.get("/api/v1/roles/permissions/catalog", headers=auth)
    assert resp.status_code == 200, resp.text
    groups = resp.json()
    assert isinstance(groups, list) and groups
    keys = [it["key"] for g in groups for it in g["items"]]
    assert keys == [
        "dashboard:view", "asset:manage", "vuln:submit", "vuln:audit", "vuln:manage",
        "import:manage", "report:manage", "special:manage", "user:manage", "system:manage",
    ]
    for g in groups:
        assert g["group"]
        for it in g["items"]:
            assert it["label"] and "desc" in it

    # 扁平接口保持兼容
    flat = await client.get("/api/v1/roles/permissions", headers=auth)
    assert flat.status_code == 200
    assert flat.json() == keys


async def test_dict_options(client: AsyncClient, auth: dict):
    """测试类型字典：预设种子 + 下拉新增持久化 + 重名拒绝。"""
    resp = await client.get("/api/v1/dict/test_type", headers=auth)
    assert resp.status_code == 200
    names = [o["name"] for o in resp.json()]
    assert names[:5] == ["加电上线", "互联网自主测试", "办公网自主测试", "CHBN项目测试", "品质测评"]

    # 新增：去空格后入库
    resp = await client.post("/api/v1/dict/test_type", headers=auth, json={"name": " 红蓝对抗 "})
    assert resp.status_code == 200, resp.text
    assert resp.json()["name"] == "红蓝对抗"

    # 重名 / 空名拒绝
    resp = await client.post("/api/v1/dict/test_type", headers=auth, json={"name": "红蓝对抗"})
    assert resp.status_code == 400
    resp = await client.post("/api/v1/dict/test_type", headers=auth, json={"name": "  "})
    assert resp.status_code == 400

    resp = await client.get("/api/v1/dict/test_type", headers=auth)
    assert "红蓝对抗" in [o["name"] for o in resp.json()]

    # 未登录拒绝
    resp = await client.get("/api/v1/dict/test_type")
    assert resp.status_code == 401


async def test_group_create_by_special_manage(client: AsyncClient, auth: dict):
    """测试计划「所属部门」下拉新增：special:manage 角色可创建组织，重名拒绝。"""
    resp = await client.post(
        "/api/v1/roles", headers=auth,
        json={"name": "专项管理角色", "permissions": ["special:manage"], "remark": ""},
    )
    role_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/users", headers=auth,
        json={"username": "special_user", "password": "Sp@123456", "realname": "专项用户",
              "email": "", "phone": "", "is_active": True, "role_id": role_id},
    )
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        "/api/v1/auth/login", data={"username": "special_user", "password": "Sp@123456"},
    )
    auth_sp = {"Authorization": f"Bearer {resp.json()['access_token']}"}

    resp = await client.post("/api/v1/groups", headers=auth_sp, json={"name": "网络安全部", "remark": ""})
    assert resp.status_code == 200, resp.text
    resp = await client.post("/api/v1/groups", headers=auth_sp, json={"name": "网络安全部", "remark": ""})
    assert resp.status_code == 400
    resp = await client.get("/api/v1/groups", headers=auth_sp)
    assert "网络安全部" in [g["name"] for g in resp.json()]


async def test_group_owner_fields(client: AsyncClient, auth: dict):
    """组织负责人信息：改由「组织成员」管理（姓名/电话/邮箱 CRUD）。"""
    resp = await client.post(
        "/api/v1/groups", headers=auth,
        json={"name": "负责人测试部", "remark": ""},
    )
    assert resp.status_code == 200, resp.text
    group = resp.json()
    assert "owner_name" not in group  # 单字段 owner 已移除，改为组织成员承载

    # 添加组织成员
    resp = await client.post(
        f"/api/v1/groups/{group['id']}/members", headers=auth,
        json={"name": "赵六", "phone": "13700000000", "email": "zhaoliu@example.com"},
    )
    assert resp.status_code == 200, resp.text
    member = resp.json()
    assert member["name"] == "赵六"
    assert member["email"] == "zhaoliu@example.com"

    # 更新组织成员
    resp = await client.put(
        f"/api/v1/groups/{group['id']}/members/{member['id']}", headers=auth,
        json={"name": "钱七", "phone": "", "email": ""},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["name"] == "钱七"
    assert resp.json()["phone"] == ""

    # 删除组织成员
    members = (await client.get(f"/api/v1/groups/{group['id']}/members", headers=auth)).json()
    assert [m["name"] for m in members] == ["钱七"]
    resp = await client.delete(
        f"/api/v1/groups/{group['id']}/members/{member['id']}", headers=auth,
    )
    assert resp.status_code == 200
    members = (await client.get(f"/api/v1/groups/{group['id']}/members", headers=auth)).json()
    assert members == []


async def test_retest_round_tracking(client: AsyncClient, auth: dict):
    """复测轮次统计：手动流转记轮、重复流转不重复计数、报告发起复测强制开新轮、闭环打完成点。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "轮次系统", "department": "轮次部门"},
    )
    plan_id = resp.json()["id"]
    assert resp.json()["retest_round_count"] == 0
    await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)

    plan = await _get_plan(client, auth, plan_id)
    body = {k: plan[k] for k in (
        "system_name", "test_type", "department", "receive_time", "first_test_done_time",
        "status", "retest_notice_time", "retest_done_time",
        "stat_critical", "stat_high", "stat_medium", "stat_low", "target_urls", "detail",
    )}

    # 手动流转到复测中：记第 1 轮（需经初测完成 30，再进入复测中 50）
    resp = await client.put(
        f"/api/v1/testing-plans/{plan_id}", headers=auth, json={**body, "status": 30},
    )
    assert resp.status_code == 200, resp.text
    resp = await client.put(
        f"/api/v1/testing-plans/{plan_id}", headers=auth, json={**body, "status": 50},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["retest_round_count"] == 1
    assert resp.json()["retest_rounds"][0]["done_time"] is None

    # 反复流转（50→30→50）：已有进行中轮次不重复计数
    await client.put(f"/api/v1/testing-plans/{plan_id}", headers=auth, json={**body, "status": 30})
    resp = await client.put(
        f"/api/v1/testing-plans/{plan_id}", headers=auth, json={**body, "status": 50},
    )
    assert resp.json()["retest_round_count"] == 1

    # 报告实际发起复测：结束上一轮并强制开新轮
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "轮次资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={
            "asset_ids": [asset_id],
            "vulns": [
                {"title": "轮次漏洞A", "level": 20, "testing_plan_id": plan_id},
                {"title": "轮次漏洞B", "level": 30, "testing_plan_id": plan_id},
            ],
        },
    )
    vul_a, vul_b = [v["id"] for v in resp.json()]
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "轮次报告", "vul_ids": [vul_a, vul_b], "testing_plan_id": plan_id},
    )
    report_id = resp.json()["id"]
    resp = await client.post(f"/api/v1/reports/{report_id}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    plan = await _get_plan(client, auth, plan_id)
    assert plan["retest_round_count"] == 2
    rounds = plan["retest_rounds"]
    assert rounds[0]["round_no"] == 1 and rounds[0]["done_time"] is not None
    assert rounds[1]["round_no"] == 2 and rounds[1]["done_time"] is None
    assert "轮次报告" in rounds[1]["source"]

    # 再次发起复测：已处于复测中的漏洞同样视为复测流程已发起（force=True），强制开新一轮
    resp = await client.post(f"/api/v1/reports/{report_id}/retest", headers=auth)
    plan = await _get_plan(client, auth, plan_id)
    assert plan["retest_round_count"] == 3

    # 全部闭环：计划复测完成，当前轮打完成点
    await client.post(
        f"/api/v1/vulns/{vul_a}/transition", headers=auth,
        json={"status": 60, "retest_html": "<p>复测通过</p>"},
    )
    await client.post(f"/api/v1/vulns/{vul_b}/transition", headers=auth, json={"status": 20})
    plan = await _get_plan(client, auth, plan_id)
    assert plan["status"] == 60
    assert plan["retest_rounds"][1]["done_time"] is not None


async def test_delete_retest_report_rolls_back_round(client: AsyncClient, auth: dict):
    """删除新发起的复测记录后复测轮数正确回退：删除复测报告移除对应轮次；
    若删除的是 force 发起的新轮，上一轮恢复进行中状态，保持轮次状态与计划状态一致。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "轮次回退系统", "department": "轮次回退部门"},
    )
    assert resp.status_code == 200, resp.text
    plan_id = resp.json()["id"]
    await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "轮次回退资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={"asset_ids": [asset_id], "vulns": [
            {"title": "回退漏洞A", "level": 20, "testing_plan_id": plan_id},
        ]},
    )
    vul_a = resp.json()[0]["id"]

    # 生成初测报告 → 漏洞进入修复中，计划进入等待复测
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "轮次回退报告", "vul_ids": [vul_a], "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    src_report = resp.json()["id"]

    # 发起复测 → 生成复测报告 R1、记录第 1 轮
    resp = await client.post(f"/api/v1/reports/{src_report}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    r1_id = resp.json()["id"]
    plan = await _get_plan(client, auth, plan_id)
    assert plan["retest_round_count"] == 1

    # 删除新发起的复测报告 R1 → 对应轮次回退，复测轮数归零，仅剩初测报告
    resp = await client.delete(f"/api/v1/reports/{r1_id}", headers=auth)
    assert resp.status_code == 200, resp.text
    plan = await _get_plan(client, auth, plan_id)
    assert plan["retest_round_count"] == 0
    assert len(plan["reports"]) == 1

    # 再次发起复测 → 重新开第 1 轮（进行中）
    resp = await client.post(f"/api/v1/reports/{src_report}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    r2_id = resp.json()["id"]
    plan = await _get_plan(client, auth, plan_id)
    assert plan["retest_round_count"] == 1
    assert plan["retest_rounds"][0]["done_time"] is None

    # 再次发起复测（漏洞已处于复测中）→ force 结束第 1 轮并开第 2 轮
    resp = await client.post(f"/api/v1/reports/{src_report}/retest", headers=auth)
    assert resp.status_code == 200, resp.text
    r3_id = resp.json()["id"]
    plan = await _get_plan(client, auth, plan_id)
    assert plan["retest_round_count"] == 2
    assert plan["retest_rounds"][0]["done_time"] is not None
    assert plan["retest_rounds"][1]["done_time"] is None

    # 删除最新复测报告 R3 → 轮数回退为 1，且上一轮恢复进行中
    resp = await client.delete(f"/api/v1/reports/{r3_id}", headers=auth)
    assert resp.status_code == 200, resp.text
    plan = await _get_plan(client, auth, plan_id)
    assert plan["retest_round_count"] == 1
    assert plan["retest_rounds"][0]["done_time"] is None

    # 删除初测报告不影响轮次（非复测报告不关联任何轮次）
    resp = await client.delete(f"/api/v1/reports/{src_report}", headers=auth)
    assert resp.status_code == 200, resp.text
    plan = await _get_plan(client, auth, plan_id)
    assert plan["retest_round_count"] == 1
    assert plan["retest_rounds"][0]["done_time"] is None


async def test_dashboard_by_department(client: AsyncClient, auth: dict):
    """安全态势部门维度：提测次数 / 发现漏洞（含手填补充） / 修复率 / 占用人天（实际人天求和）。"""
    # 无关联漏洞的计划：发现数取手填统计，修复率为空；占用人天取实际人天
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "看板系统", "department": "看板部门",
              "stat_high": 2, "stat_low": 1, "actual_mandays": 3.5},
    )
    assert resp.status_code == 200, resp.text

    stats = (await client.get("/api/v1/dashboard/stats", headers=auth)).json()
    assert "by_department" in stats
    dept = next(d for d in stats["by_department"] if d["department"] == "看板部门")
    assert dept["plans"] == 1
    assert dept["vulns"] == 3
    assert dept["fixed"] == 0
    assert dept["fix_rate"] is None
    assert dept["mandays"] == 3.5
    # 高危及以上取手填严重(0)+高危(2)；手填统计无状态概念，全部计为未闭环
    assert dept["high"] == 2
    assert dept["open"] == 3

    # 有关联漏洞的计划（test_retest_round_tracking：高危A 已修复 + 中危B 已忽略）
    dept2 = next(d for d in stats["by_department"] if d["department"] == "轮次部门")
    assert dept2["plans"] == 1
    assert dept2["vulns"] == 2
    assert dept2["fixed"] == 1
    assert dept2["fix_rate"] == 50.0
    assert isinstance(dept2["mandays"], (int, float))
    # 高危及以上仅高危A；未闭环 = 2 − 已修复1 − 已忽略1 = 0
    assert dept2["high"] == 1
    assert dept2["open"] == 0


async def test_vuln_stats_by_asset(client: AsyncClient, auth: dict):
    """/vulns/stats 按资产分组：by_asset 存在、计数正确、部门筛选联动。"""
    a1 = (await client.post(
        "/api/v1/assets", headers=auth,
        json={"name": "统计资产甲", "department": "统计部门A"},
    )).json()
    a2 = (await client.post(
        "/api/v1/assets", headers=auth,
        json={"name": "统计资产乙", "department": "统计部门B"},
    )).json()
    v1 = v2 = None
    try:
        v1 = (await client.post(
            "/api/v1/vulns", headers=auth,
            json={"title": "统计漏洞一", "level": 20, "asset_ids": [a1["id"], a2["id"]]},
        )).json()
        v2 = (await client.post(
            "/api/v1/vulns", headers=auth,
            json={"title": "统计漏洞二", "level": 30, "asset_ids": [a2["id"]]},
        )).json()

        # 全量统计：同一漏洞关联多资产时按关联关系在各资产分组重复计入
        # （session 级共享数据库，不依赖全局 total，只校验本次创建的资产）
        stats = (await client.get("/api/v1/vulns/stats", headers=auth)).json()
        assert "by_asset" in stats
        asset_count = {r["asset_id"]: r["count"] for r in stats["by_asset"]}
        assert asset_count[a1["id"]] == 1
        assert asset_count[a2["id"]] == 2
        # departments：资产部门去重列表（含全部资产部门，供组合工具选部门）
        assert "departments" in stats
        assert {"统计部门A", "统计部门B"} <= set(stats["departments"])

        # 部门筛选联动：该部门漏洞计入（跨部门关联的乙也会带出，前端以资产列表为选项源过滤）
        dept_stats = (await client.get(
            "/api/v1/vulns/stats", headers=auth, params={"department": "统计部门A"},
        )).json()
        dept_asset_count = {r["asset_id"]: r["count"] for r in dept_stats["by_asset"]}
        assert dept_asset_count[a1["id"]] == 1
    finally:
        for v in (v1, v2):
            if v:
                await client.delete(f"/api/v1/vulns/{v['id']}", headers=auth)
        await client.delete(f"/api/v1/assets/{a2['id']}", headers=auth)
        await client.delete(f"/api/v1/assets/{a1['id']}", headers=auth)


async def test_testing_plan_filter_stats_export(client: AsyncClient, auth: dict):
    """测试计划筛选（类型/部门/时间范围）、多维度统计与 Excel 双 sheet 导出。"""
    from openpyxl import load_workbook

    tag = "筛选统计专用部门"
    # 三条计划：不同状态 / 类型 / 接收时间，便于区分筛选与统计口径
    plans = [
        {"system_name": "过滤系统A", "department": tag, "test_type": "黑盒测试",
         "receive_time": "2026-01-15", "status": 20},   # 初测中
        {"system_name": "过滤系统B", "department": tag, "test_type": "白盒测试",
         "receive_time": "2026-03-20", "status": 60},   # 复测完成
        {"system_name": "过滤系统C", "department": tag, "test_type": "黑盒测试",
         "receive_time": "2026-06-10", "status": 10},   # 未测试
    ]
    for p in plans:
        resp = await client.post("/api/v1/testing-plans", headers=auth, json=p)
        assert resp.status_code == 200, resp.text

    # 部门 + 类型精确筛选
    resp = await client.get(
        "/api/v1/testing-plans", headers=auth,
        params={"department": tag, "test_type": "黑盒测试"},
    )
    names = {i["system_name"] for i in resp.json()["items"]}
    assert names == {"过滤系统A", "过滤系统C"}

    # 时间范围筛选（按需求接收时间）
    resp = await client.get(
        "/api/v1/testing-plans", headers=auth,
        params={"department": tag, "receive_from": "2026-02-01", "receive_to": "2026-04-01"},
    )
    names = {i["system_name"] for i in resp.json()["items"]}
    assert names == {"过滤系统B"}

    # 统计口径：本部门下 3 条计划
    resp = await client.get("/api/v1/testing-plans/stats", headers=auth, params={"department": tag})
    assert resp.status_code == 200, resp.text
    stats = resp.json()
    assert stats["total_plans"] == 3
    assert stats["retest_done_plans"] == 1          # 仅 B 为复测完成
    assert stats["first_test_count"] == 2           # A(20)+B(60) 达到初测中及以上
    assert stats["total_test_count"] == stats["first_test_count"] + stats["retest_count"]
    status_counts = {r["status"]: r["count"] for r in stats["by_status"]}
    assert status_counts == {10: 1, 20: 1, 60: 1}
    assert isinstance(stats["vulns_by_month"], list)

    # pending 待办流程筛选：仅 未测试(10)/初测中(20)/复测中(50)，复测完成(60) 被排除
    resp = await client.get(
        "/api/v1/testing-plans", headers=auth,
        params={"department": tag, "pending": True},
    )
    names = {i["system_name"] for i in resp.json()["items"]}
    assert names == {"过滤系统A", "过滤系统C"}
    assert all(i["status"] in (10, 20, 50) for i in resp.json()["items"])

    # pending 与单状态筛选同时传参时按 AND 处理：无交集则空结果
    resp = await client.get(
        "/api/v1/testing-plans", headers=auth,
        params={"department": tag, "pending": True, "status": 60},
    )
    assert resp.json()["total"] == 0

    # pending 统计口径与列表一致：本部门待办 2 条
    pending_stats = (await client.get(
        "/api/v1/testing-plans/stats", headers=auth, params={"department": tag, "pending": True},
    )).json()
    assert pending_stats["total_plans"] == 2
    assert pending_stats["retest_done_plans"] == 0

    # 空结果集：全 0，不报错
    empty = (await client.get(
        "/api/v1/testing-plans/stats", headers=auth, params={"department": "不存在的部门XYZ"},
    )).json()
    assert empty["total_plans"] == 0
    assert empty["total_test_count"] == 0

    # 导出：双 sheet，可被 openpyxl 回读
    resp = await client.get("/api/v1/testing-plans/export", headers=auth, params={"department": tag})
    assert resp.status_code == 200
    assert resp.content[:2] == b"PK"
    wb = load_workbook(BytesIO(resp.content))
    assert "渗透测试工单" in wb.sheetnames
    assert "统计汇总" in wb.sheetnames
    detail_rows = list(wb["渗透测试工单"].iter_rows(values_only=True))
    assert len(detail_rows) == 4  # 表头 + 3 行
    assert detail_rows[0][1] == "渗透测试工单名称"  # 表头第二列与 PLAN_EXCEL_HEADERS 一致

    # 导出同样支持 pending 筛选：仅 2 条待办计划
    resp = await client.get(
        "/api/v1/testing-plans/export", headers=auth,
        params={"department": tag, "pending": True},
    )
    assert resp.status_code == 200
    wb2 = load_workbook(BytesIO(resp.content))
    pending_rows = list(wb2["渗透测试工单"].iter_rows(values_only=True))
    assert len(pending_rows) == 3  # 表头 + 2 行
    assert "统计汇总" in wb2.sheetnames


async def test_testing_plan_mandays_and_reports(client: AsyncClient, auth: dict):
    """人天字段、人天统计（剩余预估人天=未测试计划预估人天之和）与计划反向展示关联报告。"""
    from openpyxl import load_workbook

    tag = "人天统计专用部门"
    plans = [
        {"system_name": "人天系统A", "department": tag, "status": 10,
         "est_mandays": 3, "actual_mandays": 0},     # 未测试 → 计入剩余
        {"system_name": "人天系统B", "department": tag, "status": 20,
         "est_mandays": 5, "actual_mandays": 2.5},
        {"system_name": "人天系统C", "department": tag, "status": 60,
         "est_mandays": 2, "actual_mandays": 4},
    ]
    plan_ids = []
    for p in plans:
        resp = await client.post("/api/v1/testing-plans", headers=auth, json=p)
        assert resp.status_code == 200, resp.text
        assert resp.json()["est_mandays"] == p["est_mandays"]
        assert resp.json()["actual_mandays"] == p["actual_mandays"]
        plan_ids.append(resp.json()["id"])

    # 人天统计：总预估/总实际/剩余预估（仅未测试状态）
    stats = (await client.get(
        "/api/v1/testing-plans/stats", headers=auth, params={"department": tag},
    )).json()
    assert stats["est_mandays_total"] == 10
    assert stats["actual_mandays_total"] == 6.5
    assert stats["remaining_est_mandays"] == 3

    # 导出明细含人天列，汇总含人天指标
    resp = await client.get("/api/v1/testing-plans/export", headers=auth, params={"department": tag})
    wb = load_workbook(BytesIO(resp.content))
    detail_rows = list(wb["渗透测试工单"].iter_rows(values_only=True))
    header = list(detail_rows[0])
    assert "预估人天" in header and "实际人天" in header
    est_col = header.index("预估人天")
    assert {r[est_col] for r in detail_rows[1:]} == {3, 5, 2}
    summary = {r[0]: r[1] for r in wb["统计汇总"].iter_rows(values_only=True) if r[0]}
    assert summary["预估人天总计"] == 10
    assert summary["实际人天总计"] == 6.5
    assert summary["剩余预估人天（未测试）"] == 3

    # 计划反向展示关联报告：录入漏洞→生成报告后 reports 列表可见
    plan_b = plan_ids[1]
    # 严格认领：需先认领计划才能建计划漏洞
    await client.post(f"/api/v1/testing-plans/{plan_b}/claim", headers=auth)
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "人天反向资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={"title": "人天反向漏洞", "level": 20, "asset_ids": [asset_id],
              "testing_plan_id": plan_b},
    )
    vul_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "人天反向报告", "vul_ids": [vul_id], "testing_plan_id": plan_b},
    )
    assert resp.status_code == 200, resp.text
    report_id = resp.json()["id"]

    plan = await _get_plan(client, auth, plan_b)
    assert [r["id"] for r in plan["reports"]] == [report_id]
    assert plan["reports"][0]["title"] == "人天反向报告"
    assert plan["reports"][0]["status"] == "draft"


async def test_testing_plan_actual_mandays_override(client: AsyncClient, auth: dict):
    """实际人天修正：修正后不再被初测报告自动覆盖，取消修正后恢复自动计算。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "修正人天系统", "department": "人天统计专用部门", "est_mandays": 2},
    )
    assert resp.status_code == 200, resp.text
    plan_id = resp.json()["id"]
    await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    asset_id = (await client.post(
        "/api/v1/assets", headers=auth, json={"name": "修正人天资产"},
    )).json()["id"]
    vul_id = (await client.post(
        "/api/v1/vulns", headers=auth,
        json={"title": "修正人天漏洞", "level": 20, "asset_ids": [asset_id],
              "testing_plan_id": plan_id},
    )).json()["id"]
    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "修正人天报告", "vul_ids": [vul_id], "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    report = resp.json()
    report_id = report["id"]

    # 给初测报告填写测试周期：3 天 → 计划实际人天自动 = 3
    save_body = {
        "title": report["title"], "project_name": report["project_name"],
        "author": report["author"], "revision": report["revision"],
        "test_start": "2026-08-01", "test_end": "2026-08-03",
        "sections": report["sections"],
    }
    resp = await client.put(f"/api/v1/reports/{report_id}", headers=auth, json=save_body)
    assert resp.status_code == 200, resp.text
    report = resp.json()
    assert report["actual_mandays"] == 3
    plan = await _get_plan(client, auth, plan_id)
    assert plan["actual_mandays"] == 3
    assert plan["actual_mandays_override"] is False

    # 修正：手动输入 9.5，保存后不再被初测报告自动覆盖
    body = plan
    for k in ("id", "ticket_seq", "ticket_id", "testers", "vuls", "reports",
              "retest_rounds", "retest_round_count", "create_time", "update_time"):
        body.pop(k, None)
    body["actual_mandays"] = 9.5
    body["actual_mandays_override"] = True
    resp = await client.put(f"/api/v1/testing-plans/{plan_id}", headers=auth, json=body)
    assert resp.status_code == 200, resp.text
    assert resp.json()["actual_mandays"] == 9.5
    assert resp.json()["actual_mandays_override"] is True

    # 初测报告测试周期变更为 5 天，refresh_mandays 因修正标志跳过，修正值保持不变
    save_body["revision"] = report["revision"]
    save_body["test_end"] = "2026-08-05"
    resp = await client.put(f"/api/v1/reports/{report_id}", headers=auth, json=save_body)
    assert resp.status_code == 200, resp.text
    plan = await _get_plan(client, auth, plan_id)
    assert plan["actual_mandays"] == 9.5

    # 取消修正：恢复为初测报告计算的 5 天
    body["actual_mandays"] = 99  # 该值会被自动计算覆盖
    body["actual_mandays_override"] = False
    resp = await client.put(f"/api/v1/testing-plans/{plan_id}", headers=auth, json=body)
    assert resp.status_code == 200, resp.text
    plan = resp.json()
    assert plan["actual_mandays_override"] is False
    assert plan["actual_mandays"] == 5


async def test_testing_plan_excel_import(client: AsyncClient, auth: dict):
    """测试计划 Excel 导入：模板下载、按 ID 更新、无 ID 新增、非法行报错。"""
    from openpyxl import Workbook

    # 模板下载
    resp = await client.get("/api/v1/testing-plans/import/template", headers=auth)
    assert resp.status_code == 200
    assert resp.content[:2] == b"PK"

    # 先建一条计划供导入更新
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "导入前系统", "department": "导入专用部门", "est_mandays": 1},
    )
    exist_id = resp.json()["id"]

    # 表头与模板一致（PLAN_EXCEL_HEADERS 20 列，含「渗透测试工单名称」列）
    headers_row = ["ID", "渗透测试工单名称", "测试系统", "测试类型", "所属部门", "工单ID",
                   "工单提起时间", "状态", "测试人员", "需求接收", "初测完成", "复测通知",
                   "复测完成", "预估人天", "实际人天", "超危数", "高危数", "中危数",
                   "低危数", "复测轮数"]
    wb = Workbook()
    ws = wb.active
    ws.append(headers_row)
    # 无 ID → 新增（测试人员按用户名匹配 admin）
    ws.append(["", "", "导入新增系统", "渗透测试", "导入专用部门", "", "",
               "初测中", "admin", "2025-11-01", "2025-11-05", "", "", 3.5, 1, 1, 2, 0, 0, 0])
    # 有 ID → 更新同一条计划
    ws.append([exist_id, "", "导入后系统", "白盒测试", "导入专用部门", "", "",
               "复测完成", "", "2025-10-01", "", "", "2025-12-31", 6, 5.5, 0, 0, 0, 0, 0])
    # 缺测试系统 → 失败
    ws.append(["", "", "", "黑盒测试", "导入专用部门", "", "", "", "", "", "", "", "",
               0, 0, 0, 0, 0, 0, 0])
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    resp = await client.post(
        "/api/v1/testing-plans/import", headers=auth,
        files={"file": ("plans.xlsx", buf,
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert resp.status_code == 200, resp.text
    result = resp.json()
    assert result["total"] == 3
    assert result["created"] == 1
    assert result["updated"] == 1
    assert result["failed"] == 1
    assert "测试系统为必填项" in result["errors"][0]

    # 新增行校验：状态/人天/测试人员匹配
    resp = await client.get(
        "/api/v1/testing-plans", headers=auth, params={"search": "导入新增系统"},
    )
    created = resp.json()["items"][0]
    assert created["status"] == 20
    assert created["est_mandays"] == 3.5
    assert created["stat_critical"] == 1 and created["stat_high"] == 2
    assert "admin" in [u["username"] for u in created["testers"]]

    # 更新行校验：同 ID 字段被覆盖
    updated = await _get_plan(client, auth, exist_id)
    assert updated["system_name"] == "导入后系统"
    assert updated["status"] == 60
    assert updated["est_mandays"] == 6
    assert updated["actual_mandays"] == 5.5

    # 非 xlsx 拒绝
    resp = await client.post(
        "/api/v1/testing-plans/import", headers=auth,
        files={"file": ("bad.txt", b"hello", "text/plain")},
    )
    assert resp.status_code == 400


async def test_dashboard_event_filters(client: AsyncClient, auth: dict):
    """安全态势按事件多维筛选：部门/等级/来源/时间范围。"""
    dept = "态势筛选部门"
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "态势筛选系统", "department": dept},
    )
    plan_id = resp.json()["id"]
    # 严格认领：需先认领计划才能批量建计划漏洞
    await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "态势筛选资产"})
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={
            "asset_ids": [asset_id],
            "vulns": [
                {"title": "态势高危漏洞", "level": 20, "source": 10, "testing_plan_id": plan_id},
                {"title": "态势中危漏洞", "level": 30, "source": 20, "testing_plan_id": plan_id},
            ],
        },
    )
    assert resp.status_code == 200, resp.text

    # 部门筛选：仅统计该部门计划关联的 2 条漏洞
    stats = (await client.get(
        "/api/v1/dashboard/stats", headers=auth, params={"department": dept},
    )).json()
    assert stats["total_vulns"] == 2
    assert {x["name"] for x in stats["by_level"]} == {"高危", "中危"}
    assert [d["department"] for d in stats["by_department"]] == [dept]

    # 部门 + 等级
    stats = (await client.get(
        "/api/v1/dashboard/stats", headers=auth,
        params={"department": dept, "level": 20},
    )).json()
    assert stats["total_vulns"] == 1

    # 单独录入的漏洞（未关联工单）：来源可选，且关联工单的漏洞来源恒为工单（source 强制 0）
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "态势来源资产"})
    src_asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={"title": "态势来源筛选漏洞", "level": 30, "source": 20, "asset_ids": [src_asset_id]},
    )
    assert resp.status_code == 200, resp.text

    # 来源筛选：仅命中单独录入且来源=20（春耕行动）的漏洞；关联工单的漏洞来源被强制为 0 不参与
    stats = (await client.get(
        "/api/v1/dashboard/stats", headers=auth,
        params={"source": 20},
    )).json()
    assert stats["total_vulns"] == 1

    # 时间范围：未来区间无数据；包含今天则命中
    stats = (await client.get(
        "/api/v1/dashboard/stats", headers=auth,
        params={"department": dept, "date_from": "2099-01-01", "date_to": "2099-12-31"},
    )).json()
    assert stats["total_vulns"] == 0
    stats = (await client.get(
        "/api/v1/dashboard/stats", headers=auth,
        params={"department": dept, "date_from": "2000-01-01", "date_to": "2099-12-31"},
    )).json()
    assert stats["total_vulns"] == 2


async def test_retest_record_sync_to_vul(client: AsyncClient, auth: dict):
    """复测记录增/改/删同步聚合到 Vul.retest_html（详情页/报告读取口径）。"""
    resp = await client.post(
        "/api/v1/vulns", headers=auth, json={"title": "复测同步漏洞", "level": 20},
    )
    vul_id = resp.json()["id"]

    # 新增一条记录：retest_html 同步为记录内容（标题格式「复测记录yymmdd」）
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/retests", headers=auth,
        json={"content_html": "<p>第一次复测仍存在</p>", "content_json": None},
    )
    assert resp.status_code == 200, resp.text
    rec1_id = resp.json()["id"]
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert re.search(r"复测记录\d{6}：", vul["retest_html"])
    assert "第一次复测仍存在" in vul["retest_html"]

    # 第二条记录：同日新增，标题追加 -1 后缀
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/retests", headers=auth,
        json={"content_html": "<p>第二次复测已修复</p>", "content_json": None},
    )
    rec2_id = resp.json()["id"]
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert re.search(r"复测记录\d{6}：", vul["retest_html"])
    assert re.search(r"复测记录\d{6}-1：", vul["retest_html"])
    assert "第一次复测仍存在" in vul["retest_html"]
    assert "第二次复测已修复" in vul["retest_html"]

    # 自定义标题优先于自动日期标题（手动编辑复测标题对应实际复测时间）
    resp = await client.put(
        f"/api/v1/vulns/{vul_id}/retests/{rec2_id}", headers=auth,
        json={"title": "复测记录250815", "content_html": "<p>第二次复测已修复</p>", "content_json": None},
    )
    assert resp.status_code == 200, resp.text
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert "<strong>复测记录250815：</strong>" in vul["retest_html"]
    assert "第二次复测已修复" in vul["retest_html"]

    # 清空自定义标题后回退为自动日期标题（同日追加 -1 后缀）
    resp = await client.put(
        f"/api/v1/vulns/{vul_id}/retests/{rec2_id}", headers=auth,
        json={"title": None, "content_html": "<p>第二次复测已修复</p>", "content_json": None},
    )
    assert resp.status_code == 200, resp.text
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert re.search(r"复测记录\d{6}-1：", vul["retest_html"])
    assert "<strong>复测记录250815：</strong>" not in vul["retest_html"]

    # 更新记录：聚合内容跟随变化
    resp = await client.put(
        f"/api/v1/vulns/{vul_id}/retests/{rec2_id}", headers=auth,
        json={"content_html": "<p>第二次复测部分修复</p>", "content_json": None},
    )
    assert resp.status_code == 200, resp.text
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert "第二次复测部分修复" in vul["retest_html"]
    assert "第二次复测已修复" not in vul["retest_html"]

    # 删除一条：回到单条内容（标题不带 -N 后缀）
    resp = await client.delete(f"/api/v1/vulns/{vul_id}/retests/{rec2_id}", headers=auth)
    assert resp.status_code == 200
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert re.search(r"复测记录\d{6}：", vul["retest_html"])
    assert not re.search(r"复测记录\d{6}-\d：", vul["retest_html"])
    assert "第一次复测仍存在" in vul["retest_html"]

    # 全部删除：retest_html 清空
    resp = await client.delete(f"/api/v1/vulns/{vul_id}/retests/{rec1_id}", headers=auth)
    assert resp.status_code == 200
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert not vul["retest_html"]


async def test_retest_failed_back_to_fixing(client: AsyncClient, auth: dict):
    """复测未通过打回：50 → 55 → 50 后 status=50 且 is_retest=true（前端据此展示"复测未通过"）。"""
    resp = await client.post(
        "/api/v1/vulns", headers=auth, json={"title": "复测打回漏洞", "level": 20},
    )
    vul_id = resp.json()["id"]

    for status in (50, 55):
        resp = await client.post(
            f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": status},
        )
        assert resp.status_code == 200, resp.text
    # 复测未通过回修复中：必须填写复测详情
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/transition", headers=auth,
        json={"status": 50, "retest_html": "<p>复测发现仍可利用</p>"},
    )
    assert resp.status_code == 200, resp.text

    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert vul["status"] == 50
    assert vul["is_retest"] is True


async def test_vuln_status_retest_guard(client: AsyncClient, auth: dict):
    """复测结论校验：修复中不可直接变已修复；复测中→已修复/复测未通过必须填写复测内容。"""
    resp = await client.post("/api/v1/vulns", headers=auth, json={"title": "复测守卫漏洞", "level": 20})
    vul_id = resp.json()["id"]
    # 进入修复中
    resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": 50})
    assert resp.status_code == 200

    # 报告编辑页下拉（fields 接口）从修复中直接改已修复 → 拒绝（必须经过复测中）
    resp = await client.patch(f"/api/v1/vulns/{vul_id}/fields", headers=auth, json={"status": 60})
    assert resp.status_code == 400

    # 进入复测中
    resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": 55})
    assert resp.status_code == 200

    # 复测中→已修复但未填写复测内容 → 拒绝
    resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": 60})
    assert resp.status_code == 400

    # 复测中→复测未通过（回修复中）但未填写复测详情 → 拒绝
    resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": 50})
    assert resp.status_code == 400

    # 填写复测内容后可正常闭环
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/transition", headers=auth,
        json={"status": 60, "retest_html": "<p>复测已修复</p>"},
    )
    assert resp.status_code == 200
    assert resp.json()["status"] == 60


async def test_second_round_retest_requires_new_record(client: AsyncClient, auth: dict):
    """二轮复测不得误用首轮历史复测记录放行：本轮无新记录时禁止切换为已修复/复测未通过。"""
    resp = await client.post("/api/v1/vulns", headers=auth, json={"title": "二轮复测守卫漏洞", "level": 20})
    vul_id = resp.json()["id"]

    # 首轮：50 → 55 → 新增复测记录并闭环为已修复
    for status in (50, 55):
        resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": status})
        assert resp.status_code == 200, resp.text
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/retests", headers=auth,
        json={"content_html": "<p>首轮复测通过</p>", "status": 60},
    )
    assert resp.status_code == 200, resp.text

    # 二轮：已修复重新发起复测（FIXED→RETESTING），未新增任何记录
    resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": 55})
    assert resp.status_code == 200, resp.text
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert vul["status"] == 55
    assert vul["retest_html"]  # 首轮历史记录仍聚合在 retest_html 中（旧校验会误放行）

    # 本轮无新记录：直接切换为已修复 → 拒绝
    resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": 60})
    assert resp.status_code == 400

    # 本轮无新记录：复测未通过回修复中 → 拒绝
    resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": 50})
    assert resp.status_code == 400

    # 本轮无新记录：编辑页直接点选为已修复（set_status 路径）→ 拒绝
    resp = await client.patch(f"/api/v1/vulns/{vul_id}/fields", headers=auth, json={"status": 60})
    assert resp.status_code == 400

    # 本轮新增复测记录后再流转 → 放行
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/retests", headers=auth,
        json={"content_html": "<p>二轮复测通过</p>", "status": 60},
    )
    assert resp.status_code == 200, resp.text
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert vul["status"] == 60


async def test_ticket_seq_increment_no_reuse(client: AsyncClient, auth: dict):
    """工单ID自动分配采用「当日最大编号+1」：删除/释放的编号不复用，仅手动可选用。"""
    body = {"department": "递增部门", "receive_time": "2026-01-01"}

    resp = await client.post(
        "/api/v1/testing-plans", headers=auth, json={**body, "system_name": "递增系统A"},
    )
    assert resp.status_code == 200, resp.text
    plan_a = resp.json()
    assert plan_a["ticket_id"] == "20260101-1"

    resp = await client.post(
        "/api/v1/testing-plans", headers=auth, json={**body, "system_name": "递增系统B"},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["ticket_id"] == "20260101-2"

    # 删除占用 20260101-1 的计划后，自动分配不复用空洞，继续递增到 3
    resp = await client.delete(f"/api/v1/testing-plans/{plan_a['id']}", headers=auth)
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth, json={**body, "system_name": "递增系统C"},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["ticket_id"] == "20260101-3"

    # 被释放的 20260101-1 仍可手动指定给新工单
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={**body, "system_name": "递增系统D", "ticket_id_manual": "20260101-1"},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["ticket_id"] == "20260101-1"


async def test_ticket_id_manual_occupancy_and_ghost(client: AsyncClient, auth: dict):
    """工单编号占用口径统一：自动分配跳过 manual 占用；manual 记录底层 seq 不产生幽灵占用。"""
    body = {"department": "占用部门", "receive_time": "2026-07-30"}

    # A 手动指定 20260730-2（其 ticket_seq=0，不参与自动分配）
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={**body, "system_name": "占用系统A", "ticket_id_manual": "20260730-2"},
    )
    assert resp.status_code == 200, resp.text

    # 自动工单 B：max(纯自动 seq, manual N)=2 -> 新序号 3，编号 20260730-3（跳过 A 占用的 2）
    resp = await client.post("/api/v1/testing-plans", headers=auth, json={**body, "system_name": "占用系统B"})
    assert resp.status_code == 200, resp.text
    assert resp.json()["ticket_id"] == "20260730-3"

    # 自动工单 C：max(3, 2)+1=4 -> 20260730-4（单调递增）
    resp = await client.post("/api/v1/testing-plans", headers=auth, json={**body, "system_name": "占用系统C"})
    assert resp.status_code == 200, resp.text
    assert resp.json()["ticket_id"] == "20260730-4"

    # 幽灵占用场景：D 自动得到 20260730-5 后手动改为 20260730-88（底层 seq 保留 5）
    resp = await client.post("/api/v1/testing-plans", headers=auth, json={**body, "system_name": "占用系统D"})
    plan_d = resp.json()
    assert plan_d["ticket_id"] == "20260730-5"
    resp = await client.put(
        f"/api/v1/testing-plans/{plan_d['id']}", headers=auth,
        json={**body, "system_name": "占用系统D", "ticket_id_manual": "20260730-88"},
    )
    assert resp.status_code == 200, resp.text

    # E 手动指定 20260730-5（D 显示为 20260730-88，无任何工单显示 5）→ 应放行
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={**body, "system_name": "占用系统E", "ticket_id_manual": "20260730-5"},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["ticket_id"] == "20260730-5"


async def test_create_retest_record_with_status(client: AsyncClient, auth: dict):
    """新增复测记录可同时调整漏洞状态：选结论须填复测内容，已修复/复测未通过均生效。"""
    # 准备：漏洞进入复测中
    resp = await client.post("/api/v1/vulns", headers=auth, json={"title": "复测状态联动漏洞", "level": 20})
    vul_id = resp.json()["id"]
    for status in (50, 55):
        resp = await client.post(f"/api/v1/vulns/{vul_id}/transition", headers=auth, json={"status": status})
        assert resp.status_code == 200, resp.text

    # 新增复测记录时选「已修复」但未填写内容 → 拒绝
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/retests", headers=auth,
        json={"content_html": "", "status": 60},
    )
    assert resp.status_code == 400

    # 新增复测记录时选「已修复」并填写内容 → 成功且状态流转为已修复
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/retests", headers=auth,
        json={"content_html": "<p>本次复测通过</p>", "content_json": None, "status": 60},
    )
    assert resp.status_code == 200, resp.text
    vul = (await client.get(f"/api/v1/vulns/{vul_id}", headers=auth)).json()
    assert vul["status"] == 60

    # 新漏洞复测未通过：新增复测记录时选「复测未修复」回修复中
    resp = await client.post("/api/v1/vulns", headers=auth, json={"title": "复测未通过联动漏洞", "level": 20})
    vul2 = resp.json()["id"]
    for status in (50, 55):
        resp = await client.post(f"/api/v1/vulns/{vul2}/transition", headers=auth, json={"status": status})
        assert resp.status_code == 200, resp.text
    resp = await client.post(
        f"/api/v1/vulns/{vul2}/retests", headers=auth,
        json={"content_html": "<p>复测仍存在</p>", "status": 50},
    )
    assert resp.status_code == 200, resp.text
    vul2_detail = (await client.get(f"/api/v1/vulns/{vul2}", headers=auth)).json()
    assert vul2_detail["status"] == 50
    assert vul2_detail["is_retest"] is True


async def test_report_section_contains_retest(client: AsyncClient, auth: dict):
    """含复测内容的漏洞生成报告：章节 content_html 嵌入复测详情。"""
    resp = await client.post(
        "/api/v1/vulns", headers=auth, json={"title": "报告复测漏洞", "level": 20},
    )
    vul_id = resp.json()["id"]
    resp = await client.post(
        f"/api/v1/vulns/{vul_id}/retests", headers=auth,
        json={"content_html": "<p>复测发现仍可利用</p>", "content_json": None},
    )
    assert resp.status_code == 200, resp.text

    resp = await client.post(
        "/api/v1/reports/from-vulns", headers=auth,
        json={"title": "复测内容报告", "vul_ids": [vul_id]},
    )
    assert resp.status_code == 200, resp.text
    section = resp.json()["sections"][0]
    assert "复测详情" in section["content_html"]
    assert "复测发现仍可利用" in section["content_html"]


async def test_knowledge_crud_and_from_vul(client: AsyncClient, auth: dict):
    """漏洞模板库：upsert / 按类型查询 / 存为模板 / 删除。"""
    # 新建条目（vul_type=10 SQL注入类，以 meta 字典为准）
    resp = await client.post(
        "/api/v1/knowledge", headers=auth,
        json={"vulnerability_name": "SQL注入", "vul_type": 10, "severity_level": 10,
              "description_html": "<p>标准SQL注入描述</p>",
              "harm_html": "<p>可拖库</p>", "solution_html": "<p>参数化查询</p>",
              "references": ["https://owasp.org/Top10/"]},
    )
    assert resp.status_code == 200, resp.text
    entry = resp.json()
    assert entry["vulnerability_name"] == "SQL注入"
    assert entry["vul_type"] == 10
    assert entry["severity_level"] == 10
    assert entry["references"] == ["https://owasp.org/Top10/"]

    # 同名称再提交：覆盖而非新建（POST 为整体覆盖，未传字段取默认值）
    resp = await client.post(
        "/api/v1/knowledge", headers=auth,
        json={"vulnerability_name": "SQL注入", "vul_type": 10, "severity_level": 10,
              "description_html": "<p>描述V2</p>"},
    )
    assert resp.json()["id"] == entry["id"]
    assert resp.json()["description_html"] == "<p>描述V2</p>"

    # 同类型可存多条（不同名称）
    resp = await client.post(
        "/api/v1/knowledge", headers=auth,
        json={"vulnerability_name": "SQL盲注", "vul_type": 10, "severity_level": 20},
    )
    assert resp.status_code == 200, resp.text
    resp = await client.get("/api/v1/knowledge", headers=auth)
    assert len([e for e in resp.json() if e["vul_type"] == 10]) == 2

    # 按类型查询：返回列表（前端弹窗选择），危害等级最高的在前；未知类型 404
    resp = await client.get("/api/v1/knowledge/by-type/10", headers=auth)
    assert resp.status_code == 200
    items = resp.json()
    assert isinstance(items, list)
    assert items[0]["vulnerability_name"] == "SQL注入"
    resp = await client.get("/api/v1/knowledge/by-type/9999", headers=auth)
    assert resp.status_code == 200
    assert resp.json() == []

    # 未知类型/等级/空名称/非法参考链接拒绝
    resp = await client.post("/api/v1/knowledge", headers=auth,
                             json={"vulnerability_name": "x", "vul_type": 9999})
    assert resp.status_code == 400
    resp = await client.post("/api/v1/knowledge", headers=auth,
                             json={"vulnerability_name": "x", "vul_type": 10, "severity_level": 99})
    assert resp.status_code == 400
    resp = await client.post("/api/v1/knowledge", headers=auth,
                             json={"vulnerability_name": "   ", "vul_type": 10})
    assert resp.status_code == 422
    resp = await client.post("/api/v1/knowledge", headers=auth,
                             json={"vulnerability_name": "x", "vul_type": 10,
                                   "references": ["javascript:alert(1)"]})
    assert resp.status_code == 422

    # PUT 按 ID 编辑：改名、改等级；改为已存在名称被拒
    resp = await client.put(
        f"/api/v1/knowledge/{entry['id']}", headers=auth,
        json={"vulnerability_name": "SQL注入（联合查询）", "vul_type": 10, "severity_level": 20},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["vulnerability_name"] == "SQL注入（联合查询）"
    assert resp.json()["severity_level"] == 20
    resp = await client.put(
        f"/api/v1/knowledge/{entry['id']}", headers=auth,
        json={"vulnerability_name": "SQL盲注", "vul_type": 10},
    )
    assert resp.status_code == 400

    # 从已有漏洞存为模板：按标题作为名称 upsert，携带等级与描述/修复建议
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={"title": "知识库模板源漏洞", "level": 20, "vul_type": 10,
              "description_html": "<p>高质量描述</p>", "solution_html": "<p>高质量修复建议</p>"},
    )
    vul_id = resp.json()["id"]
    resp = await client.post(f"/api/v1/knowledge/from-vul/{vul_id}", headers=auth)
    assert resp.status_code == 200, resp.text
    assert resp.json()["vulnerability_name"] == "知识库模板源漏洞"
    assert resp.json()["severity_level"] == 20
    assert resp.json()["description_html"] == "<p>高质量描述</p>"
    assert resp.json()["solution_html"] == "<p>高质量修复建议</p>"

    # 未登录拒绝
    resp = await client.get("/api/v1/knowledge")
    assert resp.status_code == 401

    # 删除
    resp = await client.delete(f"/api/v1/knowledge/{entry['id']}", headers=auth)
    assert resp.status_code == 200
    resp = await client.get("/api/v1/knowledge", headers=auth)
    assert entry["id"] not in [e["id"] for e in resp.json()]


async def test_knowledge_batch_import_and_delete(client: AsyncClient, auth: dict):
    """漏洞模板库：批量导入（按名称 upsert）与批量删除。"""
    items = [
        {"vulnerability_name": "批量-SSRF", "vul_type": 75, "severity_level": 20,
         "description_html": "<p>SSRF描述</p>", "references": ["https://portswigger.net/web-security/ssrf"]},
        {"vulnerability_name": "批量-垂直越权", "vul_type": 40, "severity_level": 10},
        {"vulnerability_name": "批量-弱口令", "vul_type": 65, "severity_level": 20},
    ]
    resp = await client.post("/api/v1/knowledge/batch-import", headers=auth, json={"items": items})
    assert resp.status_code == 200, resp.text
    assert resp.json() == {"created": 3, "updated": 0, "total": 3}

    # 再次导入：同名覆盖，新名新增
    items[0]["description_html"] = "<p>SSRF描述V2</p>"
    items.append({"vulnerability_name": "批量-未授权访问", "vul_type": 40, "severity_level": 20})
    resp = await client.post("/api/v1/knowledge/batch-import", headers=auth, json={"items": items})
    assert resp.json() == {"created": 1, "updated": 3, "total": 4}
    resp = await client.get("/api/v1/knowledge", headers=auth)
    batch_rows = [e for e in resp.json() if e["vulnerability_name"].startswith("批量-")]
    assert len(batch_rows) == 4
    assert next(e for e in batch_rows if e["vulnerability_name"] == "批量-SSRF")["description_html"] == "<p>SSRF描述V2</p>"

    # 批内重名 / 字典码非法：整批拒绝
    resp = await client.post("/api/v1/knowledge/batch-import", headers=auth, json={"items": [
        {"vulnerability_name": "重名", "vul_type": 10}, {"vulnerability_name": "重名", "vul_type": 15},
    ]})
    assert resp.status_code == 400
    resp = await client.post("/api/v1/knowledge/batch-import", headers=auth, json={"items": [
        {"vulnerability_name": "非法类型", "vul_type": 9999},
    ]})
    assert resp.status_code == 400
    resp = await client.post("/api/v1/knowledge/batch-import", headers=auth, json={"items": []})
    assert resp.status_code == 422

    # 批量删除（含不存在的 ID，忽略）
    ids = [e["id"] for e in batch_rows]
    resp = await client.post("/api/v1/knowledge/batch-delete", headers=auth,
                             json={"ids": ids + [999999]})
    assert resp.status_code == 200, resp.text
    assert resp.json()["deleted"] == 4
    resp = await client.get("/api/v1/knowledge", headers=auth)
    assert not [e for e in resp.json() if e["vulnerability_name"].startswith("批量-")]


async def test_word_import_knowledge_backfill(client: AsyncClient, auth: dict):
    """Word 导入确认入库：描述/修复建议为空时自动套用知识库模板。"""
    # 准备 vul_type=15 的模板
    resp = await client.post(
        "/api/v1/knowledge", headers=auth,
        json={"vulnerability_name": "存储型XSS", "vul_type": 15, "severity_level": 20,
              "description_html": "<p>回填标准描述</p>",
              "harm_html": "<p>回填危害说明</p>", "solution_html": "<p>回填修复建议</p>"},
    )
    assert resp.status_code == 200, resp.text

    # 上传官方模板样例并等待解析
    resp = await client.get("/api/v1/imports/template", headers=auth)
    resp = await client.post(
        "/api/v1/imports", headers=auth,
        files={"file": ("知识库回填样例.docx", BytesIO(resp.content),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    batch_id = resp.json()["id"]
    detail = await _wait_batch(client, auth, batch_id)
    rec = detail["records"][0]

    # 清空描述/修复建议并改为模板类型，确认入库触发回填
    resp = await client.put(
        f"/api/v1/imports/records/{rec['id']}", headers=auth,
        json={"title": "知识库回填漏洞", "vul_type": 15,
              "description_html": "", "solution_html": ""},
    )
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        f"/api/v1/imports/{batch_id}/confirm", headers=auth,
        json={"record_ids": [rec["id"]]},
    )
    assert resp.status_code == 200, resp.text

    resp = await client.get("/api/v1/vulns", headers=auth, params={"search": "知识库回填漏洞"})
    vul = resp.json()["items"][0]
    assert "回填标准描述" in vul["description_html"]
    assert "危害说明" in vul["description_html"]
    assert "回填危害说明" in vul["description_html"]
    assert vul["solution_html"] == "<p>回填修复建议</p>"


async def test_vuln_list_sorting(client: AsyncClient, auth: dict):
    """漏洞列表按 level 升序排序结果非降序；sort 非法字段回退默认排序不报错。"""
    # 造三条不同等级的漏洞
    for lv in (40, 10, 30):
        resp = await client.post(
            "/api/v1/vulns", headers=auth,
            json={"title": f"排序用例漏洞-{lv}", "level": lv, "vul_type": 30},
        )
        assert resp.status_code == 200, resp.text

    resp = await client.get(
        "/api/v1/vulns", headers=auth, params={"sort": "level", "order": "asc", "size": 100},
    )
    assert resp.status_code == 200, resp.text
    levels = [v["level"] for v in resp.json()["items"]]
    assert levels == sorted(levels), levels

    # 降序
    resp = await client.get(
        "/api/v1/vulns", headers=auth, params={"sort": "level", "order": "desc", "size": 100},
    )
    levels = [v["level"] for v in resp.json()["items"]]
    assert levels == sorted(levels, reverse=True), levels

    # 非法排序字段回退默认排序（不报错）
    resp = await client.get(
        "/api/v1/vulns", headers=auth, params={"sort": "drop table", "order": "asc"},
    )
    assert resp.status_code == 200, resp.text


async def test_knowledge_default_sorting(client: AsyncClient, auth: dict):
    """知识库列表默认按 severity_level 升序为主、vul_type 升序为次。"""
    entries = [
        {"vulnerability_name": "排序模板A", "vul_type": 40, "severity_level": 30},
        {"vulnerability_name": "排序模板B", "vul_type": 20, "severity_level": 30},
        {"vulnerability_name": "排序模板C", "vul_type": 60, "severity_level": 10},
    ]
    for e in entries:
        resp = await client.post("/api/v1/knowledge", headers=auth, json=e)
        assert resp.status_code == 200, resp.text

    resp = await client.get("/api/v1/knowledge", headers=auth)
    assert resp.status_code == 200, resp.text
    rows = resp.json()
    # 主关键字 severity_level 升序，同级时 vul_type 升序
    keys = [(r["severity_level"], r["vul_type"]) for r in rows]
    assert keys == sorted(keys), keys


async def test_plan_complete_no_vuln_flow(client: AsyncClient, auth: dict):
    """无漏洞闭环：无漏洞完结 → 测试通过 + 无漏洞报告 → 重复确认拒绝 → 补录漏洞自动重开。"""
    # meta 字典包含新增的「测试通过」状态
    meta = (await client.get("/api/v1/meta", headers=auth)).json()
    assert meta["testing_plan_status"]["70"] == "测试通过"

    # 初测中且无漏洞的计划
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "无漏洞闭环系统", "test_type": "渗透测试",
              "receive_time": "2026-08-01", "status": 20},
    )
    assert resp.status_code == 200, resp.text
    plan_id = resp.json()["id"]

    # 确认无漏洞完结：状态流转、初测完成打点、结论记录、无漏洞报告生成
    resp = await client.post(
        f"/api/v1/testing-plans/{plan_id}/complete-no-vuln", headers=auth,
        json={"conclusion": "覆盖 OWASP Top 10 主要攻击面，未发现安全漏洞",
              "generate_report": True, "title": ""},
    )
    assert resp.status_code == 200, resp.text
    plan = resp.json()
    assert plan["status"] == 70
    assert plan["no_vul_conclusion"] == "覆盖 OWASP Top 10 主要攻击面，未发现安全漏洞"
    assert plan["first_test_done_time"]
    assert len(plan["reports"]) == 1
    report_brief = plan["reports"][0]
    assert "无漏洞闭环系统" in report_brief["title"]
    assert "渗透测试报告（无漏洞）" in report_brief["title"]
    # 无漏洞报告按初测报告口径计入计划实际人天（接收日期至确认当天）
    assert plan["actual_mandays"] >= 1

    # 报告单章节「测试结论」，无漏洞关联，结论含安全测试通过文案与补充说明
    resp = await client.get(f"/api/v1/reports/{report_brief['id']}", headers=auth)
    assert resp.status_code == 200, resp.text
    detail = resp.json()
    assert detail["testing_plan_id"] == plan_id
    assert len(detail["sections"]) == 1
    section = detail["sections"][0]
    assert section["title"] == "测试结论"
    assert section["vul_id"] is None
    assert "未发现" in section["content_html"]
    assert "安全测试通过" in section["content_html"]
    assert "OWASP" in section["content_html"]

    # 重复确认被拒绝
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/complete-no-vuln", headers=auth, json={})
    assert resp.status_code == 400

    # 补录漏洞后计划自动重开为「初测中」（关联需先认领）
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        "/api/v1/vulns", headers=auth, json={"title": "无漏洞闭环补录漏洞", "level": 30},
    )
    assert resp.status_code == 200, resp.text
    vul_id = resp.json()["id"]
    resp = await client.post(
        f"/api/v1/testing-plans/{plan_id}/attach-vulns", headers=auth, json={"vul_ids": [vul_id]},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["status"] == 20
    # 重开后不可再确认无漏洞（存在关联漏洞），无漏洞测试结论保留以便追溯
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/complete-no-vuln", headers=auth, json={})
    assert resp.status_code == 400
    plan = (await client.get(f"/api/v1/testing-plans/{plan_id}", headers=auth)).json()
    assert plan["no_vul_conclusion"] == "覆盖 OWASP Top 10 主要攻击面，未发现安全漏洞"


async def test_plan_complete_no_vuln_requires_no_vulns(client: AsyncClient, auth: dict):
    """存在关联漏洞的计划不能确认无漏洞；直接创建漏洞到已通过计划同样触发重开。"""
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "有漏洞系统", "receive_time": "2026-08-01", "status": 20},
    )
    assert resp.status_code == 200, resp.text
    plan_id = resp.json()["id"]
    # 录入/关联漏洞需先认领计划
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/claim", headers=auth)
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={"title": "有漏洞系统-漏洞1", "level": 40, "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    vid = resp.json()["id"]

    # 有关联漏洞时拒绝无漏洞完结
    resp = await client.post(f"/api/v1/testing-plans/{plan_id}/complete-no-vuln", headers=auth, json={})
    assert resp.status_code == 400

    # 移除漏洞后不生成报告仅流转状态与记录结论
    resp = await client.delete(f"/api/v1/vulns/{vid}", headers=auth)
    assert resp.status_code == 200
    resp = await client.post(
        f"/api/v1/testing-plans/{plan_id}/complete-no-vuln", headers=auth,
        json={"conclusion": "仅完结不生成报告", "generate_report": False},
    )
    assert resp.status_code == 200, resp.text
    plan = resp.json()
    assert plan["status"] == 70
    assert plan["reports"] == []

    # 直接向已通过计划创建漏洞（不经 attach-vulns）同样自动重开
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={"title": "有漏洞系统-漏洞2", "level": 40, "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    plan = (await client.get(f"/api/v1/testing-plans/{plan_id}", headers=auth)).json()
    assert plan["status"] == 20


# ---------- refresh token 空闲 24h 滑动过期 ----------
async def test_refresh_rotation_and_expiry(client: AsyncClient):
    """refresh 轮换下发新令牌对（滑动重置）；过期 refresh 被拒绝。"""
    resp = await client.post(
        "/api/v1/auth/login", data={"username": "admin", "password": "admin123"}
    )
    refresh = resp.json()["refresh_token"]

    # 轮换：新 access 可用，且再次下发新 refresh（计时重置的实现）
    resp = await client.post("/api/v1/auth/refresh", json={"refresh_token": refresh})
    assert resp.status_code == 200, resp.text
    pair = resp.json()
    me = await client.get("/api/v1/auth/me", headers={"Authorization": f"Bearer {pair['access_token']}"})
    assert me.status_code == 200
    resp = await client.post("/api/v1/auth/refresh", json={"refresh_token": pair["refresh_token"]})
    assert resp.status_code == 200

    # 空闲超 24 小时的 refresh（构造已过期令牌）被拒绝
    from datetime import timedelta

    from app.core import security

    expired = security._create_token("1", "refresh", timedelta(hours=-1), 0)
    resp = await client.post("/api/v1/auth/refresh", json={"refresh_token": expired})
    assert resp.status_code == 401


# ---------- F7 审计 ----------
async def test_audit_login_and_operation(client: AsyncClient, auth: dict):
    """登录成败写审计；敏感操作（建漏洞）写审计；查询端点按类目过滤。"""
    # 失败登录 → login_failure
    await client.post("/api/v1/auth/login", data={"username": "admin", "password": "nope"})

    resp = await client.get("/api/v1/audit/logs", headers=auth, params={"category": "login", "size": 50})
    assert resp.status_code == 200, resp.text
    logs = resp.json()
    assert logs["total"] >= 1
    actions = {i["action"] for i in logs["items"]}
    assert "login_success" in actions and "login_failure" in actions
    sample = logs["items"][0]
    assert sample["username"] and sample["ip"] != "" and sample["create_time"]

    # 操作日志：建漏洞 → vuln_create
    await client.post("/api/v1/vulns", headers=auth, json={"title": "审计测试漏洞", "level": 30})
    resp = await client.get("/api/v1/audit/logs", headers=auth, params={"category": "operation", "size": 50})
    actions = {i["action"] for i in resp.json()["items"]}
    assert "vuln_create" in actions

    # 筛选：按动作精确定位
    resp = await client.get("/api/v1/audit/logs", headers=auth, params={"action": "login_failure"})
    assert all(i["action"] == "login_failure" for i in resp.json()["items"])


async def test_meta_audit_and_notify_dicts(client: AsyncClient, auth: dict):
    resp = await client.get("/api/v1/meta", headers=auth)
    meta = resp.json()
    assert meta["audit_actions"]["login_success"] == "登录成功"
    assert set(meta["notify_channel_types"]) == {"wecom", "dingtalk", "email"}
    assert "retest_completed" in meta["notify_events"]


# ---------- F6 PAT 与开放 API ----------
async def test_pat_lifecycle_and_open_api(client: AsyncClient, auth: dict):
    # 非法档位被拒
    resp = await client.post("/api/v1/pats", headers=auth, json={"name": "非法档位", "expire_days": 15})
    assert resp.status_code == 422

    # 创建：明文仅此一次返回
    resp = await client.post(
        "/api/v1/pats", headers=auth, json={"name": "看板令牌", "expire_days": 30}
    )
    assert resp.status_code == 200, resp.text
    pat = resp.json()
    assert pat["token"].startswith("tlp_") and pat["prefix"]
    pat_id, plaintext = pat["id"], pat["token"]

    # 列表不含明文
    resp = await client.get("/api/v1/pats", headers=auth)
    items = resp.json()["items"]
    assert all("token" not in i for i in items)
    assert any(i["id"] == pat_id and i["name"] == "看板令牌" for i in items)

    pat_headers = {"Authorization": f"Bearer {plaintext}"}
    # 开放 API：PAT 可查询漏洞与统计
    resp = await client.get("/api/v1/open/vulns", headers=pat_headers, params={"size": 5})
    assert resp.status_code == 200, resp.text
    assert "items" in resp.json() and "total" in resp.json()
    resp = await client.get("/api/v1/open/stats", headers=pat_headers)
    assert resp.status_code == 200, resp.text
    assert "total_vulns" in resp.json()

    # 认证边界：JWT 访问开放 API 被拒；PAT 访问站内端点被拒
    resp = await client.get("/api/v1/open/vulns", headers=auth)
    assert resp.status_code == 401
    resp = await client.get("/api/v1/vulns", headers=pat_headers)
    assert resp.status_code == 401

    # 吊销后 PAT 失效
    resp = await client.delete(f"/api/v1/pats/{pat_id}", headers=auth)
    assert resp.status_code == 200
    resp = await client.get("/api/v1/open/vulns", headers=pat_headers)
    assert resp.status_code == 401


async def test_pat_expired_rejected(client: AsyncClient, auth: dict):
    """过期 PAT 被拒绝（直接落库一条已过期令牌）。"""
    from datetime import timedelta

    from app.core.timeutil import now
    from app.db import async_session_maker
    from app.models import PersonalAccessToken

    from app.api.v1.pats import generate_pat

    token, token_hash, prefix = generate_pat()
    async with async_session_maker() as session:
        session.add(PersonalAccessToken(
            user_id=1, name="过期令牌", token_hash=token_hash, prefix=prefix,
            expires_at=now() - timedelta(minutes=1),
        ))
        await session.commit()

    resp = await client.get("/api/v1/open/vulns", headers={"Authorization": f"Bearer {token}"})
    assert resp.status_code == 401
    assert "过期" in resp.json()["detail"]


async def test_pat_rate_limit(client: AsyncClient, auth: dict, monkeypatch):
    from app.core.config import settings

    monkeypatch.setattr(settings, "PAT_RATE_LIMIT", 3)
    resp = await client.post("/api/v1/pats", headers=auth, json={"name": "限流令牌", "expire_days": 7})
    plaintext = resp.json()["token"]
    headers = {"Authorization": f"Bearer {plaintext}"}
    codes = [(await client.get("/api/v1/open/stats", headers=headers)).status_code for _ in range(4)]
    assert codes[:3] == [200, 200, 200]
    assert codes[3] == 429


async def test_open_api_plan_read(client: AsyncClient, auth: dict):
    """开放 API 工单查询：PAT 可读渗透测试工单与漏扫基线工单，JWT 与站内端点边界不变。"""
    resp = await client.post(
        "/api/v1/pats", headers=auth, json={"name": "工单只读令牌", "expire_days": 30},
    )
    pat = {"Authorization": f"Bearer {resp.json()['token']}"}

    # 列表与详情（先在站内造一条，保证有数据）
    created = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "开放API只读系统", "test_type": "渗透测试", "status": 10},
    )
    assert created.status_code == 200, created.text
    plan_id = created.json()["id"]

    resp = await client.get("/api/v1/open/testing-plans", headers=pat, params={"size": 5})
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert "total" in body and "items" in body

    resp = await client.get(f"/api/v1/open/testing-plans/{plan_id}", headers=pat)
    assert resp.status_code == 200, resp.text
    for key in ("id", "ticket_id", "testers", "vuls", "reports", "retest_rounds"):
        assert key in resp.json()

    # 筛选参数生效：按部门精确过滤
    resp = await client.get(
        "/api/v1/open/testing-plans", headers=pat, params={"search": "开放API只读系统"},
    )
    assert [p["id"] for p in resp.json()["items"]] == [plan_id]

    # 漏扫基线工单列表可读
    assert (await client.get("/api/v1/open/nonpen-plans", headers=pat)).status_code == 200

    # 认证边界：JWT 访问开放工单接口被拒；PAT 访问站内工单端点被拒；不存在 404
    assert (await client.get("/api/v1/open/testing-plans", headers=auth)).status_code == 401
    assert (await client.get("/api/v1/testing-plans", headers=pat)).status_code == 401
    assert (await client.get("/api/v1/open/testing-plans/999999", headers=pat)).status_code == 404
    assert (await client.get("/api/v1/open/nonpen-plans/999999", headers=pat)).status_code == 404

    # 分页参数越界 422
    assert (
        await client.get("/api/v1/open/testing-plans", headers=pat, params={"size": 500})
    ).status_code == 422


async def test_open_api_plan_write(client: AsyncClient, auth: dict):
    """开放 API 工单写入：创建 / 更新渗透测试工单与漏扫基线工单，写操作受 special:manage 约束。"""
    resp = await client.post(
        "/api/v1/pats", headers=auth, json={"name": "工单写入令牌", "expire_days": 30},
    )
    pat = {"Authorization": f"Bearer {resp.json()['token']}"}

    # 创建渗透测试工单（工单ID按需求接收日期自动生成）
    resp = await client.post(
        "/api/v1/open/testing-plans", headers=pat,
        json={
            "system_name": "开放API写入系统", "test_type": "渗透测试",
            "department": "研发一部", "receive_time": "2026-09-03", "status": 10,
        },
    )
    assert resp.status_code == 200, resp.text
    plan = resp.json()
    plan_id = plan["id"]
    assert plan["ticket_id"].startswith("20260903-")
    assert plan["department"] == "研发一部"

    # 全量更新：改部门 + 状态流转 10 → 20（未测试 → 初测中）
    update_body = {
        "system_name": "开放API写入系统", "test_type": "渗透测试", "department": "研发二部",
        "receive_time": "2026-09-03", "status": 20,
    }
    resp = await client.put(f"/api/v1/open/testing-plans/{plan_id}", headers=pat, json=update_body)
    assert resp.status_code == 200, resp.text
    assert resp.json()["department"] == "研发二部"
    assert resp.json()["status"] == 20

    # 非法状态流转 400（初测中 20 不能直接到复测完成 60）
    resp = await client.put(
        f"/api/v1/open/testing-plans/{plan_id}", headers=pat, json={**update_body, "status": 60},
    )
    assert resp.status_code == 400

    # 必填校验 422（system_name 为空）
    resp = await client.post(
        "/api/v1/open/testing-plans", headers=pat, json={"system_name": ""},
    )
    assert resp.status_code == 422

    # 更新不存在的工单 404
    resp = await client.put(
        "/api/v1/open/testing-plans/999999", headers=pat, json=update_body,
    )
    assert resp.status_code == 404

    # 漏扫基线工单：创建 + 更新（测试项勾选生效）
    resp = await client.post(
        "/api/v1/open/nonpen-plans", headers=pat,
        json={
            "system_name": "开放API漏扫系统", "department": "研发一部",
            "receive_time": "2026-09-03", "test_items": ["baseline", "host"],
        },
    )
    assert resp.status_code == 200, resp.text
    nonpen = resp.json()
    nonpen_id = nonpen["id"]
    assert nonpen["ticket_id"].startswith("20260903-")
    assert nonpen["items"]["baseline"]["status"] == "not_started"
    assert nonpen["items"]["web"]["status"] == "ignored"

    resp = await client.put(
        f"/api/v1/open/nonpen-plans/{nonpen_id}", headers=pat,
        json={
            "system_name": "开放API漏扫系统", "department": "研发二部",
            "receive_time": "2026-09-03", "test_items": ["web"],
        },
    )
    assert resp.status_code == 200, resp.text
    updated = resp.json()
    assert updated["department"] == "研发二部"
    # 仅合并勾选变化：web 由 ignored 变 not_started，baseline 取消勾选保留 ignored
    assert updated["items"]["web"]["status"] == "not_started"
    assert updated["items"]["baseline"]["status"] == "ignored"

    # 工单ID必须存在来源（无接收日期且未手动指定）→ 422
    resp = await client.post(
        "/api/v1/open/nonpen-plans", headers=pat, json={"system_name": "缺少工单ID来源"},
    )
    assert resp.status_code == 422

    # 无 special:manage 权限的账号，其令牌写操作 403（读仍放行）
    resp = await client.post(
        "/api/v1/roles", headers=auth,
        json={"name": "无专项写入权限", "permissions": ["vuln:submit"], "remark": ""},
    )
    role_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/users", headers=auth,
        json={"username": "open_api_writer", "password": "Writer@123", "realname": "开放API写入",
              "email": "", "phone": "", "is_active": True, "role_id": role_id},
    )
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        "/api/v1/auth/login", data={"username": "open_api_writer", "password": "Writer@123"},
    )
    weak_auth = {"Authorization": f"Bearer {resp.json()['access_token']}"}
    resp = await client.post(
        "/api/v1/pats", headers=weak_auth, json={"name": "无权限令牌", "expire_days": 7},
    )
    weak_pat = {"Authorization": f"Bearer {resp.json()['token']}"}

    assert (await client.get("/api/v1/open/testing-plans", headers=weak_pat)).status_code == 200
    resp = await client.post(
        "/api/v1/open/testing-plans", headers=weak_pat, json={"system_name": "无权限创建"},
    )
    assert resp.status_code == 403
    assert "special:manage" in resp.json()["detail"]
    resp = await client.put(
        f"/api/v1/open/testing-plans/{plan_id}", headers=weak_pat, json=update_body,
    )
    assert resp.status_code == 403


async def test_batch_import_confirm(client: AsyncClient, auth: dict):
    """批量确认入库：多批次统一关联工单 → 逐批确认并返回 report_ids；
    已确认批次重复关联跳过；单批失败隔离不影响其余批次；权限与参数校验。"""
    # 创建统一工单
    resp = await client.post(
        "/api/v1/testing-plans", headers=auth,
        json={"system_name": "批量确认系统", "test_type": "渗透测试"},
    )
    assert resp.status_code == 200, resp.text
    plan_id = resp.json()["id"]

    # 上传两份报告格式批次
    b1 = await _upload_report_batch(client, auth, "批量确认系统A", "20260801批量确认A渗透测试报告.docx")
    b2 = await _upload_report_batch(client, auth, "批量确认系统B", "20260801批量确认B渗透测试报告.docx")

    # 正常批量确认：统一关联工单，两个批次都成功，报告自动生成并挂到该工单
    resp = await client.post(
        "/api/v1/imports/batch-confirm", headers=auth,
        json={"batch_ids": [b1, b2], "testing_plan_id": plan_id},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["confirmed"] == 2
    assert body["skipped"] == 0
    assert body["failed"] == 0
    assert len(body["report_ids"]) == 2
    assert all(d["status"] == "confirmed" for d in body["details"])
    for rid in body["report_ids"]:
        report = (await client.get(f"/api/v1/reports/{rid}", headers=auth)).json()
        assert report["testing_plan_id"] == plan_id

    # 重复关联：已确认批次再批量确认（含 batch_ids 内重复）→ 全部 skipped
    resp = await client.post(
        "/api/v1/imports/batch-confirm", headers=auth,
        json={"batch_ids": [b1, b1, b2, b1]},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["confirmed"] == 0
    assert body["skipped"] == 2
    assert body["failed"] == 0
    assert all(d["status"] == "skipped" for d in body["details"])

    # 部分失败隔离：一个有效批次 + 一个不存在批次 → confirmed=1, failed=1，有效批次照常入库
    b3 = await _upload_report_batch(client, auth, "批量确认系统C", "20260801批量确认C渗透测试报告.docx")
    resp = await client.post(
        "/api/v1/imports/batch-confirm", headers=auth,
        json={"batch_ids": [b3, 999999]},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["confirmed"] == 1
    assert body["failed"] == 1
    assert len(body["report_ids"]) == 1
    failed_detail = next(d for d in body["details"] if d["status"] == "failed")
    assert failed_detail["batch_id"] == 999999
    assert failed_detail["detail"] == "导入批次不存在"

    # 空 batch_ids 拒绝
    resp = await client.post(
        "/api/v1/imports/batch-confirm", headers=auth, json={"batch_ids": []},
    )
    assert resp.status_code == 400

    # 非法工单拒绝
    resp = await client.post(
        "/api/v1/imports/batch-confirm", headers=auth,
        json={"batch_ids": [b1], "testing_plan_id": 999999},
    )
    assert resp.status_code == 400

    # 无 import:manage 权限用户 403
    resp = await client.post(
        "/api/v1/roles", headers=auth,
        json={"name": "无导入权限", "permissions": ["vuln:submit"], "remark": ""},
    )
    assert resp.status_code == 200, resp.text
    role_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/users", headers=auth,
        json={"username": "no_import_perm", "password": "Tester@123", "realname": "无导入",
              "email": "", "phone": "", "is_active": True, "role_id": role_id},
    )
    assert resp.status_code == 200, resp.text
    resp = await client.post(
        "/api/v1/auth/login", data={"username": "no_import_perm", "password": "Tester@123"},
    )
    auth2 = {"Authorization": f"Bearer {resp.json()['access_token']}"}
    resp = await client.post(
        "/api/v1/imports/batch-confirm", headers=auth2,
        json={"batch_ids": [b1]},
    )
    assert resp.status_code == 403


async def test_batch_confirm_chrono_order(client: AsyncClient, auth: dict):
    """批量确认时序：同一工单同一漏洞跨多份复测报告去重合并时，
    必须按报告日期从旧到新处理，保证最新报告的「已修复」最终生效而非被旧报告覆盖。"""
    system_name = "批量时序系统"
    target_url = "http://10.9.9.9/chrono"
    docx_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    async def upload(filename: str, fixed: bool) -> int:
        doc = _build_report_docx(
            system_name, target_url, "10.9.9.9",
            sections=[("XSS时序漏洞", "中危", "XSS跨站", fixed)],
        )
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        resp = await client.post(
            "/api/v1/imports", headers=auth,
            files={"file": (filename, buf, docx_mime)},
        )
        assert resp.status_code == 200, resp.text
        bid = resp.json()["id"]
        detail = await _wait_batch(client, auth, bid)
        assert detail["batch"]["status"] == "parsed", detail
        assert detail["batch"]["doc_kind"] == "report", detail
        return bid

    # 旧复测报告 XSS 未修复，新复测报告 XSS 已修复
    older = await upload("20260723批量时序系统渗透测试复测报告.docx", False)
    newer = await upload("20260731批量时序系统渗透测试复测报告.docx", True)

    # 以「新→旧」倒序批量确认（复现列表 newest-first 选中提交的顺序）
    resp = await client.post(
        "/api/v1/imports/batch-confirm", headers=auth,
        json={"batch_ids": [newer, older]},
    )
    assert resp.status_code == 200, resp.text
    assert resp.json()["confirmed"] == 2

    # 同一工单同一标题应去重合并为一条漏洞，最终状态为「已修复」(60)
    vul = await _find_vuln(client, auth, "XSS时序漏洞")
    assert vul["status"] == 60, f"期望已修复(60)，实际 {vul['status']}"


# ---------- F3 通知渠道 ----------
async def test_notify_channel_crud_and_validation(client: AsyncClient, auth: dict):
    # 类型非法 / webhook 地址缺失被拒
    resp = await client.post(
        "/api/v1/notify-channels", headers=auth,
        json={"name": "坏类型", "type": "sms", "config": {}, "events": ["vuln_created"]},
    )
    assert resp.status_code == 422
    resp = await client.post(
        "/api/v1/notify-channels", headers=auth,
        json={"name": "缺地址", "type": "wecom", "config": {}, "events": ["vuln_created"]},
    )
    assert resp.status_code == 422
    # 邮箱渠道缺收件人被拒
    resp = await client.post(
        "/api/v1/notify-channels", headers=auth,
        json={"name": "缺收件人", "type": "email", "config": {}, "events": ["vuln_created"]},
    )
    assert resp.status_code == 422

    # 正常创建 / 编辑 / 测试发送 / 删除
    resp = await client.post(
        "/api/v1/notify-channels", headers=auth,
        json={
            "name": "安全群机器人", "type": "wecom",
            "config": {"url": "https://qyapi.example.com/hook"},
            "events": ["vuln_created", "retest_completed"], "is_active": True,
        },
    )
    assert resp.status_code == 200, resp.text
    channel = resp.json()
    assert channel["events"] == ["vuln_created", "retest_completed"]

    resp = await client.put(
        f"/api/v1/notify-channels/{channel['id']}", headers=auth,
        json={
            "name": "安全群机器人", "type": "wecom",
            "config": {"url": "https://qyapi.example.com/hook2"},
            "events": ["vuln_transition"], "is_active": False,
        },
    )
    assert resp.status_code == 200
    assert resp.json()["is_active"] is False

    resp = await client.post(f"/api/v1/notify-channels/{channel['id']}/test", headers=auth)
    assert resp.status_code == 200
    resp = await client.delete(f"/api/v1/notify-channels/{channel['id']}", headers=auth)
    assert resp.status_code == 200


async def test_notify_emit_on_vuln_created(client: AsyncClient, auth: dict, monkeypatch):
    """漏洞创建事件触发渠道分发（monkeypatch dispatch 捕获，不出站）。"""
    calls: list[tuple] = []

    async def fake_dispatch(app, func_name, *args):
        calls.append((func_name, args))

    import app.services.notify_service as notify_service

    monkeypatch.setattr(notify_service, "dispatch", fake_dispatch)

    resp = await client.post(
        "/api/v1/notify-channels", headers=auth,
        json={
            "name": "邮件渠道", "type": "email",
            "config": {"recipients": ["sec@example.com"]},
            "events": ["vuln_created"], "is_active": True,
        },
    )
    assert resp.status_code == 200

    resp = await client.post("/api/v1/vulns", headers=auth, json={"title": "通知触发漏洞", "level": 30})
    assert resp.status_code == 200

    notify_calls = [c for c in calls if c[0] == "send_notify_task"]
    assert notify_calls, "漏洞创建应触发通知分发"
    func_name, args = notify_calls[0]
    assert args[0] == "email" and args[1]["recipients"] == ["sec@example.com"]
    assert "[Talos] 新漏洞创建" in args[2]


# ---------- F4 CVSS ----------
async def test_vuln_cvss_fields(client: AsyncClient, auth: dict):
    vector = "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:H/I:H/A:H"
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={"title": "CVSS评分漏洞", "level": 10, "score": 9.8, "cvss_vector": vector},
    )
    assert resp.status_code == 200, resp.text
    vul = resp.json()
    assert vul["score"] == 9.8
    assert vul["cvss_vector"] == vector

    # 编辑改分
    resp = await client.put(
        f"/api/v1/vulns/{vul['id']}", headers=auth,
        json={"title": "CVSS评分漏洞", "level": 10, "score": 5.3,
              "cvss_vector": "CVSS:3.1/AV:N/AC:L/PR:N/UI:N/S:U/C:L/I:N/A:N"},
    )
    assert resp.status_code == 200
    assert resp.json()["score"] == 5.3


async def test_knowledge_cvss_vector(client: AsyncClient, auth: dict):
    vector = "CVSS:3.1/AV:N/AC:L/PR:L/UI:N/S:C/C:H/I:H/A:H"
    resp = await client.post(
        "/api/v1/knowledge", headers=auth,
        json={"vulnerability_name": "SSRF向量模板", "vul_type": 75, "severity_level": 20, "cvss_vector": vector},
    )
    assert resp.status_code == 200, resp.text
    entry = resp.json()
    assert entry["cvss_vector"] == vector

    # from-vul：漏洞向量随「存为模板」带入知识库
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={"title": "存模板向量漏洞", "level": 20, "cvss_vector": vector},
    )
    vul_id = resp.json()["id"]
    resp = await client.post(f"/api/v1/knowledge/from-vul/{vul_id}", headers=auth)
    assert resp.status_code == 200, resp.text
    assert resp.json()["cvss_vector"] == vector





async def test_vuln_search_by_system_name(client: AsyncClient, auth: dict):
    """关键词搜索支持系统名称：标题/URL 命中之外，还应命中关联资产（系统）的名称。"""
    resp = await client.post(
        "/api/v1/assets", headers=auth, json={"name": "检索定位专用系统XQ"},
    )
    assert resp.status_code == 200, resp.text
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns", headers=auth,
        json={
            "title": "与系统名毫无关联的标题XQ", "level": 30, "vul_type": 30,
            "affected_url": "http://xq.example.com/api",
            "description_html": "<p>xq</p>",
            "asset_ids": [asset_id],
        },
    )
    assert resp.status_code == 200, resp.text
    vul_id = resp.json()["id"]

    # 系统名称命中
    resp = await client.get(
        "/api/v1/vulns", headers=auth, params={"search": "检索定位专用系统XQ"},
    )
    assert resp.status_code == 200, resp.text
    assert vul_id in [v["id"] for v in resp.json()["items"]]

    # 标题命中依旧可用（回归确认）
    resp = await client.get(
        "/api/v1/vulns", headers=auth, params={"search": "与系统名毫无关联的标题XQ"},
    )
    assert resp.status_code == 200, resp.text
    assert vul_id in [v["id"] for v in resp.json()["items"]]


async def test_global_search(client: AsyncClient, auth: dict):
    """⌘K 全局搜索：空关键字返回空分组；按标题/名称模糊命中漏洞与资产。"""
    resp = await client.get("/api/v1/search", headers=auth, params={"q": "  "})
    assert resp.status_code == 200
    assert resp.json() == {"vulns": [], "assets": [], "plans": [], "reports": []}

    # 造数据：资产 + 挂在其上的漏洞
    resp = await client.post("/api/v1/assets", headers=auth, json={"name": "搜索目标系统Alpha"})
    assert resp.status_code == 200, resp.text
    asset_id = resp.json()["id"]
    resp = await client.post(
        "/api/v1/vulns/batch", headers=auth,
        json={
            "asset_ids": [asset_id],
            "vulns": [{"title": "搜索专用SQL注入漏洞XYZ", "level": 20, "vul_type": 10}],
        },
    )
    assert resp.status_code == 200, resp.text

    resp = await client.get("/api/v1/search", headers=auth, params={"q": "XYZ"})
    assert resp.status_code == 200
    body = resp.json()
    assert [v["title"] for v in body["vulns"]] == ["搜索专用SQL注入漏洞XYZ"]
    assert body["assets"] == []

    # 按资产名搜索命中资产分区
    resp = await client.get("/api/v1/search", headers=auth, params={"q": "Alpha"})
    assert resp.status_code == 200
    body = resp.json()
    assert any(a["name"] == "搜索目标系统Alpha" for a in body["assets"])
    assert body["vulns"] == []
