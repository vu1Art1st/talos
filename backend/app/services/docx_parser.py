"""Word 导入解析：支持两种格式。

1. 固定模板：每个漏洞为一张两列信息表格（/api/v1/imports/template 可下载示例）
   - 表格第一列为字段标签：漏洞名称 / 漏洞等级 / 漏洞类型 / 漏洞链接 / 漏洞描述 / 漏洞证明 / 修复建议
     （标签命名与报告模板「风险问题详情」章节一致；旧模板的 影响URL/复现步骤 仍兼容）
   - 第二列为内容，描述/证明/建议单元格支持多段落与截图
   - 文档中可包含任意数量的漏洞表格，其余段落内容会被忽略
2. 平台报告格式：平台导出的渗透测试（复测）报告，如「20260729综合办公系统渗透测试复测报告.docx」
   - 通过「风险问题详情」标题 + 「风险问题汇总」表识别
   - 解析封面/测试目标表得到系统名、报告日期、是否复测，用于自动生成测试计划
"""
import html as html_mod
import re
import uuid
from pathlib import Path

# XXE 评估结论：python-docx 的 oxml 解析器以 resolve_entities=False 构造，
# 不解析外部/内部实体，故 .docx 导入不存在 XXE 风险；xlsx 侧 openpyxl 同样
# 使用 resolve_entities=False，且安装 defusedxml 后自动启用其安全解析器。
from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell

from app.constants import IMPORT_LABEL_MAP, VUL_LEVEL_EXPORT, VUL_LEVEL_REVERSE, VUL_TYPE_REVERSE

_A_BLIP = qn("a:blip")
_R_EMBED = qn("r:embed")


def _norm_label(text: str) -> str:
    return re.sub(r"[\s:：]+", "", text or "")


def _save_image(part, rid: str, image_dir: Path, url_prefix: str) -> str | None:
    """把 run 中引用的图片落盘，返回可访问 URL。"""
    try:
        image_part = part.related_parts[rid]
    except KeyError:
        return None
    ext = Path(str(image_part.partname)).suffix or ".png"
    name = f"{uuid.uuid4().hex}{ext}"
    image_dir.mkdir(parents=True, exist_ok=True)
    (image_dir / name).write_bytes(image_part.blob)
    return f"{url_prefix}/{name}"


def _run_to_html(run, part, image_dir: Path, url_prefix: str) -> str:
    pieces: list[str] = []
    # 图片：w:drawing 内的 a:blip r:embed 引用
    for blip in run._element.iter(_A_BLIP):
        rid = blip.get(_R_EMBED)
        if rid:
            url = _save_image(part, rid, image_dir, url_prefix)
            if url:
                pieces.append(f'<img src="{url}">')
    text = html_mod.escape(run.text or "")
    if text:
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        pieces.append(text)
    return "".join(pieces)


def _cell_to_html(cell: _Cell, part, image_dir: Path, url_prefix: str) -> str:
    paragraphs: list[str] = []
    for p in cell.paragraphs:
        inner = "".join(_run_to_html(r, part, image_dir, url_prefix) for r in p.runs)
        if inner.strip():
            paragraphs.append(f"<p>{inner}</p>")
    return "".join(paragraphs)


def _cell_text(cell: _Cell) -> str:
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _map_level(text: str) -> int | None:
    text = _norm_label(text)
    if text in VUL_LEVEL_REVERSE:
        return VUL_LEVEL_REVERSE[text]
    for name, code in VUL_LEVEL_REVERSE.items():
        if text and (text in name or name in text):
            return code
    return None


def _map_type(text: str) -> int | None:
    text = _norm_label(text)
    if text in VUL_TYPE_REVERSE:
        return VUL_TYPE_REVERSE[text]
    for name, code in VUL_TYPE_REVERSE.items():
        if text and (text in name or name in text):
            return code
    return None


def parse_docx(file_path: str, image_dir: str, image_url_prefix: str) -> list[dict]:
    """解析文档，返回记录列表。单条记录解析失败不影响整批。

    每条记录: {title, level, vul_type, affected_url,
               description_html, reproduce_html, solution_html, errors: [str]}
    """
    doc = Document(file_path)
    part = doc.part
    img_dir = Path(image_dir)
    results: list[dict] = []

    for table in doc.tables:
        fields: dict[str, _Cell] = {}
        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = _norm_label(_cell_text(row.cells[0]))
            key = IMPORT_LABEL_MAP.get(label)
            if key and key not in fields:
                fields[key] = row.cells[1]

        if "title" not in fields:
            continue  # 非漏洞信息表格，跳过

        record: dict = {
            "title": _cell_text(fields["title"]),
            "level": 30,
            "vul_type": 75,
            "affected_url": "",
            "description_html": "",
            "reproduce_html": "",
            "solution_html": "",
            "errors": [],
        }
        if not record["title"]:
            record["errors"].append("漏洞名称为空")

        if "level" in fields:
            level = _map_level(_cell_text(fields["level"]))
            if level is None:
                record["errors"].append(f"无法识别漏洞等级「{_cell_text(fields['level'])}」，已按中危处理")
            else:
                record["level"] = level
        else:
            record["errors"].append("缺少「漏洞等级」行，已按中危处理")

        if "vul_type" in fields:
            vtype = _map_type(_cell_text(fields["vul_type"]))
            if vtype is None:
                record["errors"].append(f"无法识别漏洞类型「{_cell_text(fields['vul_type'])}」，已按其他处理")
            else:
                record["vul_type"] = vtype

        if "affected_url" in fields:
            record["affected_url"] = _cell_text(fields["affected_url"])[:512]

        for key in ("description_html", "reproduce_html", "solution_html"):
            if key in fields:
                record[key] = _cell_to_html(fields[key], part, img_dir, image_url_prefix)

        results.append(record)

    return results


# ---------- 平台报告格式（渗透测试/复测报告）解析 ----------

_SUMMARY_HEADERS = ("问题等级", "风险类型", "风险问题", "修复状态")
_LEVEL_EXPORT_REVERSE = {v: k for k, v in VUL_LEVEL_EXPORT.items()}  # 超危→10 等报告口径
# 报告文件名：日期 + 公司/系统名 + 渗透测试(复测)报告 + 可选轮次后缀 -N（同日重复发起复测自动追加 -1/-2）
# 例：20250917中移系统集成有限公司综合办公系统渗透测试复测报告-1.docx → 第二轮复测
_REPORT_NAME_RE = re.compile(r"^(\d{8})?(.*?)(?:渗透测试)?(复测)?报告(?:-(\d+))?$")
# 文件名兜底系统名剥离公司前缀：取最后一个「有限公司」之后的部分（如「中移系统集成有限公司综合办公系统」→「综合办公系统」）
_COMPANY_SUFFIX_RE = re.compile(r"(.*?有限公司)(.*)$")
_COVER_DATE_RE = re.compile(r"^(\d{4})年(\d{1,2})月(\d{1,2})日$")
_RETEST_LABEL_RE = re.compile(r"^\d*漏洞复测$")
# H3 章节内的字段标签行（归一化后） -> 记录字段
_SECTION_LABELS = {
    "测试状态": "status",
    "漏洞等级": "level",
    "漏洞链接": "affected_url",
    "漏洞描述": "description_html",
    "漏洞证明": "reproduce_html",
    "修复建议": "solution_html",
}


def _find_summary_table(doc: Document):
    """定位「风险问题汇总」表：表头含 问题等级/风险类型/风险问题/修复状态。"""
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [_norm_label(c.text) for c in table.rows[0].cells]
        if all(h in headers for h in _SUMMARY_HEADERS):
            return table
    return None


def is_report_docx(doc: Document) -> bool:
    """判定是否为平台导出的渗透测试（复测）报告格式。"""
    has_detail = any(
        p.style.name.startswith("Heading") and "风险问题详情" in p.text for p in doc.paragraphs
    )
    return has_detail and _find_summary_table(doc) is not None


def parse_report_filename(filename: str) -> dict:
    """从文件名提取报告日期/系统名/是否复测/复测轮次序号。

    文件名示例：20260729综合办公系统渗透测试报告.docx、20251011综合办公系统渗透测试复测报告-1.docx。
    复测轮次序号 retest_round_seq：初测=0；首份复测（无后缀）=1；同日重复复测带 -N 后缀=N+1（如 -1 即第二轮复测）。
    系统名剥离常见公司前缀（最后一个「有限公司」之后），保证与工单 system_name（纯系统名）匹配。
    """
    out = {"report_date": "", "system_name": "", "is_retest": False, "retest_round_seq": 0}
    m = _REPORT_NAME_RE.match(Path(filename or "").stem.strip())
    if not m:
        return out
    if m.group(1):
        d = m.group(1)
        out["report_date"] = f"{d[:4]}-{d[4:6]}-{d[6:8]}"
    system_name = m.group(2).strip()
    cm = _COMPANY_SUFFIX_RE.match(system_name)
    if cm and cm.group(2).strip():
        system_name = cm.group(2).strip()
    out["system_name"] = system_name
    out["is_retest"] = bool(m.group(3))
    out["retest_round_seq"] = (int(m.group(4)) + 1) if m.group(4) else (1 if m.group(3) else 0)
    return out


def _parse_cover(doc: Document) -> dict:
    """封面段落（首个标题之前）：公司 / 系统名 / 「渗透测试(复测)报告」 / 日期。"""
    meta = {"system_name": "", "report_date": "", "is_retest": False}
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            break
        if p.text.strip():
            lines.append(p.text.strip())
    for i, t in enumerate(lines):
        if t.endswith("报告") and ("渗透测试" in t or "复测" in t):
            meta["is_retest"] = "复测" in t
            if i >= 1:
                meta["system_name"] = lines[i - 1]
        m = _COVER_DATE_RE.match(t)
        if m:
            meta["report_date"] = f"{m.group(1)}-{int(m.group(2)):02d}-{int(m.group(3)):02d}"
    return meta


def _parse_target_table(doc: Document) -> dict:
    """「测试目标」表（首列含 业务系统名称）：系统名 / 被测URL / 被测IP / 被测测试账号。"""
    for table in doc.tables:
        labels = {}
        for row in table.rows:
            if not row.cells:
                continue
            labels.setdefault(_norm_label(row.cells[0].text), row.cells[-1].text.strip())
        if "业务系统名称" in labels:
            return {
                "system_name": labels.get("业务系统名称", ""),
                "target_url": labels.get("被测系统URL", ""),
                "target_ip": labels.get("被测系统IP", ""),
                "test_account": labels.get("被测测试账号", ""),
            }
    return {}


def _parse_schedule_table(doc: Document) -> dict:
    """「时间与人员」表（模板表5）：起始/结束时间 + 参测人员姓名。

    模板结构（含导出报告）：row0 测试工作时间段；row1 起始时间|<值>|结束时间|<值>；
    row3 表头 参测人员|所属部门|人员角色|人员分工；row4+ 各参测人员姓名列。
    解析结果用于回填报告测试周期与工单测试人员（按姓名映射系统账号）。
    """
    meta = {"test_start": "", "test_end": "", "testers": []}
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 5:
            continue
        header = [_norm_label(c.text) for c in rows[3].cells]
        if "参测人员" not in header:
            continue
        for row in rows:
            # 起始/结束时间位于同一行的不同列（如 起始时间|<值>|结束时间|<值>），按标签逐列定位取值
            cells = [_norm_label(c.text) for c in row.cells]
            for i, label in enumerate(cells):
                if label == "起始时间" and i + 1 < len(row.cells):
                    meta["test_start"] = _cell_text(row.cells[i + 1]).strip()
                elif label == "结束时间" and i + 1 < len(row.cells):
                    meta["test_end"] = _cell_text(row.cells[i + 1]).strip()
        for row in rows[4:]:
            name = _cell_text(row.cells[0]).strip()
            if name and name not in meta["testers"]:
                meta["testers"].append(name)
        break
    return meta


def _parse_summary_rows(table) -> list[dict]:
    rows: list[dict] = []
    if table is None:
        return rows
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 4 or not cells[2]:
            continue
        rows.append({
            "level_text": cells[0],
            "type_text": cells[1],
            "title": cells[2],
            "fixed": "已修复" in cells[3],
        })
    return rows


def _map_level_report(text: str) -> int | None:
    """报告口径等级：先按字典（严重/高危…），再按导出口径（超危→10）。"""
    code = _map_level(text)
    if code is not None:
        return code
    t = _norm_label(text)
    for name, c in _LEVEL_EXPORT_REVERSE.items():
        if t and (t in name or name in t):
            return c
    return None


def _fuzzy_type(text: str) -> int:
    """类型模糊匹配：精确/包含未命中时按公共字符数取最优（如 权限跨越→权限绕过）。"""
    code = _map_type(text)
    if code is not None:
        return code
    t = _norm_label(text)
    best, best_score = 75, 0
    for name, c in VUL_TYPE_REVERSE.items():
        score = len(set(t) & set(name))
        if score > best_score:
            best, best_score = c, score
    return best if best_score >= 2 else 75


def _normalize_vul_title(raw_title: str) -> tuple[str, bool]:
    """归一化报告中的漏洞标题，并判定该条记录是否「已修复」。

    - 剥掉结尾括号里的修复/整改状态后缀（含修饰词，如「（部分未修复）」「（基本已修复）」），
      使同一漏洞在初测/多轮复测报告中标题一致，从而支持跨报告去重合并。
    - fixed 判定：仅当标题明确「已修复/已整改」且不含「部分/未修复/未整改」等否定或部分修饰时才为 True。
    """
    raw_title = (raw_title or "").strip()
    title = raw_title
    # 循环剥掉结尾的修复/整改状态括号（兼容嵌套，如「（部分未修复）（未修复）」）
    while True:
        stripped = re.sub(r"[（(][^（）()]*(?:修复|整改)[^（）()]*[)）]\s*$", "", title).strip()
        if stripped == title:
            break
        title = stripped
    fixed = (
        ("已修复" in raw_title or "已整改" in raw_title)
        and "部分" not in raw_title
        and "基本" not in raw_title
        and "大部分" not in raw_title
        and "未修复" not in raw_title
        and "未整改" not in raw_title
    )
    return title or raw_title, fixed


def _match_summary(summary: list[dict], title: str) -> dict | None:
    """按标题包含关系匹配汇总表行，未命中时取公共字符最多的行。"""
    t = _norm_label(title)
    for row in summary:
        rt = _norm_label(row["title"])
        if t and rt and (t in rt or rt in t):
            return row
    best, best_score = None, 0
    for row in summary:
        score = len(set(t) & set(_norm_label(row["title"])))
        if score > best_score:
            best, best_score = row, score
    return best if best_score >= 4 else None


def parse_report_docx(file_path: str, image_dir: str, image_url_prefix: str,
                      filename: str = "") -> tuple[dict, list[dict]]:
    """解析平台报告格式文档，返回 (meta, records)。

    meta: {system_name, report_date, is_retest, target_url, target_ip}
    records 字段与 parse_docx 一致，另含 retest_html / fixed。
    """
    doc = Document(file_path)
    part = doc.part
    img_dir = Path(image_dir)

    # meta：封面优先，文件名兜底，测试目标表补齐系统名与 URL/IP，
    # 时间与人员表补齐测试周期与参测人员。
    # 封面/文件名系统名含乱码（U+FFFD）时视为无效，回退到下一数据源，保证能匹配现有工单
    meta = {"system_name": "", "report_date": "", "is_retest": False, "target_url": "", "target_ip": "",
            "test_account": "", "retest_round_seq": 0, "test_start": "", "test_end": "", "testers": []}
    from_name = parse_report_filename(filename or Path(file_path).name)
    cover = _parse_cover(doc)
    target = _parse_target_table(doc)
    schedule = _parse_schedule_table(doc)

    def _clean(name: str) -> str:
        name = (name or "").strip()
        return "" if "\ufffd" in name else name

    meta["system_name"] = (_clean(cover["system_name"])
                           or _clean(from_name["system_name"])
                           or _clean(target.get("system_name", "")))
    meta["report_date"] = cover["report_date"] or from_name["report_date"]
    meta["is_retest"] = cover["is_retest"] or from_name["is_retest"]
    meta["retest_round_seq"] = from_name.get("retest_round_seq", 0)
    meta["target_url"] = target.get("target_url", "")
    meta["target_ip"] = target.get("target_ip", "")
    meta["test_account"] = target.get("test_account", "")
    meta["test_start"] = schedule.get("test_start", "")
    meta["test_end"] = schedule.get("test_end", "")
    meta["testers"] = schedule.get("testers", [])

    # 切分「风险问题详情」下的 H3 漏洞章节
    sections: list[tuple[str, list]] = []
    in_detail = False
    current: list | None = None
    for p in doc.paragraphs:
        style = p.style.name or ""
        if style.startswith("Heading"):
            if "风险问题详情" in p.text:
                in_detail = True
                current = None
            elif in_detail and style == "Heading 3":
                current = []
                sections.append((p.text.strip(), current))
            else:
                in_detail = False
                current = None
            continue
        if current is not None:
            current.append(p)

    summary = _parse_summary_rows(_find_summary_table(doc))
    records: list[dict] = []
    for idx, (raw_title, paras) in enumerate(sections):
        title, fixed = _normalize_vul_title(raw_title)
        record: dict = {
            "title": title,
            "level": 30,
            "vul_type": 75,
            "affected_url": "",
            "description_html": "",
            "reproduce_html": "",
            "solution_html": "",
            "retest_html": "",
            "fixed": fixed,
            "errors": [],
        }
        bucket: str | None = None
        for p in paras:
            norm = _norm_label(p.text)
            if norm in _SECTION_LABELS:
                bucket = _SECTION_LABELS[norm]
                continue
            if _RETEST_LABEL_RE.match(norm):
                bucket = "retest_html"
                continue
            if bucket in ("description_html", "reproduce_html", "solution_html", "retest_html"):
                inner = "".join(_run_to_html(r, part, img_dir, image_url_prefix) for r in p.runs)
                if inner.strip():
                    record[bucket] += f"<p>{inner}</p>"
            elif bucket == "affected_url" and not record["affected_url"] and p.text.strip():
                record["affected_url"] = p.text.strip()[:512]
            # status/level 桶为模板选项占位行（如【超危】【高危】…），以汇总表为准

        # 等级/类型/修复状态优先取汇总表：数量一致按序配对，否则按标题匹配
        row = summary[idx] if len(summary) == len(sections) else _match_summary(summary, title)
        if row is not None:
            level = _map_level_report(row["level_text"])
            if level is None:
                record["errors"].append(f"无法识别漏洞等级「{row['level_text']}」，已按中危处理")
            else:
                record["level"] = level
            record["vul_type"] = _fuzzy_type(row["type_text"])
            record["fixed"] = record["fixed"] or row["fixed"]
        else:
            record["errors"].append("未在「风险问题汇总」表中匹配到该漏洞，等级已按中危处理")
        records.append(record)

    return meta, records


def parse_any_docx(file_path: str, image_dir: str, image_url_prefix: str,
                   filename: str = "") -> tuple[str, dict | None, list[dict]]:
    """自动识别文档格式并解析。

    返回 (doc_kind, meta, records)：报告格式为 ("report", meta, records)，
    否则按固定模板解析为 ("template", None, records)。
    """
    doc = Document(file_path)
    if is_report_docx(doc):
        meta, records = parse_report_docx(file_path, image_dir, image_url_prefix, filename)
        return "report", meta, records
    return "template", None, parse_docx(file_path, image_dir, image_url_prefix)


def build_import_template() -> Document:
    """生成标准导入模板文档（含填写示例），字段与报告模板「风险问题详情」章节一致。"""
    doc = Document()
    doc.add_heading("漏洞导入模板", level=0)
    doc.add_paragraph("填写说明：每个漏洞使用一张下方格式的表格，可复制多份；"
                      "表格以外的内容不会被导入。漏洞等级取值：严重/高危/中危/低危。"
                      "字段名称与导出报告的「风险问题详情」章节一致，导入后可直接生成报告。")

    rows = [
        ("漏洞名称", "示例：后台登录接口存在SQL注入"),
        ("漏洞等级", "高危"),
        ("漏洞类型", "SQL注入漏洞"),
        ("漏洞链接", "https://example.com/api/login"),
        ("漏洞描述", "在此填写漏洞描述，可包含多段文字与截图。"),
        ("漏洞证明", "1. 打开登录页\n2. 在用户名输入 ' or 1=1--\n3. 观察返回（可粘贴截图作为证明）"),
        ("修复建议", "使用参数化查询，过滤特殊字符。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    return doc
