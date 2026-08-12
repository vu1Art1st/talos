"""报告 Word 构建：以渗透测试报告模板为基底，填充封面 / 版本记录 / 测试目标 / 汇总统计 / 漏洞详情。

模板锚点依赖 backend/app/templates/report_template.docx 的固定结构：
表0 封面装饰 | 表1 版本变更记录 | 表2 适用性声明 | 表3 目录(TOC 域)
表4 测试目标 | 表5 时间与人员 | 表6 风险问题汇总
"""
import copy
import html as _html_mod
import re
from datetime import datetime
from pathlib import Path
from urllib.parse import unquote, urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Length, Pt, RGBColor
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
from htmldocx import HtmlToDocx
from pygments import lex
from pygments.lexers import get_lexer_by_name, guess_lexer
from pygments.lexers.special import TextLexer
from pygments.token import Token

from app.constants import VUL_LEVEL_EXPORT, VUL_STATUS, VUL_TYPE, VulStatus
from app.core.config import settings
from app.core.timeutil import now as tznow  # 系统本地时间（UTC+8）；别名避免遮蔽模块内局部变量 now

_STORAGE_SRC = re.compile(r'src="/storage/([^"]+)"')
_IMG_SRC_RE = re.compile(r'<img\b[^>]*\bsrc="([^"]+)"', re.IGNORECASE)
_HEADING_OPEN = re.compile(r"<h[1-6][^>]*>", re.IGNORECASE)
_HEADING_CLOSE = re.compile(r"</h[1-6]>", re.IGNORECASE)
# 章节快照中「测试状态：」后的初测/复测标记，导出时按漏洞最新 is_retest 重写
_TEST_STATE = re.compile(r"(测试状态：</strong>)\s*(?:初测|复测)")
# 代码块：<pre>/<code> 中的换行转 <br/>、缩进空格转 &nbsp;，保证 Word 中保留格式。
# 同时匹配已转义的 &lt;code&gt;/&lt;pre&gt;（富文本粘贴/Word 导入等场景会把标签存成实体），
# 避免 Word 中直接显示 `<code>21</code>` 原始标签。
_REAL_PRE_BLOCK = re.compile(r"<pre[^>]*>(.*?)</pre>", re.DOTALL | re.IGNORECASE)
_REAL_CODE_BLOCK = re.compile(r"<code[^>]*>(.*?)</code>", re.DOTALL | re.IGNORECASE)
_ESC_PRE_BLOCK = re.compile(r"&lt;pre[^&]*&gt;(.*?)&lt;/pre&gt;", re.DOTALL | re.IGNORECASE)
_ESC_CODE_BLOCK = re.compile(r"&lt;code[^&]*&gt;(.*?)&lt;/code&gt;", re.DOTALL | re.IGNORECASE)
# 超链接标签剥离：仅保留链接文字（需求13：URL 以普通文本展示）
_LINK_OPEN = re.compile(r"<a\s+[^>]*>", re.IGNORECASE)
_LINK_CLOSE = re.compile(r"</a>", re.IGNORECASE)

# 汇总表「问题等级」字体颜色，与模板统计段落（超危/高危/中危/低危漏洞N个）保持一致
_LEVEL_COLORS = {
    "超危": RGBColor(0xC0, 0x00, 0x00),
    "高危": RGBColor(0xFF, 0x00, 0x00),
    "中危": RGBColor(0xFF, 0xC0, 0x00),
    "低危": RGBColor(0x00, 0x70, 0xC0),
}

# 汇总表「问题状态」已修复漏洞的绿色
FIXED_STATUS_COLOR = RGBColor(0x00, 0xB0, 0x50)

# 模板中用于定位替换的封面/统计段落文案
_TPL_COVER_TITLE = "标准名称（邮件、台账）-系统名称（网页）"
_TPL_COVER_DATE = "2026年xx月xx日"
_TPL_COMPANY = "中移系统集成有限公司"


def _localize_images(html: str) -> str:
    """把 /storage/xx 图片 URL 替换为本地文件路径，供 htmldocx 内嵌图片。

    兼容多种 src 形态：单双引号、属性实体转义（&quot;/&amp;）、URL 编码字符（%20 等）、
    绝对 URL 与协议相对（//）前缀。仅接受解析后仍位于 storage 根目录内的路径，
    防止 ../ 路径遍历读取任意文件；文件不存在时按文件名在 storage 下搜索兜底；
    仍无法解析时保持原样，由 _drop_unresolvable_images 后续移除，杜绝链接式占位导出。"""
    base = settings.storage_path.resolve()

    def repl(m: re.Match) -> str:
        tag = m.group(0)
        quote = m.group(1)
        orig = m.group(2)
        # 注意：本正则 (['\"'])(.*?)\1 中 group(1) 是引号、group(2) 才是 src 值，
        # 取错分组会把路径变成引号导致所有图片无法本地化（曾在 v1.1.0 引发图片全部丢失）
        raw = _html_mod.unescape(orig)  # 仅还原 src 值内的属性实体，不做整段 unescape
        path = unquote(raw)
        # 剥掉绝对 URL / 协议相对前缀，归一化为 /storage/... 相对路径
        for prefix in ("/storage/", "//", "http://", "https://"):
            idx = path.find(prefix)
            if idx >= 0:
                path = path[idx:]
                break
        rel = re.sub(r"^/", "", path)
        if not rel or ".." in rel:
            return tag

        def local_src() -> str | None:
            candidate = (settings.storage_path / rel).resolve()
            if candidate.is_relative_to(base) and candidate.is_file():
                return candidate.as_posix()
            # 文件不存在：按文件名在 storage 下搜索兜底（规避 URL 前缀差异导致的漏匹配）
            name = Path(rel).name
            if name and base.is_dir():
                hit = next(base.rglob(name), None)
                if hit is not None:
                    return hit.as_posix()
            return None

        local = local_src()
        if local is None:
            return tag
        # 只替换 src 属性值、保留 <img> 标签其余部分：
        # 若直接返回 src="..." 会丢 <img 与 > 使标签残缺（v1.1.0 曾因此图片全部丢失）
        return tag.replace(f"src={quote}{orig}{quote}", f'src="{local}"')

    return re.sub(
        r"<img\b[^>]*?\bsrc\s*=\s*(['\"])(.*?)\1[^>]*>",
        repl,
        html or "",
        flags=re.IGNORECASE | re.DOTALL,
    )


def _compress_image_file(src: Path) -> str:
    """用 Pillow 压缩单张图片，从源头减小 docx 体积：
    - 最长边超过 REPORT_IMAGE_MAX_PX 时等比重采样（仅降不升，小图不放大）；
    - 含透明通道 → 保存 PNG optimize（仅优化编码，保留透明）；
    - 无透明通道 → 转 RGB 存 JPEG（quality=REPORT_IMAGE_QUALITY, optimize）。
    产物写入系统临时目录（随机文件名），不污染 storage；docx 内嵌后即可丢弃。"""
    import tempfile
    import uuid

    from PIL import Image

    max_px = settings.REPORT_IMAGE_MAX_PX
    with Image.open(src) as img:
        w, h = img.size
        longest = max(w, h)
        if longest > max_px:
            ratio = max_px / longest
            img = img.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
        has_alpha = img.mode in ("RGBA", "LA") or (
            img.mode == "P" and "transparency" in img.info
        )
        tmp = Path(tempfile.gettempdir()) / f"talos_img_{uuid.uuid4().hex}"
        if has_alpha:
            out = tmp.with_suffix(".png")
            img.save(out, "PNG", optimize=True)
            return str(out)
        if img.mode != "RGB":
            img = img.convert("RGB")
        out = tmp.with_suffix(".jpg")
        img.save(out, "JPEG", quality=settings.REPORT_IMAGE_QUALITY, optimize=True, progressive=True)
        return str(out)


def _compress_images(html: str) -> str:
    """压缩 HTML 中 <img src="本地文件"> 指向的图片，替换 src 为压缩后文件。

    仅处理本地存在的文件（http 等远程 / 越界路径保持原样）；
    Pillow 不可用或压缩失败时保持原图，不中断导出。"""
    def repl(m: re.Match) -> str:
        tag = m.group(0)
        src = m.group(1)
        if not src or not Path(src).exists():
            return tag
        try:
            new_src = _compress_image_file(Path(src))
        except Exception:
            return tag
        if new_src == src:
            return tag
        return tag.replace(f'src="{src}"', f'src="{new_src}"')

    return _IMG_SRC_RE.sub(repl, html or "")


# img 标签匹配（src 单双引号均可，仅用于缺失判定，取值含引号内完整 src）
_IMG_SRC_ANY_QUOTE = re.compile(
    r"<img\b[^>]*?\bsrc\s*=\s*(['\"])(.*?)\1", re.IGNORECASE | re.DOTALL
)


def _drop_unresolvable_images(html: str) -> str:
    """移除 src 既非本地文件也非 http(s) URL 的 img 标签，杜绝链接式占位导出。

    htmldocx 对无法读取的本地图片会输出 `<image: 文件名>` 占位段落（即「图片以链接形式导出」）。
    此处先过滤掉本地化/压缩后仍解析失败的 img；http(s) 远程图片保留交由 htmldocx 下载，
    下载失败产生的占位由 build_report_docx 末尾的安全网统一清理。"""
    def repl(m: re.Match) -> str:
        src = m.group(2)
        if src.startswith(("http://", "https://")):
            return m.group(0)
        if Path(src).exists():
            return m.group(0)
        return ""

    return _IMG_SRC_ANY_QUOTE.sub(repl, html or "")


def _demote_headings(html: str) -> str:
    """正文富文本中的 h1-h6 降级为加粗段落，确保目录层级仅由模板章节标题构成
    （一级：二、测试结果综述 / 二级：2.1、风险问题汇总 / 三级：2.2.1、漏洞标题）。"""
    html = _HEADING_OPEN.sub("<p><strong>", html or "")
    return _HEADING_CLOSE.sub("</strong></p>", html)


# ---------- 代码块：pygments 语法高亮 + 等宽样式（导出到 Word 保持格式） ----------

# 代码块等宽字体；eastAsia 用于代码中的中文注释，避免中文字符走默认字体
_CODE_FONT = "Consolas"
_CODE_EAST_FONT = "Microsoft YaHei"
# 代码块底纹与边框（GitHub 风格浅灰）
_CODE_BG = "F6F8FA"
_CODE_BORDER = "D0D7DE"
# 高亮配色（GitHub 风格）：(父 token 类型, 十六进制颜色, 是否斜体)
_CODE_TOKEN_STYLE = [
    (Token.Comment, "6A737D", True),          # 注释：灰 + 斜体
    (Token.Literal.String, "032F62", False),  # 字符串：深蓝
    (Token.Literal.Number, "005CC5", False),  # 数字：蓝
    (Token.Keyword, "D73A49", False),         # 关键字：红
    (Token.Name.Function, "6F42C1", False),   # 函数名：紫
    (Token.Name.Class, "6F42C1", False),      # 类名：紫
    (Token.Name.Builtin, "005CC5", False),    # 内置名：蓝
    (Token.Name.Namespace, "6F42C1", False),
    (Token.Name.Constant, "005CC5", False),
    (Token.Name.Decorator, "D73A49", False),  # 装饰器：红
    (Token.Operator, "D73A49", False),        # 运算符：红
    (Token.Generic, "24292E", False),
]
_CODE_DEFAULT_COLOR = "24292E"

_LANG_CLASS = re.compile(
    r'class\s*=\s*["\']?(?:language|lang|brush)[:-]\s*([A-Za-z0-9_+#-]+)', re.IGNORECASE
)


def _extract_lang(tag: str) -> str:
    """从 <pre>/<code> 标签 class 中提取语言名，如 language-python → python。"""
    m = _LANG_CLASS.search(tag)
    return m.group(1).lower() if m else ""


def _token_style(ttype) -> tuple[str, bool]:
    """按 pygments token 类型返回 (颜色, 是否斜体)，未匹配则用默认黑色。"""
    for parent, color, italic in _CODE_TOKEN_STYLE:
        if ttype in parent:
            return color, italic
    return _CODE_DEFAULT_COLOR, False


def _resolve_lexer(code: str, lang: str):
    """解析代码块语言：
    - 显式语言标记优先；否则代码足够长（>60 字符）时用 guess_lexer 猜测，
      避免把 `' or 1=1--` 之类的短片段误判；均失败时退回纯文本 lexer。
    """
    if lang:
        try:
            return get_lexer_by_name(lang, stripnl=False)
        except Exception:
            pass
    if len(code.strip()) > 60:
        try:
            return guess_lexer(code, stripnl=False)
        except Exception:
            pass
    return TextLexer(stripnl=False)


def _split_token_lines(tokens: list[tuple]) -> list[list[tuple]]:
    """把 pygments token 流按换行拆分为多行。"""
    lines: list[list[tuple]] = [[]]
    for ttype, value in tokens:
        while True:
            idx = value.find("\n")
            if idx < 0:
                if value:
                    lines[-1].append((ttype, value))
                break
            head, value = value[:idx], value[idx + 1:]
            if head:
                lines[-1].append((ttype, head))
            lines.append([])
    return lines


def _style_code_run(run, ttype) -> None:
    """代码 run：等宽字体 + 按 token 类型着色。"""
    color, italic = _token_style(ttype)
    run.font.name = _CODE_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), _CODE_EAST_FONT)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)
    run.italic = italic


def _code_block_indents(doc: Document) -> tuple[Length, Length]:
    """计算代码块左右缩进，使代码块显示宽度统一为 REPORT_IMAGE_WIDTH_CM（默认 14cm）且水平居中。

    以正文所在 section 的内容区宽度为基准：内容区宽于目标宽度时左右缩进均分，
    宽度不足时退化为占满内容区（缩进 0）。全程以整数 EMU 计算避免舍入误差。"""
    section = doc.sections[-1]
    content_w = (
        int(section.page_width or 0)
        - int(section.left_margin or 0)
        - int(section.right_margin or 0)
    )
    target = int(Cm(settings.REPORT_IMAGE_WIDTH_CM))
    if content_w <= target:
        return Cm(0), Cm(0)
    each = (content_w - target) // 2  # EMU 整数
    return Length(each), Length(each)


def _style_code_para(p: Paragraph, left: Length, right: Length) -> None:
    """代码块段落：浅灰底纹 + 单线边框 + 紧凑行距，宽度固定（默认 14cm）。

    块体水平居中靠对称缩进实现（左右缩进均分），段落保持默认左对齐，
    代码内容文本不随块体居中。"""
    pf = p.paragraph_format
    pf.left_indent = left
    pf.right_indent = right
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "3")
        b.set(qn("w:color"), _CODE_BORDER)
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _CODE_BG)
    pPr.append(shd)


def _add_code_block(doc: Document, code: str, lang: str) -> None:
    """把代码渲染为一个带边框底纹的段落：行内换行用 w:br，token 逐个着色。"""
    lexer = _resolve_lexer(code, lang)
    try:
        tokens = list(lex(code, lexer))
    except Exception:
        tokens = [(Token.Text, code)]
    lines = _split_token_lines(tokens)
    left, right = _code_block_indents(doc)
    p = doc.add_paragraph()
    _style_code_para(p, left, right)
    first = True
    for line in lines:
        if not first:
            p.add_run().add_break()
        first = False
        for ttype, value in line:
            run = p.add_run(value)
            _style_code_run(run, ttype)


def _esc_code_to_real(m: re.Match) -> str:
    """转义实体 &lt;code&gt;：多行还原为 <pre> 代码块，单行还原为真 <code> 标签。"""
    content = m.group(1)
    if "\n" in content or "<br" in content.lower():
        return f"<pre>{content}</pre>"
    return f"<code>{content}</code>"


def _strip_code_wrap(inner: str) -> str:
    """去掉 <pre> 内容外层可能存在的 <code> 包裹标签。"""
    inner = re.sub(r"^\s*<code[^>]*>", "", inner, count=1)
    inner = re.sub(r"</code>\s*$", "", inner, count=1)
    return inner.strip()


def _split_inline_codes(fragment: str) -> list[tuple[str, str, str]]:
    """把 <pre> 之外的片段按多行 <code> 再分割（提升为代码块）。

    单行行内 <code> 不分割、保留在 html 片段中（htmldocx 以等宽字体渲染），
    避免 tiptap/Word 常见 `<pre><code>...</code></pre>` 结构被二次处理。"""
    parts: list[tuple[str, str, str]] = []
    last = 0
    for m in _REAL_CODE_BLOCK.finditer(fragment):
        content = m.group(1)
        if "\n" not in content and "<br" not in content.lower():
            continue  # 单行 code 留在 html 段
        if m.start() > last:
            parts.append(("html", fragment[last:m.start()], ""))
        parts.append(("code", _html_mod.unescape(content), _extract_lang(m.group(0))))
        last = m.end()
    if not parts:
        return [("html", fragment, "")]
    if last < len(fragment):
        parts.append(("html", fragment[last:], ""))
    return parts


def _split_blocks(html: str) -> list[tuple[str, str, str]]:
    """把富文本 HTML 分割为交替的 html / code 片段。

    返回 [(kind, content, lang)]，kind ∈ {"html", "code"}：
    - 真标签 <pre>、多行 <code> 以及转义实体 &lt;pre&gt;/&lt;code&gt;（多行）识别为代码块；
    - 单行行内 <code> 保留在 html 片段中，由 htmldocx 以等宽字体渲染。
    """
    html = html or ""
    # 1) 转义实体还原（富文本粘贴 / Word 导入会把标签存成 &lt;...&gt;）
    html = _ESC_PRE_BLOCK.sub(lambda m: f"<pre>{m.group(1)}</pre>", html)
    html = _ESC_CODE_BLOCK.sub(_esc_code_to_real, html)
    # 2) 按 <pre> 分割（<pre> 内部的 <code> 在提取时剥掉，避免嵌套 pre）
    parts: list[tuple[str, str, str]] = []
    pos = 0
    for m in _REAL_PRE_BLOCK.finditer(html):
        if m.start() > pos:
            parts.extend(_split_inline_codes(html[pos:m.start()]))
        parts.append((
            "code",
            _html_mod.unescape(_strip_code_wrap(m.group(1))),
            _extract_lang(m.group(0)),
        ))
        pos = m.end()
    if pos < len(html):
        parts.extend(_split_inline_codes(html[pos:]))
    return parts or [("html", html, "")]


def _strip_links(html: str) -> str:
    """剥离 <a href> 超链接标签，仅保留链接文字（需求13：URL 以普通文本展示）。"""
    return _LINK_CLOSE.sub("", _LINK_OPEN.sub("", html or ""))


def _color_vuln_levels(html: str) -> str:
    """漏洞章节中的「漏洞等级：」文字着色，颜色与风险汇总一致（需求14）。

    使用 htmldocx 支持的 <span style="color: #rrggbb"> 内联样式（其 color 处理
    支持 #hex 形式），将「漏洞等级：高危」等文字染上与风险汇总一致的颜色。"""
    for label, color in _LEVEL_COLORS.items():
        # 仅处理「漏洞等级：」后紧跟该等级文字的片段，避免误伤正文
        html = re.sub(
            rf"(<strong>漏洞等级：</strong>)\s*{label}",
            rf'\g<1><span style="color: #{color}">{label}</span>',
            html,
        )
    return html


def _number_vuln_urls(html: str) -> str:
    """漏洞章节「漏洞链接：」内容排版（需求14/需求3）。

    - 链接一律另起一行展示（不在「漏洞链接：」标签后紧跟）；
    - 单个链接直接换行展示；多个链接自动编号并按行排列。"""
    def _repl(m: re.Match) -> str:
        urls = [u.strip() for u in re.split(r"<br\s*/?>", m.group(2)) if u.strip()]
        if not urls:
            return m.group(0)
        if len(urls) == 1:
            # 单个链接：另起一行，不带编号
            return f"{m.group(1)}<br/>{urls[0]}"
        numbered = "<br/>".join(f"{i}. {u}" for i, u in enumerate(urls, 1))
        return f"{m.group(1)}<br/>{numbered}"

    return re.sub(
        r"(<strong>漏洞链接：</strong>)(.*?)(</p>)",
        _repl,
        html or "",
        flags=re.DOTALL,
    )


def _add_html(doc: Document, html: str) -> None:
    # 图片统一宽度在 docx 层面由 _normalize_image_width 处理（htmldocx 忽略 HTML width）；
    # 先本地化 /storage 路径 → 压缩图像数据（重采样+JPEG）→ 过滤仍无法解析的 img，
    # 避免 htmldocx 对缺失图片输出 <image: 文件名> 链接占位
    html = _localize_images(_demote_headings(html or ""))
    html = _drop_unresolvable_images(_compress_images(html))
    for kind, content, lang in _split_blocks(html):
        if kind == "code":
            try:
                _add_code_block(doc, content, lang)
            except Exception:
                # 高亮失败时降级为无高亮的等宽纯文本代码块
                _add_code_block(doc, content, "")
            continue
        text = _strip_links(content)
        if not text.strip():
            continue
        parser = HtmlToDocx()
        try:
            parser.add_html_to_document(text, doc)
        except Exception:
            # 富文本转换失败时降级为纯文本，避免导出整体失败
            plain = re.sub(r"<[^>]+>", "", text)
            doc.add_paragraph(plain)


def _set_para_text(para: Paragraph, text: str) -> None:
    """替换段落文本并保留首个 run 的字体格式。"""
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _set_cell_text(cell, text: str) -> None:
    """替换单元格文本，保留原格式；多余段落清除。"""
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    _set_para_text(cell.paragraphs[0], text)


def _clone_row(table: Table, src_row: _Row) -> _Row:
    """深拷贝样例行（保留边框/字体等格式）并追加到表尾。"""
    new_tr = copy.deepcopy(src_row._tr)
    src_row._tr.getparent().append(new_tr)
    return table.rows[-1]


def _fill_cover(doc: Document, meta: dict, now: datetime) -> None:
    system_name = meta.get("project_name") or meta.get("title") or "渗透测试报告"
    company_done = False
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == _TPL_COVER_TITLE:
            # 封面第二行 = 系统名称（上一行为公司名、下两行为“渗透测试/报告”）
            _set_para_text(para, system_name)
        elif text == "渗透测试报告":
            # 首页第三行标题：复测报告自动变更为「渗透测试复测报告」
            if meta.get("is_retest"):
                _set_para_text(para, "渗透测试复测报告")
        elif text == _TPL_COVER_DATE:
            _set_para_text(para, now.strftime("%Y年%m月%d日"))
        elif not company_done and text == _TPL_COMPANY and meta.get("customer"):
            # 只替换封面首个公司名段落（不能提前 break，否则跳过其后的系统名/标题占位），
            # 统计段等由各自逻辑处理
            _set_para_text(para, meta["customer"])
            company_done = True


def _fill_applicability(doc: Document, meta: dict) -> None:
    """适用性声明（模板表2）：将系统名称占位符「xxxxx系统」替换为实际系统名称。

    模板中占位符为独立 run，直接替换即可保留其余文字与格式；若占位符跨 run
    （模板结构变化），则整段合并重写并保留首 run 格式。"""
    system_name = meta.get("project_name") or meta.get("title") or ""
    if not system_name:
        return
    para = doc.tables[2].rows[1].cells[0].paragraphs[0]
    for run in para.runs:
        if "xxxxx系统" in run.text:
            run.text = run.text.replace("xxxxx系统", system_name)
            return
    # 跨 run 兜底：拼接整段替换后重写
    full = "".join(r.text for r in para.runs)
    if "xxxxx系统" in full:
        _set_para_text(para, full.replace("xxxxx系统", system_name))


def _version_records(meta: dict, now: datetime, vulns: list[dict]) -> list[dict]:
    """按测试阶段生成版本变更记录序列：
    初测 = V1.0「初测创建」；第一轮复测 = V2.0「复测创建」；
    第二轮起：复测报告创建人与上一份相同则次版本 +1（V2.1），变更则升级主版本（V3.0）。

    版本清单（meta.report_records）取自计划实际关联的报告（每份报告含
    is_retest / creator_name / date），版本号严格与复测报告数量一致——
    手动流转、报告导入产生的无报告复测轮次不会虚增版本号。

    每条记录时间取该版对应报告自身的 date，缺失回退当前导出时间；修改人取
    发起导出报告的账号（meta.generator），无则回退报告作者。版本号递增仍按
    报告创建人比较，展示值与导出账号无关。"""
    export_date = now.strftime("%Y-%m-%d")
    report_records = meta.get("report_records") or []
    modifier = meta.get("generator") or meta.get("author", "")

    # 初测 V1.0：取报告清单中首份非复测报告的日期；无报告/全为复测时回退导出时间
    first_report = next((r for r in report_records if not r.get("is_retest")), None)
    records = [{
        "date": (first_report.get("date") if first_report else "") or export_date,
        "version": "V1.0",
        "note": "初测创建",
        "author": modifier,
    }]
    if not report_records and any(v.get("is_retest") for v in vulns):
        # 无报告清单但存在复测漏洞时，V1.0 标记为复测更新（兼容旧行为）
        records[0]["note"] = "复测更新"
    major, minor = 1, 0
    prev_creator = ""
    for idx, r in enumerate((x for x in report_records if x.get("is_retest")), start=1):
        creator = (r.get("creator_name") or "").strip()
        if idx == 1:
            major, minor = 2, 0
            note = "复测创建"
        else:
            if creator and creator == prev_creator:
                minor += 1
            else:
                major, minor = major + 1, 0
            note = "复测更新"
        records.append({
            "date": r.get("date") or export_date,
            "version": f"V{major}.{minor}",
            "note": note,
            "author": modifier,
        })
        prev_creator = creator
    return records


def _fill_version_table(doc: Document, meta: dict, vulns: list[dict], now: datetime) -> None:
    table = doc.tables[1]
    records = _version_records(meta, now, vulns)
    sample = table.rows[2]
    for i, rev in enumerate(records):
        if i < len(table.rows) - 2:
            row = table.rows[2 + i]
        else:
            row = _clone_row(table, sample)
        values = [rev["date"], rev["version"], rev["note"], "内部使用", rev["author"]]
        for cell, value in zip(row.cells, values):
            _set_cell_text(cell, value)


def _fill_target_table(doc: Document, meta: dict, assets: list[dict]) -> None:
    table = doc.tables[4]
    urls: list[str] = []
    for a in assets:
        urls.extend(u.get("url", "") for u in a.get("public_urls", []) if isinstance(u, dict))
        urls.extend(a.get("internal_urls", []))
    urls = [u for u in dict.fromkeys(urls) if u]
    domains = list(dict.fromkeys(
        h for u in urls if (h := urlparse(u if "://" in u else f"http://{u}").hostname)
    ))
    system_names = list(dict.fromkeys(a.get("name", "") for a in assets if a.get("name")))

    _set_cell_text(table.rows[0].cells[1], meta.get("project_name") or "、".join(system_names) or meta.get("title", ""))
    _set_cell_text(table.rows[1].cells[1], "\n".join(urls))
    _set_cell_text(table.rows[2].cells[1], "\n".join(domains))
    _set_cell_text(table.rows[3].cells[1], meta.get("target_ip", ""))
    # rows[4] 被测测试账号：系统无对应数据，留空由人工补充


def _fill_schedule_table(doc: Document, meta: dict) -> None:
    table = doc.tables[5]
    _set_cell_text(table.rows[1].cells[1], meta.get("test_start", ""))
    _set_cell_text(table.rows[1].cells[3], meta.get("test_end", ""))
    # 参测人员：第一行取发起导出账号姓名（meta.generator，即当前登录账号），
    # 其余取报告作者（按 、,， 切分去空白）；generator 与作者重复时不重复；
    # 总人数超过模板行数（3 行）时自动克隆新增行；generator 缺失时仅列作者
    lead_name = (meta.get("generator") or "").strip()
    authors = [n.strip() for n in re.split(r"[、,，]", meta.get("author") or "") if n.strip()]
    names = [lead_name] + [a for a in authors if a != lead_name] if lead_name else authors
    if not names:
        return
    sample = table.rows[5]
    rows = list(table.rows[4:])
    for i, name in enumerate(names):
        if i >= len(rows):
            rows.append(_clone_row(table, sample))
        _set_cell_text(rows[i].cells[0], name)
    # 清空多余模板行的参测人员名，避免残留
    for row in rows[len(names):]:
        _set_cell_text(row.cells[0], "")


def _set_summary_first_para(para: Paragraph, company: str, system_name: str) -> None:
    """2.1 风险问题汇总首段：「经本次测试，{company}{system_name}共发现：」
    仅系统名称加粗，其余文字保持模板原格式。

    模板中系统名（xxx系统）为独立 run，找到后仅替换其文本并设置加粗；
    若 run 结构不符合预期（未找到系统名 run），兜底整段重写（不保证加粗）。"""
    if not para.runs:
        para.add_run(f"经本次测试，{company}{system_name}共发现：")
        return
    para.runs[0].text = f"经本次测试，{company}"
    for run in para.runs[1:]:
        if "系统" in run.text:
            run.text = system_name
            run.bold = True
            return
    _set_para_text(para, f"经本次测试，{company}{system_name}共发现：")


def _fill_summary(doc: Document, meta: dict, vulns: list[dict]) -> None:
    """重写「风险问题汇总」统计段落与汇总表。"""
    counts = {"超危": 0, "高危": 0, "中危": 0, "低危": 0}
    for v in vulns:
        # 统计段仅计入当前未修复漏洞（状态非已修复）
        if v.get("status") == VulStatus.FIXED:
            continue
        label = VUL_LEVEL_EXPORT.get(v.get("level"))
        if label in counts:
            counts[label] += 1
    system_name = meta.get("project_name") or meta.get("title", "")
    company = meta.get("customer") or _TPL_COMPANY

    lines = list(counts.items())
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("经本次测试"):
            _set_summary_first_para(para, company, system_name)
        else:
            for i, (label, num) in enumerate(lines):
                if text.startswith(f"{label}漏洞"):
                    tail = "。" if i == len(lines) - 1 else "，"
                    _set_para_text(para, f"{label}漏洞{num}个{tail}")
                    break

    table = doc.tables[6]
    sample = table.rows[1]
    for v in vulns:
        row = _clone_row(table, sample)
        values = [
            VUL_LEVEL_EXPORT.get(v.get("level"), "-"),
            VUL_TYPE.get(v.get("vul_type"), "其他"),
            v.get("title", ""),
            VUL_STATUS.get(v.get("status"), "-"),
        ]
        for cell, value in zip(row.cells, values):
            _set_cell_text(cell, value)
        # 问题等级字体颜色对齐上方统计段落中对应等级的颜色
        color = _LEVEL_COLORS.get(values[0])
        if color is not None:
            for run in row.cells[0].paragraphs[0].runs:
                run.font.color.rgb = color
        # 已修复漏洞状态列显示为绿色
        if v.get("status") == VulStatus.FIXED:
            for run in row.cells[3].paragraphs[0].runs:
                run.font.color.rgb = FIXED_STATUS_COLOR
    if not vulns:
        # 无漏洞报告：汇总表占位字段统一用反斜杠填充
        for cell in sample.cells:
            _set_cell_text(cell, "\\")
    else:
        sample._tr.getparent().remove(sample._tr)


def _remove_sample_details(doc: Document) -> None:
    """删除模板「风险问题详情」小节下的样例漏洞内容（H3 起至文末）。"""
    body = doc.element.body
    found_h2 = False
    removing = False
    for child in list(body.iterchildren()):
        if child.tag == qn("w:sectPr"):
            continue
        if not removing and child.tag == qn("w:p"):
            para = Paragraph(child, doc)
            if not found_h2 and para.text.strip() == "风险问题详情":
                found_h2 = True
                continue
            if found_h2 and para.style.name == "Heading 3":
                removing = True
        if removing:
            body.remove(child)


def _apply_field_spacing(doc: Document, start: int) -> None:
    """风险问题详情章节正文段落统一 1.5 倍行距（章节标题、代码块除外）。

    覆盖漏洞描述/证明/修复建议/复测详情等具体内容段落，保证阅读舒适；
    代码块带边框底纹，保持自身紧凑行距不被覆盖。"""
    for para in doc.paragraphs[start:]:
        if para.style.name.startswith("Heading"):
            continue
        # 代码块段落（带 pBdr 边框）跳过，保持 1.0 紧凑行距
        pPr = para._p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:pBdr")) is not None:
            continue
        if para.text.strip():
            para.paragraph_format.line_spacing = 1.5


def _append_details(doc: Document, vulns: list[dict], sections: list[dict]) -> None:
    by_id = {v.get("id"): v for v in vulns}
    if not sections:
        doc.add_paragraph("本次报告未关联漏洞记录。")
        return
    for section in sections:
        vul = by_id.get(section.get("vul_id"))
        title = section.get("title") or "未命名章节"
        content_html = section.get("content_html", "")
        if vul is not None:
            # 测试状态：按漏洞最新 is_retest 重写快照（复测未通过也属于复测）
            state = "复测" if vul.get("is_retest") else "初测"
            content_html = _TEST_STATE.sub(rf"\g<1>{state}", content_html)
            # 漏洞等级颜色与风险汇总一致、漏洞链接逐条编号另起一行（需求14）
            content_html = _color_vuln_levels(content_html)
            content_html = _number_vuln_urls(content_html)
            # 修复中且经历过复测 = 复测未通过打回，展示层区分（状态码不变）
            status_name = VUL_STATUS.get(vul.get("status"), "-")
            if vul.get("status") == VulStatus.FIXING and vul.get("is_retest"):
                status_name = "复测未通过"
            title = f"{title}（{status_name}）"
        doc.add_paragraph(title, style="Heading 3")
        body_start = len(doc.paragraphs)
        _add_html(doc, content_html)
        # 章节快照未含复测详情时，追加漏洞最新复测内容（避免与生成时嵌入的快照重复）
        retest = (vul or {}).get("retest_html", "")
        if retest and "复测详情" not in (content_html or ""):
            _add_html(doc, f"<p><strong>复测详情：</strong></p>{retest}")
        _apply_field_spacing(doc, body_start)


def _build_toc_field(doc: Document) -> None:
    """恢复模板表3中的 TOC 域（回退方案：目录由 Word/WPS/LibreOffice 打开时刷新）。

    模板表3 结构：row0 =「目录」标题，row1 = TOC 域单元格。这里清空 row1 中
    缓存的旧目录条目，重建一个干净的 TOC 域（w:fldChar begin/instrText/separate/end）：
    - 域指令 `TOC \\o "1-3" \\h \\z \\u`：收录 1-3 级内置标题、条目带超链接、
      隐藏制表符前导符、按大纲级别收集；
    - separate 后保留一段提示占位文本，打开文档更新域后会被真实目录替换；
    - 配合 _enable_update_fields（w:updateFields=true）让 Word/WPS 打开时自动刷新。
    """
    table = doc.tables[3]
    if len(table.rows) < 2:
        return
    toc_cell = table.rows[1].cells[0]
    # 清空原 TOC 域所在单元格的全部段落（含缓存的目录条目）
    for p in list(toc_cell.paragraphs):
        p._element.getparent().remove(p._element)
    para = toc_cell.add_paragraph()
    run = para.add_run()
    r_el = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    # 未更新域前的占位提示；打开文档按 F9 / 右键目录「更新域」即替换为真实目录
    placeholder = OxmlElement("w:t")
    placeholder.set(qn("xml:space"), "preserve")
    placeholder.text = "（目录占位：打开文档后右键目录 → 更新域，或按 F9 刷新生成完整目录）"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, separate, placeholder, end):
        r_el.append(el)


def _enable_update_fields(doc: Document) -> None:
    """标记打开文档时刷新域，Word 会自动更新目录页码。

    w:updateFields 必须位于 CT_Settings 规定的早期位置（在 hdrShapeDefaults / footnotePr /
    endnotePr / compat / rsids 等之前），否则 Word 会忽略该设置导致打开不刷新目录。"""
    element = doc.settings.element
    if element.find(qn("w:updateFields")) is not None:
        return
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    # 定位首个“应排在 updateFields 之后”的元素，插入其前；均不存在时追加到末尾
    anchor_tags = (
        "w:hdrShapeDefaults", "w:footnotePr", "w:endnotePr",
        "w:compat", "w:rsids", "m:mathPr", "w:themeFontLang",
    )
    anchor = None
    for child in element:
        if child.tag in {qn(t) for t in anchor_tags}:
            anchor = child
            break
    if anchor is not None:
        anchor.addprevious(update)
    else:
        element.append(update)


def _shape_in_table(shape) -> bool:
    """判断内联图片是否位于表格内（封面 logo / 装饰图所在位置）。"""
    node = shape._inline
    while node is not None:
        if node.tag == qn("w:tbl"):
            return True
        node = node.getparent()
    return False


def _normalize_image_width(doc: Document, width_cm: float | None = None) -> None:
    """把正文内联图片宽度统一为 settings.REPORT_IMAGE_WIDTH_CM（默认 14cm）。

    - 正文图片全部等比缩放到统一宽度，保证版式一致（高分辨率截图冗余像素较多）；
    - 跳过模板封面表格内的小图（logo / 装饰图），避免放大失真；
    - 跳过宽度 < 5cm 的其它小图，作为兜底保护；
    - 恰好为目标宽度的图片跳过，避免无谓改动。"""
    target = Cm(width_cm or settings.REPORT_IMAGE_WIDTH_CM)
    for shape in doc.inline_shapes:
        w, h = shape.width, shape.height
        if not w or w < Cm(5) or w == target or _shape_in_table(shape):
            continue
        shape.width = target
        if h:
            shape.height = int(h * (int(target) / w))


def _center_body_images(doc: Document) -> None:
    """将正文内联图片所在段落设为水平居中（跳过模板封面表格内的 logo / 装饰图）。

    与 _normalize_image_width 配合：图片统一 14cm 宽度后再整段居中，
    保证图片在报告中水平居中对齐。"""
    for shape in doc.inline_shapes:
        if _shape_in_table(shape):
            continue
        # 沿 inline 节点向上找所属段落 <w:p>
        node = shape._inline
        while node is not None and node.tag != qn("w:p"):
            node = node.getparent()
        if node is None:
            continue
        Paragraph(node, doc).alignment = WD_ALIGN_PARAGRAPH.CENTER


def _remove_image_placeholders(doc: Document) -> None:
    """清理 htmldocx 对缺失/下载失败图片输出的 `<image: xxx>` 占位段落（链接式图片导出）。

    作为图片导出修复的最终安全网：无论前端过滤是否遗漏（如远程图片下载失败），
    凡正文段落以 `<image:` 开头的占位一律删除，保证导出文档不含链接式图片。"""
    for para in list(doc.paragraphs):
        if para.text.strip().startswith("<image:"):
            para._element.getparent().remove(para._element)


def build_report_docx(
    meta: dict,
    vulns: list[dict],
    sections: list[dict],
    out_path: str,
    assets: list[dict] | None = None,
) -> str:
    """基于模板生成报告 docx。

    meta: 报告元信息；vulns: 漏洞数据（含 id/level/status 等）；
    sections: [{title, content_html, vul_id}] 有序章节；assets: 关联资产聚合。
    """
    template = Path(settings.REPORT_TEMPLATE)
    if not template.exists():
        raise FileNotFoundError(f"报告模板不存在: {template}")

    doc = Document(str(template))
    now = tznow()

    _fill_cover(doc, meta, now)
    _fill_applicability(doc, meta)
    _fill_version_table(doc, meta, vulns, now)
    _fill_target_table(doc, meta, assets or [])
    _fill_schedule_table(doc, meta)
    _fill_summary(doc, meta, vulns)
    _remove_sample_details(doc)
    _append_details(doc, vulns, sections)
    _build_toc_field(doc)
    _normalize_image_width(doc)
    _center_body_images(doc)
    _remove_image_placeholders(doc)
    _enable_update_fields(doc)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return out_path
